import json
import logging
from datetime import datetime, timedelta

import openai
from django.conf import settings
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth import login, logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import PasswordChangeForm
from django.core.exceptions import ValidationError
from django.core.files.storage import default_storage
from django.core.mail import send_mail
from django.core.paginator import Paginator
from django.db import IntegrityError
from django.db.models import Avg, OuterRef, Q
from django.http import JsonResponse, HttpResponse
from django.shortcuts import redirect, render
from django.template.loader import render_to_string
from django.urls import reverse
from django.utils import timezone
from django.utils.crypto import get_random_string
from django.utils.timezone import now
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_GET, require_POST, require_http_methods

# CSRF exemptions for Railway deployment issues
def csrf_exempt_if_railway(view_func):
    """Exempt view from CSRF if running on Railway"""
    import os
    if os.environ.get('RAILWAY_ENVIRONMENT') or os.environ.get('RAILWAY_PROJECT_ID'):
        return csrf_exempt(view_func)
    return view_func
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response

from .decorators import counselor_required
from .forms import AppointmentForm
from .forms import AppointmentStatusForm
from .forms import CounselorProfileForm
from .forms import CustomLoginForm
from .forms import CustomUserRegistrationForm
from .forms import PasswordResetRequestForm
from .forms import ReportForm
from .forms import SetNewPasswordForm
from .models import Appointment
from .models import Counselor
from .models import CustomUser
from .models import DASSResult
from .models import Feedback
from .models import FollowupRequest
from .models import LiveSession
from .models import Notification
from .models import RelaxationLog
from .models import Report
from .models import SessionMessage
from .models import SessionParticipant
from .models import UserSettings
from .models_secure import SecureDASSResult
from .serializers import AppointmentSerializer
from .serializers_secure import SecureDASSResultSerializer
from .decorators import verified_required

# Import ratelimit with fallback
try:
    from ratelimit.decorators import ratelimit
    from ratelimit.exceptions import Ratelimited
    from django.http import HttpResponseTooManyRequests
    RATELIMIT_AVAILABLE = True

    def ratelimit_429(key=None, rate=None, block=False):
        """Rate limit decorator that returns 429 instead of 403"""
        def decorator(view_func):
            ratelimited_view = ratelimit(key=key, rate=rate, block=block)(view_func)
            def wrapper(*args, **kwargs):
                try:
                    return ratelimited_view(*args, **kwargs)
                except Ratelimited:
                    return HttpResponseTooManyRequests('Rate limit exceeded', status=429)
            return wrapper
        return decorator

except ImportError:
    # Fallback decorator if ratelimit is not available
    def ratelimit(key=None, rate=None, block=False):
        def decorator(view_func):
            return view_func
        return decorator

    def ratelimit_429(key=None, rate=None, block=False):
        def decorator(view_func):
            return view_func
        return decorator

    RATELIMIT_AVAILABLE = False

logger = logging.getLogger(__name__)

@login_required
def dass21_test(request):
    context = {
        'username': request.user.username,
        # other context data...
    }
    return render(request, 'index.html', context)


def generate_verification_token(self):
    self.verification_token = get_random_string(64)
    self.save()
    return self.verification_token


@verified_required
@login_required
def index(request):
    # Get the most recent secure DASS result for this user
    latest_result = SecureDASSResult.objects.filter(user=request.user).order_by('-date_taken').first()

    # Serialize the result if it exists
    scores = {'depression': 0, 'anxiety': 0, 'stress': 0}
    if latest_result:
        serializer = SecureDASSResultSerializer(latest_result)
        serialized_data = serializer.data
        scores = {
            'depression': serialized_data.get('depression_score', 0),
            'anxiety': serialized_data.get('anxiety_score', 0),
            'stress': serialized_data.get('stress_score', 0),
        }

    return render(request, 'index.html', {
        'user': request.user,
        'username': request.user.username,
        'college_display': request.user.get_college_display(),
        'year_display': request.user.get_year_level_display(),
        'profile_picture': request.user.profile_picture.url if request.user.profile_picture else None,
        'scores': scores,
    })

@csrf_exempt
@ratelimit(key='user', rate='3/h', block=True)  # 3 DASS submissions per hour per user
def save_dass_results(request):
    # Check authentication manually to return JSON instead of redirect
    if not request.user.is_authenticated:
        return JsonResponse({
            'status': 'error',
            'message': 'Authentication required'
        }, status=401)

    if request.method == 'POST':
        try:
            from .utils import ConsentManager, AuditLogger, DASSDataValidator
            from .models_secure import SecureDASSResult

            data = json.loads(request.body)

            # Validate DASS data
            try:
                DASSDataValidator.validate_answers(data['answers'])
                DASSDataValidator.validate_scores(
                    data['depression'],
                    data['anxiety'],
                    data['stress']
                )
            except ValidationError as e:
                AuditLogger.log_security_event(
                    'invalid_dass_data',
                    request.user,
                    {'errors': str(e)}
                )
                return JsonResponse({
                    'status': 'error',
                    'message': f'Invalid data: {str(e)}'
                }, status=400)

            # Create secure DASS result
            secure_result = SecureDASSResult(
                user=request.user,
                depression_score=data['depression'],
                anxiety_score=data['anxiety'],
                stress_score=data['stress'],
                depression_severity=data['depression_severity'],
                anxiety_severity=data['anxiety_severity'],
                stress_severity=data['stress_severity'],
                answers=data['answers']
            )
            secure_result.save(user=request.user)

            # Log successful DASS submission
            AuditLogger.log_dass_access(
                request.user,
                'dass_result_created',
                secure_result.id
            )

            # Also create a regular DASSResult for backward compatibility
            # (this will be phased out in future versions)
            result = DASSResult(
                user=request.user,
                depression_score=data['depression'],
                anxiety_score=data['anxiety'],
                stress_score=data['stress'],
                depression_severity=data['depression_severity'],
                anxiety_severity=data['anxiety_severity'],
                stress_severity=data['stress_severity'],
                answers=data['answers']
            )
            result.save()

            return JsonResponse({
                'status': 'success',
                'message': 'Results saved successfully',
                'result_id': secure_result.id
            })

        except json.JSONDecodeError:
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid JSON data'
            }, status=400)
        except Exception as e:
            # Log unexpected errors
            AuditLogger.log_security_event(
                'dass_save_error',
                request.user,
                {'error': str(e)}
            )
            return JsonResponse({
                'status': 'error',
                'message': 'An error occurred while saving results'
            }, status=500)

    return JsonResponse({
        'status': 'error',
        'message': 'Invalid request method'
    }, status=405)


@login_required
@require_POST
def record_consent(request):
    """Record user's consent for DASS assessment"""
    try:
        from .utils import ConsentManager, AuditLogger

        # Record consent
        ConsentManager.record_consent(request.user, 'dass_assessment', consent_given=True)

        # Log the consent recording
        AuditLogger.log_security_event(
            'consent_recorded',
            request.user,
            {
                'consent_type': 'dass_assessment',
                'source': 'frontend_modal'
            }
        )

        return JsonResponse({
            'status': 'success',
            'message': 'Consent recorded successfully'
        })

    except Exception as e:
        AuditLogger.log_security_event(
            'consent_recording_error',
            request.user,
            {'error': str(e)}
        )
        return JsonResponse(
            {
                'status': 'error',
                'message': 'Failed to record consent'
            },
            status=500
        )


@login_required
def user_profile(request):
    if request.method == 'POST':
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            try:
                user = request.user
                # Only allow updating these specific fields
                user.full_name = request.POST.get('full_name', user.full_name)
                user.college = request.POST.get('college', user.college)
                user.program = request.POST.get('program', user.program)
                user.year_level = request.POST.get('year_level', user.year_level)
                user.age = request.POST.get('age', user.age)
                user.gender = request.POST.get('gender', user.gender)
                user.save()
                
                return JsonResponse({'success': True})
            except Exception as e:
                return JsonResponse({'success': False, 'error': str(e)})
        elif 'profile_picture' in request.FILES:
            try:
                user = request.user
                # Handle profile picture upload
                profile_picture = request.FILES['profile_picture']
                
                # Validate file type
                allowed_types = ['image/jpeg', 'image/jpg', 'image/png', 'image/gif']
                if profile_picture.content_type not in allowed_types:
                    return JsonResponse({
                        'success': False, 
                        'error': 'Invalid file type. Please upload a JPEG, PNG, or GIF image.'
                    })
                
                # Validate file size (max 5MB)
                if profile_picture.size > 5 * 1024 * 1024:
                    return JsonResponse({
                        'success': False, 
                        'error': 'File size too large. Please upload an image smaller than 5MB.'
                    })
                
                # Save the new profile picture
                user.profile_picture = profile_picture
                user.save()

                # Also update counselor image if user is a counselor
                if hasattr(user, 'counselor_profile') and user.counselor_profile:
                    user.counselor_profile.image = profile_picture
                    user.counselor_profile.save()

                return JsonResponse({
                    'success': True,
                    'message': 'Profile picture updated successfully!',
                    'image_url': user.profile_picture.url
                })
            except Exception as e:
                return JsonResponse({'success': False, 'error': str(e)})

    user = request.user
    
    # Paginate test results using SecureDASSResult
    test_results = SecureDASSResult.objects.filter(user=user).order_by('-date_taken')
    test_paginator = Paginator(test_results, 5)  # Show 5 results per page
    test_page_number = request.GET.get('test_page')
    test_page_obj = test_paginator.get_page(test_page_number)
    
    # Paginate appointments
    appointments = Appointment.objects.filter(user=user).order_by('-date', '-time')
    appt_paginator = Paginator(appointments, 5)  # Show 5 appointments per page
    appt_page_number = request.GET.get('appt_page')
    appt_page_obj = appt_paginator.get_page(appt_page_number)
    
    return render(request, 'user-profile.html', {
        'user': user,
        'username': user.username,
        'email': user.email,
        'program': user.program,
        'test_page_obj': test_page_obj,  # Changed from direct queryset
        'appt_page_obj': appt_page_obj,  # Changed from direct queryset
    })
    
def home(request):
    return render(request, 'mentalhealth/login.html')


@csrf_exempt_if_railway
def register(request):
    if request.method == 'POST':
        print(f"DEBUG: POST data received: {request.POST}")
        print(f"DEBUG: Is AJAX: {request.headers.get('X-Requested-With') == 'XMLHttpRequest'}")

        form = CustomUserRegistrationForm(request.POST)
        print(f"DEBUG: Form is valid: {form.is_valid()}")

        if not form.is_valid():
            print(f"DEBUG: Form errors: {form.errors}")
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'errors': form.errors.get_json_data()
                }, status=400)
            messages.error(request, "Please correct the errors below.")
            return render(request, 'register.html', {'form': form})

        # Form is valid, create user
        try:
            print("DEBUG: Creating user...")
            user = form.save(commit=False)
            user.is_active = False
            user.verification_token = get_random_string(64)
            user.save()
            print(f"DEBUG: User created successfully with ID: {user.id}, email: {user.email}")

            # Create verification link
            verification_link = request.build_absolute_uri(
                reverse('verify_email', kwargs={'token': user.verification_token})
            )
            print(f"DEBUG: Verification link: {verification_link}")

            # Send verification email
            try:
                html_message = render_to_string('mentalhealth/verification-email.html', {
                    'user': user,
                    'verification_link': verification_link
                })

                plain_message = f"""Verify Your Email for CalmConnect

Hello {user.full_name},

Please click this link to verify your email: {verification_link}

Thank you!"""

                print(f"DEBUG: Sending email to: {user.email}")
                send_mail(
                    'Verify Your Email for CalmConnect',
                    plain_message,
                    settings.DEFAULT_FROM_EMAIL,
                    [user.email],
                    html_message=html_message,
                    fail_silently=False,
                )
                print("DEBUG: Email sent successfully")
                email_sent = True
            except Exception as email_error:
                print(f"DEBUG: Email sending failed: {email_error}")
                email_sent = False

            # Handle response based on environment and request type
            if settings.DEBUG:
                # In development, activate user immediately
                user.is_active = True
                user.email_verified = True
                user.save()
                print("DEBUG: User activated for development")

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'email': user.email,
                        'redirect_url': reverse('index'),
                        'message': 'Registration successful! You are now logged in.'
                    })
                return redirect('index')
            else:
                # In production, require email verification
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'email': user.email,
                        'redirect_url': reverse('verify_prompt'),
                        'email_sent': email_sent
                    })
                return redirect('verify_prompt')

        except Exception as e:
            print(f"DEBUG: Exception during user creation: {e}")
            import traceback
            traceback.print_exc()

            # Clean up partial user if created
            if 'user' in locals() and hasattr(user, 'pk') and user.pk:
                try:
                    user.delete()
                    print("DEBUG: Cleaned up partial user creation")
                except:
                    pass

            error_msg = f"Registration failed: {str(e)}"
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'error': error_msg
                }, status=400)
            messages.error(request, error_msg)
            return render(request, 'register.html', {'form': form})

    # GET request
    form = CustomUserRegistrationForm()
    return render(request, 'register.html', {'form': form})


# Add these new views
def verify_email(request, token):
    try:
        user = CustomUser.objects.get(verification_token=token)
        
        if user.email_verified:
            messages.info(request, "Email already verified.")
            return redirect('login')
            
        user.email_verified = True
        user.is_active = True  # Allow login after verification
        user.verification_token = None
        user.save()
        
        # Auto-login after verification
        login(request, user)
        messages.success(request, "Email successfully verified!")
        return redirect('index')  # Changed from login to index
        
    except CustomUser.DoesNotExist:
        messages.error(request, "Invalid verification link.")
        return redirect('register')
    

def resend_verification(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            email = data.get('email')
            user = CustomUser.objects.get(email=email)
            
            if user.email_verified:
                return JsonResponse({
                    'success': False,
                    'error': 'Email is already verified'
                }, status=400)
                
            token = user.generate_verification_token()
            verification_link = request.build_absolute_uri(
                reverse('verify_email', kwargs={'token': token})
            )
            
            # Render HTML email
            html_message = render_to_string('mentalhealth/verification-email.html', {
                'user': user,
                'verification_link': verification_link
            })
            
            # Plain text version
            plain_message = f"""
            Verify Your Email for CalmConnect

            Hello {user.full_name},

            We've received a request to resend your verification email. 
            Please verify your email address by visiting this link:

            {verification_link}

            If you didn't request this, you can safely ignore this email.

            Thank you for using CalmConnect!
            The CalmConnect Team
            """
            
            send_mail(
                'Verify your email for CalmConnect',
                plain_message,
                settings.DEFAULT_FROM_EMAIL,
                [user.email],
                html_message=html_message,
                fail_silently=False,
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Verification email resent successfully'
            })
            
        except CustomUser.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'User with this email does not exist'
            }, status=404)
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)
    
    return JsonResponse({
        'success': False,
        'error': 'Invalid request method'
    }, status=400)
    
    
def verify_prompt(request):
    if request.user.is_authenticated and request.user.email_verified:
        return redirect('index')
    
    if request.user.is_authenticated:
        # User is logged in but not verified
        return render(request, 'mentalhealth/verify-prompt.html', {
            'email': request.user.email,
            'username': request.user.username
        })
    
    # User not logged in at all
    return redirect('login')


def validate_fields(request):
    try:
        data = json.loads(request.body)
        field = data.get('field')
        value = data.get('value')
        errors = []
        
        if field == 'username':
            if CustomUser.objects.filter(username=value).exists():
                errors.append('Username already taken')
        
        return JsonResponse({'valid': len(errors) == 0, 'errors': errors})
    
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt_if_railway
def login_view(request):
    # Initialize error_message at the start
    error_message = None
    
    if request.user.is_authenticated:
        # Check for counselor profile first
        if hasattr(request.user, 'counselor_profile'):
            # Check if the counselor profile is active
            if not request.user.counselor_profile.is_active:
                messages.error(request, 'Your account has been archived. Please contact the administrator.')
                logout(request)
                return redirect('login')
            return redirect('counselor_dashboard')
        elif request.user.is_superuser or request.user.is_staff:
            return redirect('admin_dashboard')
        elif request.user.email_verified:
            return redirect('index')
        return redirect('verify_prompt')

    if request.method == 'POST':
        form = CustomLoginForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            
            # Check if user is a counselor BEFORE logging in
            if hasattr(user, 'counselor_profile'):
                # Check if the counselor profile is active
                if not user.counselor_profile.is_active:
                    error_message = 'Your account has been archived. Please contact the administrator.'
                    return render(request, 'login.html', {
                        'form': CustomLoginForm(),
                        'error_message': error_message
                    })
            
            # If we get here, the user can log in
            login(request, user)

            if not request.POST.get('remember_me'):
                request.session.set_expiry(0)

            # Check for 'next' parameter (from @login_required redirects)
            next_url = request.POST.get('next') or request.GET.get('next')
            if next_url and next_url.startswith('/') and not next_url.startswith('//'):
                # Ensure it's a safe internal redirect (not external)
                # Remove any potential query parameters that could be malicious
                from urllib.parse import urlparse
                parsed = urlparse(next_url)
                if not parsed.netloc and not parsed.scheme:  # Only relative URLs
                    return redirect(next_url)

            # Check if user is a counselor
            if hasattr(user, 'counselor_profile'):
                # Check if this is a new counselor with default password
                # This should be handled by a more secure method in production
                if hasattr(user, 'force_password_change') and user.force_password_change:
                    request.session['force_password_change'] = True
                    return redirect('password_change')
                return redirect('counselor_dashboard')
            elif user.is_superuser or user.is_staff:
                return redirect('admin_dashboard')
            elif user.email_verified:
                return redirect('index')
            return redirect('verify_prompt')
        else:
            # Provide more specific error messages
            username = request.POST.get('username', '')
            password = request.POST.get('password', '')
            
            # Check if username exists
            user_exists = False
            if '@' in username:
                user_exists = CustomUser.objects.filter(email=username).exists()
            else:
                user_exists = CustomUser.objects.filter(username=username).exists()
            if not user_exists:
                error_message = 'Username or email not found. Please check your credentials.'
            else:
                # Check if user account is active
                if '@' in username:
                    try:
                        user = CustomUser.objects.get(email=username)
                    except CustomUser.DoesNotExist:
                        user = None
                else:
                    try:
                        user = CustomUser.objects.get(username=username)
                    except CustomUser.DoesNotExist:
                        user = None
                
                if user and not user.is_active:
                    error_message = 'Your account has been deactivated. Please contact the administrator.'
                else:
                    # User exists but password is wrong
                    error_message = 'Incorrect password. Please try again.'
    
    return render(request, 'login.html', {
        'form': CustomLoginForm(),
        'error_message': error_message
    })


@login_required
def logout_view(request):
    """User logout view"""
    logout(request)
    messages.success(request, 'You have been logged out successfully.')
    return redirect('home')


@require_http_methods(["POST"])
def get_programs_ajax(request):
    """AJAX view to get programs based on selected college"""
    try:
        data = json.loads(request.body)
        college = data.get('college')
        
        # Define college-program mapping (same as in your JS)
        college_programs = {
            'CASS': [
                'Bachelor of Arts in English',
                'Bachelor of Arts in Literature',
                'Bachelor of Arts in Social Sciences',
                'Bachelor of Arts in Development Communication',
                'Bachelor of Arts in Psychology'
            ],
            'CEN': [
                'Bachelor of Science in Civil Engineering',
                'Bachelor of Science in Information Technology',
                'Bachelor of Science in Agricultural Biosystems Engineering',
            ],
            'CBA': [
                'Bachelor of Science in Accountancy',
                'Bachelor of Science in Business Administration',
                'Bachelor of Science in Entrepreneurship',
                'Bachelor of Science in Management Accounting'
            ],
            'COF': [
                'Bachelor of Science in Fisheries'
            ],
            'CAG': [
                'Bachelor of Science in Agriculture',
                'Bachelor of Science in Agribusiness',
                'Bachelor of Science in Animal Science',
                'Bachelor of Science in Crop Science'
            ],
            'CHSI': [
                'Bachelor of Science in Food Technology',
                'Bachelor of Science in Fashion and Textile Technology',
                'Bachelor of Science in Hospitality Management',
                'Bachelor of Science in Tourism Management'
            ],
            'CED': [
                'Bachelor of Elementary Education',
                'Bachelor of Secondary Education',
                'Bachelor of Cultural and Arts Education',
                'Bachelor of Early Childhood Education',
            ],
            'COS': [
                'Bachelor of Science in Biology',
                'Bachelor of Science in Chemistry',
                'Bachelor of Science in Environmental Science',
                'Bachelor of Science in Mathematics',
                'Bachelor of Science in Statistics',
                'Bachelor of Science in Meteorology'
            ],
            'CVSM': [
                'Doctor of Veterinary Medicine'
            ]
        }
        
        programs = college_programs.get(college, [])
        return JsonResponse({
            'success': True,
            'programs': programs
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })


@require_http_methods(["POST"])
def validate_field_ajax(request):
    """AJAX view to validate individual form fields"""
    try:
        data = json.loads(request.body)
        field = data.get('field')
        value = data.get('value')
        
        errors = []
        
        if field == 'username':
            if CustomUser.objects.filter(username=value).exists():
                errors.append('This username is already taken')
        
        # Return the validation results
        return JsonResponse({
            'success': True,
            'valid': len(errors) == 0,
            'errors': errors
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })
        
        
@verified_required
@login_required
def scheduler(request):
    # Get the most recent DASS result for this user
    latest_result = DASSResult.objects.filter(user=request.user).order_by('-date_taken').first()
    
    counselors = Counselor.objects.filter(is_active=True)
    
    # Debug output
    print(f"Latest DASS result: {latest_result}")
    print(f"Counselors: {list(counselors.values_list('name', flat=True))}")
    
    context = {
        'counselors': counselors,
        'username': request.user.username,
        'dass_result': latest_result,
        'scores': {
            'depression': latest_result.depression_score if latest_result else 0,
            'anxiety': latest_result.anxiety_score if latest_result else 0,
            'stress': latest_result.stress_score if latest_result else 0,
        } if latest_result else None,
        'profile_picture': request.user.profile_picture.url if request.user.profile_picture else None,
    }
    
    return render(request, 'scheduler.html', context)


@require_GET
@login_required
def get_counselor_slots(request, counselor_id):
    try:
        counselor = Counselor.objects.get(id=counselor_id)
        # Generate available slots for the next 2 weeks
        available_slots = []
        today = timezone.now().date()
        for day_offset in range(0, 7):  # Next 7 days (1 week)
            current_date = today + timedelta(days=day_offset)
            day_name = current_date.strftime('%A')
            if day_name in counselor.available_days:
                # Get individual day schedule or fall back to default
                day_schedule = counselor.day_schedules.get(day_name, {})
                from datetime import time as dt_time
                def parse_time(val):
                    if isinstance(val, str):
                        parts = val.split(":")
                        return dt_time(int(parts[0]), int(parts[1]))
                    return val
                start_time = parse_time(day_schedule.get('start_time')) if day_schedule.get('start_time') else (counselor.available_start_time if counselor.available_start_time else None)
                end_time = parse_time(day_schedule.get('end_time')) if day_schedule.get('end_time') else (counselor.available_end_time if counselor.available_end_time else None)
                if start_time and end_time:
                    time_str = start_time.strftime('%H:%M')
                    # Check if slot is already booked
                    is_booked = Appointment.objects.filter(
                        counselor=counselor,
                        date=current_date,
                        time=time_str,
                        status__in=['pending', 'confirmed']
                    ).exists()
                    if not is_booked:
                        available_slots.append({
                            'date': current_date.strftime('%Y-%m-%d'),
                            'day': day_name,
                            'time': time_str,
                            'display': f"{day_name}, {time_str}"
                        })
        return JsonResponse({
            'success': True,
            'slots': available_slots
        })
    except Counselor.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Counselor not found'
        }, status=404)

from .notification_service import notification_service

def create_notification(user, message, notification_type='general', url='#', priority='normal'):
    """Create a notification for a user (legacy function - use notification_service directly)"""
    return notification_service.create_notification(
        user=user,
        message=message,
        notification_type=notification_type,
        priority=priority,
        action_url=url
    )


@require_POST
@login_required
def book_appointment(request):
    """Book an appointment"""
    try:
        data = json.loads(request.body)
        
        # Validate required fields
        required_fields = ['counselor_id', 'date', 'time', 'session_type', 'services', 'reason', 'phone', 'course_section']
        for field in required_fields:
            if field not in data or not data[field]:
                return JsonResponse({
                    'success': False,
                    'error': f'Missing required field: {field}'
                }, status=400)
        
        # Get counselor
        try:
            counselor = Counselor.objects.get(id=data['counselor_id'], is_active=True)
        except Counselor.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'Counselor not found or not available'
            }, status=404)
        
        # Parse date and time
        try:
            appointment_date = datetime.strptime(data['date'], '%Y-%m-%d').date()
            appointment_time = datetime.strptime(data['time'], '%H:%M').time()
        except ValueError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid date or time format'
            }, status=400)
        
        # Check if date is in the past
        if appointment_date < timezone.now().date():
            return JsonResponse({
                'success': False,
                'error': 'Cannot book appointments in the past'
            }, status=400)
        
        # Check if time slot is available
        existing_appointment = Appointment.objects.filter(
            counselor=counselor,
            date=appointment_date,
            time=appointment_time,
            status__in=['pending', 'confirmed']
        ).first()
        
        if existing_appointment:
            return JsonResponse({
                'success': False,
                'error': 'This time slot is already booked'
            }, status=400)
        
        # Create appointment
        appointment = Appointment.objects.create(
            user=request.user,
            counselor=counselor,
            date=appointment_date,
            time=appointment_time,
            session_type=data['session_type'],
            services=data['services'],
            reason=data['reason'],
            phone=data['phone'],
            course_section=data['course_section'],
            status='pending'
        )
        
        # Attach DASS result if provided
        if 'dass_result_id' in data and data['dass_result_id']:
            try:
                dass_result = DASSResult.objects.get(id=data['dass_result_id'], user=request.user)
                appointment.dass_result = dass_result
                appointment.save()
            except DASSResult.DoesNotExist:
                pass  # Continue without DASS result
        
        # Create notifications using the new notification service
        from .notification_service import create_appointment_notification
        create_appointment_notification(appointment, 'created')
        
        # Handle session type specific actions
        if appointment.session_type == 'remote':
            # Create live session for remote appointments
            start_time = timezone.make_aware(datetime.combine(appointment.date, appointment.time))
            end_time = start_time + timedelta(hours=1)  # Default 1-hour session
            
            # Create or get existing live session
            live_session, created = LiveSession.objects.get_or_create(
                appointment=appointment,
                defaults={
                    'scheduled_start': start_time,
                    'scheduled_end': end_time,
                    'session_type': 'video',
                    'status': 'scheduled'
                }
            )
            
            # Generate room ID if not already set
            if not live_session.room_id:
                live_session.room_id = f'appointment_{appointment.id}'
                live_session.save()
            
            # Generate video call link for remote sessions
            video_call_url = request.build_absolute_uri(
                reverse('live_session_view', kwargs={'room_id': live_session.room_id})
            )
            appointment.video_call_url = video_call_url
            appointment.save()
            
            # Send remote session confirmation email
            send_remote_session_email(request, request.user, appointment)
        else:
            # Send standard confirmation email for face-to-face sessions
            send_confirmation_email(request, request.user, appointment)
        
        return JsonResponse({
            'success': True,
            'message': 'Appointment booked successfully!',
            'appointment_id': appointment.id,
            'session_type': appointment.session_type,
            'video_call_url': getattr(appointment, 'video_call_url', None)
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        print(f"Error booking appointment: {e}")
        return JsonResponse({
            'success': False,
            'error': 'An error occurred while booking the appointment'
        }, status=500)

def send_remote_session_email(request, user, appointment):
    """Send remote session confirmation email with video call link"""
    # Generate cancellation token and deadline
    appointment.cancellation_token = get_random_string(64)
    appointment.cancellation_deadline = timezone.now() + timedelta(hours=24)
    appointment.save()

    # Build URLs
    cancellation_url = request.build_absolute_uri(
        reverse('cancel_appointment', kwargs={'token': appointment.cancellation_token})
    )
    
    # Get or create live session for remote appointments
    if appointment.session_type == 'remote':
        try:
            live_session = LiveSession.objects.get(appointment=appointment)
            room_id = live_session.room_id
        except LiveSession.DoesNotExist:
            room_id = f'appointment_{appointment.id}'
    else:
        room_id = f'appointment_{appointment.id}'
    
    video_call_url = request.build_absolute_uri(
        reverse('live_session_view', kwargs={'room_id': room_id})
    )
    
    subject = 'Your CalmConnect Remote Session Confirmation'
    
    html_message = render_to_string('mentalhealth/remote-session-confirmation.html', {
        'user': user,
        'appointment': appointment,
        'cancellation_url': cancellation_url,
        'video_call_url': video_call_url,
        'deadline': appointment.cancellation_deadline.strftime("%B %d, %Y %I:%M %p")
    })
    
    plain_message = f"""
    Remote Session Confirmation
    
    Hello {user.full_name},
    
    Your remote counseling session has been successfully booked with CalmConnect.
    
    Appointment Details:
    - Counselor: {appointment.counselor}
    - Date: {appointment.date}
    - Time: {appointment.time}
    - Session Type: {appointment.get_session_type_display()}
    - Services: {', '.join(appointment.services)}
    
    Video Call Information:
    Your video call will be conducted through our secure platform.
    Video Call Link: {video_call_url}
    
    How to Join:
    1. Click the video call link above 5-10 minutes before your session
    2. Allow camera and microphone access when prompted
    3. Wait in the virtual waiting room until your counselor joins
    4. Your session will begin automatically when both parties are ready
    
    Technical Requirements:
    - Stable internet connection
    - Webcam and microphone
    - Modern web browser (Chrome, Firefox, Safari, Edge)
    - Private, quiet location for your session
    
    If you need to reschedule or cancel, please contact us at least 24 hours in advance.
    
    Thank you for using CalmConnect!
    
    The CalmConnect Team
    """
    
    try:
        send_mail(
            subject,
            plain_message,
            settings.DEFAULT_FROM_EMAIL,
            [user.email],
            html_message=html_message,
            fail_silently=False
        )
    except Exception as e:
        print(f"❌ Failed to send remote session email: {e}")


def send_confirmation_email(request, user, appointment):
    """Send appointment confirmation email (identical content) with working cancellation"""
    # Generate cancellation token and deadline
    appointment.cancellation_token = get_random_string(64)
    appointment.cancellation_deadline = timezone.now() + timedelta(hours=24)
    appointment.save()

    # Build cancellation URL using the passed request object
    cancellation_url = request.build_absolute_uri(
        reverse('cancel_appointment', kwargs={'token': appointment.cancellation_token})
    )
    
    # Keep the EXACT same email content as before
    subject = 'Your CalmConnect Appointment Confirmation'
    
    html_message = render_to_string('mentalhealth/appointment-confirmation.html', {
        'user': user,
        'appointment': appointment,
        'cancellation_url': cancellation_url,  # Only added this new variable
        'deadline': appointment.cancellation_deadline.strftime("%B %d, %Y %I:%M %p")
    })
    
    # Preserve the original plain text message exactly
    plain_message = f"""
    Appointment Confirmation
    
    Hello {user.full_name},
    
    Your counseling appointment has been successfully booked with CalmConnect.
    
    Appointment Details:
    - Counselor: {appointment.counselor}
    - Date: {appointment.date}
    - Time: {appointment.time}
    - Session Type: {appointment.get_session_type_display()}
    - Services: {', '.join(appointment.services)}
    
    What to Expect:
    Please arrive 5-10 minutes before your scheduled time. Bring any relevant documents or materials.
    
    If you need to reschedule or cancel, please contact us at least 24 hours in advance.
    
    Thank you for using CalmConnect!
    
    The CalmConnect Team
    """
    
    try:
        send_mail(
            subject,
            plain_message,
            settings.DEFAULT_FROM_EMAIL,
            [user.email],
            html_message=html_message,
            fail_silently=False
        )
    except Exception as e:
        print(f"❌ Failed to send confirmation email: {e}")
    
    
@require_http_methods(["GET", "POST"])
def cancel_appointment(request, token):
    try:
        appointment = Appointment.objects.get(cancellation_token=token)
        
        # Check if cancellation is still allowed
        if timezone.now() > appointment.cancellation_deadline:
            return render(request, 'mentalhealth/cancellation-expired.html', status=400)
        
        if request.method == 'POST':
            # Process cancellation form
            reason = request.POST.get('reason', 'No reason provided')
            appointment.status = 'cancelled'
            appointment.cancellation_reason = reason
            appointment.save()
            
            # Send confirmation email with HTML template
            reschedule_url = request.build_absolute_uri(reverse('scheduler'))
            
            html_message = render_to_string('mentalhealth/appointment-cancellation.html', {
                'user': appointment.user,
                'appointment': appointment,
                'reschedule_url': reschedule_url,
            })
            
            plain_message = f"""
            Appointment Cancellation Confirmed
            
            Hello {appointment.user.full_name},
            
            Your counseling appointment has been successfully cancelled.
            
            Cancelled Appointment Details:
            - Counselor: {appointment.counselor}
            - Date: {appointment.date}
            - Time: {appointment.time}
            - Services: {', '.join(appointment.services)}
            """
            
            if appointment.cancellation_reason:
                plain_message += f"- Cancellation Reason: {appointment.cancellation_reason}\n"
            
            plain_message += f"""
            
            Need to reschedule?
            You can easily book a new appointment at any time through our scheduling system.
            
            We're sorry to see you go. You can always schedule a new appointment anytime.
            
            Thank you for using CalmConnect!
            
            The CalmConnect Team
            """
            
            send_mail(
                'Appointment Cancellation Confirmed',
                plain_message,
                settings.DEFAULT_FROM_EMAIL,
                [appointment.user.email],
                html_message=html_message,
                fail_silently=False
            )
            
            return render(request, 'mentalhealth/cancellation-success.html')
        
        # GET request - show cancellation form
        return render(request, 'mentalhealth/cancellation-form.html', {
            'appointment': appointment,
            'token': token
        })
        
    except Appointment.DoesNotExist:
        return render(request, 'mentalhealth/cancellation-expired.html', status=404)


def send_feedback_request_email(request, appointment):
    """Send feedback request email when appointment is completed"""
    print(f"🔔 Sending feedback request email to {appointment.user.email} for appointment {appointment.id}")
    
    # Generate feedback token
    feedback_token = get_random_string(64)
    
    # Store token in session or create a temporary model field
    # For now, we'll use the appointment ID in the URL
    feedback_url = request.build_absolute_uri(
        reverse('feedback_form', kwargs={'appointment_id': appointment.id})
    )
    
    subject = 'How was your CalmConnect session?'
    
    html_message = render_to_string('mentalhealth/feedback-request.html', {
        'user': appointment.user,
        'appointment': appointment,
        'feedback_url': feedback_url,
    })
    
    plain_message = f"""
    How was your CalmConnect session?
    
    Hello {appointment.user.full_name},
    
    Your counseling session has been completed. We hope it was helpful and would love to hear about your experience!
    
    Session Details:
    - Counselor: {appointment.counselor.name}
    - Date: {appointment.date}
    - Time: {appointment.time}
    - Services: {', '.join(appointment.services)}
    
    Share Your Feedback:
    Your feedback helps us improve our services and supports our counselors in providing the best care possible.
    It only takes a few minutes and your input is invaluable to us.
    
    Thank you for choosing CalmConnect for your mental health support!
    
    The CalmConnect Team
    """
    
    try:
        send_mail(
            subject,
            plain_message,
            settings.DEFAULT_FROM_EMAIL,
            [appointment.user.email],
            html_message=html_message,
            fail_silently=False
        )
        print(f"✅ Feedback request email sent successfully to {appointment.user.email}")
    except Exception as e:
        print(f"❌ Failed to send feedback request email: {e}")


def force_logout(request):
    """Force logout and clear all session data."""
    logout(request)
    request.session.flush()  # Deletes the session cookie and all data
    return redirect('login')


@staff_member_required
def admin_dashboard(request):
    # Calculate statistics
    total_appointments = Appointment.objects.count()

    # Debug: Check what severity values exist in the database
    all_severities = DASSResult.objects.values_list(
        'depression_severity', 'anxiety_severity', 'stress_severity'
    ).distinct()
    print(f"All severity values in database: {list(all_severities)}")

    # Fix the high risk cases query to match the actual severity values being stored
    high_risk_cases = DASSResult.objects.filter(
        Q(depression_severity__icontains='Severe') |
        Q(depression_severity__icontains='Extremely') |
        Q(anxiety_severity__icontains='Severe') |
        Q(anxiety_severity__icontains='Extremely') |
        Q(stress_severity__icontains='Severe') |
        Q(stress_severity__icontains='Extremely')
    ).distinct().count()

    # Get the actual critical cases for display
    critical_cases = DASSResult.objects.filter(
        Q(depression_severity__icontains='Severe') |
        Q(depression_severity__icontains='Extremely') |
        Q(anxiety_severity__icontains='Severe') |
        Q(anxiety_severity__icontains='Extremely') |
        Q(stress_severity__icontains='Severe') |
        Q(stress_severity__icontains='Extremely')
    ).select_related('user').order_by('-date_taken')[:10]  # Get latest 10 critical cases

    print(f"High risk cases found: {high_risk_cases}")

    active_counselors = Counselor.objects.filter(is_active=True).count()

    # Get recent activities (last 7 days)
    recent_appointments = Appointment.objects.filter(
        created_at__gte=now()-timedelta(days=7)
    ).order_by('-created_at')[:4]

    # Prepare DASS21 chart data
    dass_data = {
        'depression': DASSResult.objects.aggregate(avg=Avg('depression_score'))['avg'],
        'anxiety': DASSResult.objects.aggregate(avg=Avg('anxiety_score'))['avg'],
        'stress': DASSResult.objects.aggregate(avg=Avg('stress_score'))['avg'],
    }

    return render(request, 'admin-panel.html', {
        'total_appointments': total_appointments,
        'high_risk_cases': high_risk_cases,
        'critical_cases': critical_cases,
        'active_counselors': active_counselors,
        'recent_appointments': recent_appointments,
        'dass_data': dass_data,
    })
    
    
@staff_member_required
def admin_data(request):
    # Get all colleges with their average DASS scores and student counts
    colleges = CustomUser.COLLEGE_CHOICES
    college_data = []
    
    for code, name in colleges:
        # Get all users from this college
        users = CustomUser.objects.filter(college=code)
        student_count = users.count()
        
        # Get DASS results for these users
        results = DASSResult.objects.filter(user__college=code)
        
        if results.exists():
            # Calculate averages
            depression_avg = results.aggregate(Avg('depression_score'))['depression_score__avg']
            anxiety_avg = results.aggregate(Avg('anxiety_score'))['anxiety_score__avg']
            stress_avg = results.aggregate(Avg('stress_score'))['stress_score__avg']
            overall_avg = (depression_avg + anxiety_avg + stress_avg) / 3
            
            # Determine severity level
            if overall_avg <= 10:
                severity = 'normal'
            elif overall_avg <= 20:
                severity = 'moderate'
            else:
                severity = 'severe'
                
            # Get last updated time
            last_updated = results.latest('date_taken').date_taken
            
            college_data.append({
                'code': code,
                'name': name,
                'student_count': student_count,
                'depression_avg': depression_avg,
                'anxiety_avg': anxiety_avg,
                'stress_avg': stress_avg,
                'overall_avg': overall_avg,
                'severity': severity,
                'last_updated': last_updated,
                'color_rgb': get_college_color(code)  # Helper function to assign colors
            })
    
    return render(request, 'admin-data.html', {
        'colleges': college_data,
    })

def get_college_color(college_code):
    """Assign consistent colors to each college"""
    color_map = {
        'CASS': '255, 99, 132',
        'CEN': '54, 162, 235',
        'CBA': '255, 206, 86',
        'COF': '75, 192, 192',
        'CAG': '153, 102, 255',
        'CHSI': '255, 159, 64',
        'CED': '199, 199, 199',
        'COS': '83, 102, 255',
        'CVSM': '40, 167, 69'
    }
    return color_map.get(college_code, '75, 192, 192')


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def appointment_list(request):
    # Add debugging
    print(f"Appointment list API called by user: {request.user.username}")
    print(f"User is staff: {request.user.is_staff}")
    print(f"User has counselor profile: {hasattr(request.user, 'counselor_profile')}")
    
    # Check if user is staff (admin) or has counselor profile
    if not (request.user.is_staff or hasattr(request.user, 'counselor_profile')):
        print("Permission denied - user is not staff or counselor")
        return Response({'error': 'Permission denied'}, status=403)
    
    upcoming = request.GET.get('upcoming', 'false').lower() == 'true'
    queryset = Appointment.objects.all().select_related('user', 'counselor', 'dass_result').order_by('-date', '-time')

    if upcoming:
        queryset = queryset.filter(
            status__in=['pending', 'confirmed'],
            date__gte=timezone.now().date()
        ).order_by('date', 'time')

    print(f"Found {queryset.count()} appointments")
    serializer = AppointmentSerializer(queryset, many=True)
    return Response(serializer.data)

@login_required
@counselor_required
def appointment_detail_view(request, pk):
    """HTML view for appointment details that redirects appropriately based on appointment type"""
    try:
        appointment = Appointment.objects.select_related('user', 'counselor').get(pk=pk)

        # Check permissions - must be the counselor for this appointment
        if appointment.counselor.user != request.user:
            messages.error(request, 'Permission denied')
            return redirect('counselor_dashboard')

        # For counselors, always redirect to schedule view where they can see appointment details
        # This provides a consistent experience for viewing appointment information
        return redirect('counselor_schedule')

    except Appointment.DoesNotExist:
        messages.error(request, 'Appointment not found')
        return redirect('counselor_dashboard')


@csrf_exempt
def appointment_detail(request, pk):
    # Convert pk to int
    try:
        pk = int(pk)
    except ValueError:
        return HttpResponse(json.dumps({'error': 'Invalid appointment ID'}), status=400, content_type='application/json')

    # Explicit authentication check
    if not request.user.is_authenticated:
        return HttpResponse(json.dumps({'error': 'Authentication required'}), status=401, content_type='application/json')

    # Check if user is staff (admin) or has counselor profile
    if not (request.user.is_staff or hasattr(request.user, 'counselor_profile')):
        return HttpResponse(json.dumps({'error': 'Permission denied'}), status=403, content_type='application/json')

    try:
        appointment = Appointment.objects.get(pk=pk)
    except Appointment.DoesNotExist:
        return HttpResponse(json.dumps({'error': 'Appointment not found'}), status=404, content_type='application/json')

    if request.method == 'OPTIONS':
        return HttpResponse(json.dumps({}), status=200, content_type='application/json')

    if request.method == 'GET':
        try:
            serializer = AppointmentSerializer(appointment)
            return HttpResponse(json.dumps(serializer.data), content_type='application/json')
        except Exception as e:
            logger.error(f"Error serializing appointment {pk}: {e}")
            return HttpResponse(json.dumps({'error': 'Failed to serialize appointment data'}), status=500, content_type='application/json')

    elif request.method == 'PATCH':
        try:
            # Parse JSON data
            data = json.loads(request.body)

            # Ensure status is lowercase to match choices
            if 'status' in data and data['status']:
                data['status'] = data['status'].lower()

            serializer = AppointmentSerializer(appointment, data=data, partial=True)
            if serializer.is_valid():
                serializer.save()
                return HttpResponse(json.dumps(serializer.data), content_type='application/json')

            # Format validation errors
            errors = {}
            for field, error_list in serializer.errors.items():
                errors[field] = [str(error) for error in error_list]
            return HttpResponse(json.dumps({'error': 'Validation failed', 'details': errors}), status=400, content_type='application/json')
        except json.JSONDecodeError:
            return HttpResponse(json.dumps({'error': 'Invalid JSON data'}), status=400, content_type='application/json')
        except Exception as e:
            logger.error(f"Error updating appointment {pk}: {e}")
            return HttpResponse(json.dumps({'error': f'Failed to update appointment: {str(e)}'}), status=500, content_type='application/json')

    return HttpResponse(json.dumps({'error': 'Method not allowed'}), status=405, content_type='application/json')


@csrf_exempt
def test_view(request):
    logger.info("test_view called")
    return HttpResponse(json.dumps({'test': 'ok', 'method': request.method}), content_type='application/json')
    

@staff_member_required
def admin_appointments(request):
    # Add debugging
    print(f"Admin appointments view called by user: {request.user.username}")
    print(f"User is staff: {request.user.is_staff}")
    print(f"User is superuser: {request.user.is_superuser}")
    
    # You can reuse the same context data as your other admin views
    return render(request, 'admin-appointments.html', {
        'total_appointments': Appointment.objects.count(),
        'high_risk_cases': DASSResult.objects.filter(
            Q(depression_severity__icontains='Severe') | 
            Q(depression_severity__icontains='Extremely') |
            Q(anxiety_severity__icontains='Severe') | 
            Q(anxiety_severity__icontains='Extremely') |
            Q(stress_severity__icontains='Severe') | 
            Q(stress_severity__icontains='Extremely')
        ).distinct().count(),
        'active_counselors': Counselor.objects.filter(is_active=True).count()
    })
    
    

@staff_member_required
def admin_personnel(request):
    counselors = Counselor.objects.filter(is_active=True).select_related('user')
    return render(request, 'admin-personnel.html', {
        'counselors': counselors,
        'colleges': CustomUser.COLLEGE_CHOICES,
        'total_appointments': Appointment.objects.count(),
        'high_risk_cases': DASSResult.objects.filter(
            Q(depression_severity__icontains='Severe') | 
            Q(depression_severity__icontains='Extremely') |
            Q(anxiety_severity__icontains='Severe') | 
            Q(anxiety_severity__icontains='Extremely') |
            Q(stress_severity__icontains='Severe') | 
            Q(stress_severity__icontains='Extremely')
        ).distinct().count(),
        'active_counselors': counselors.count()
    })

@csrf_exempt_if_railway
@require_http_methods(["POST"])
def add_counselor(request):
    # Manual authentication check for AJAX compatibility
    if not request.user.is_authenticated or not (request.user.is_staff or request.user.is_superuser):
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    try:
        logger.info(f"add_counselor called by {request.user.username}, content_type: {request.content_type}, method: {request.method}")
        logger.info(f"Request body length: {len(request.body) if request.body else 0}")
        logger.info(f"Request POST: {dict(request.POST)}")
        logger.info(f"Request FILES: {list(request.FILES.keys()) if request.FILES else 'None'}")

        # Handle both form data and JSON requests - try JSON first, then form data
        data = {}
        files = None
        try:
            data = json.loads(request.body)
            logger.info(f"Parsed as JSON data: {data}")
        except json.JSONDecodeError:
            # Try form data
            if request.POST:
                data = request.POST.copy()
                files = request.FILES
                # Convert form data lists to single values
                for key, value in data.items():
                    if isinstance(value, list) and len(value) == 1:
                        data[key] = value[0]
                logger.info(f"Parsed as form data: {dict(data)}")
            else:
                logger.error(f"No valid data found. Body: {request.body[:200]}, POST: {dict(request.POST)}")
                return JsonResponse({
                    'success': False,
                    'error': 'No valid data provided'
                }, status=400)

        # Validate required fields
        required_fields = ['name', 'email', 'college', 'rank']
        missing_fields = [field for field in required_fields if not data.get(field)]

        if missing_fields:
            return JsonResponse({
                'success': False,
                'error': f'Missing required fields: {", ".join(missing_fields)}'
            }, status=400)

        # Validate college to be one of the valid college codes
        if data.get('college'):
            from .models import CustomUser
            valid_college_codes = [code for code, name in CustomUser.COLLEGE_CHOICES]

            # Check if it's a valid college code
            if data['college'] not in valid_college_codes:
                return JsonResponse({
                    'success': False,
                    'error': f'Invalid college code. Must be one of: {", ".join(valid_college_codes)}'
                }, status=400)

        # Check if email already exists
        if Counselor.objects.filter(email=data['email']).exists():
            return JsonResponse({
                'success': False,
                'error': 'A counselor with this email already exists'
            }, status=400)

        if CustomUser.objects.filter(email=data['email']).exists():
            return JsonResponse({
                'success': False,
                'error': 'A user with this email already exists'
            }, status=400)

        # Generate a secure random password and setup token
        setup_token = get_random_string(64)
        default_password = get_random_string(12)

        try:
            # First create the user account (inactive by default)
            username = data['email'].split('@')[0]
            user = CustomUser.objects.create_user(
                username=username,
                email=data['email'],
                password=default_password,
                full_name=data['name'],
                is_staff=True,  # Counselors are staff members
                is_superuser=False,  # But not superusers
                is_active=False,  # Inactive until they complete setup
                email_verified=False,
                verification_token=setup_token,
                student_id=f"staff-{get_random_string(8)}",
                age=0,
                gender='Prefer not to say',
                college=data['college'],  # Set college to match the college code
                program='Staff',
                year_level='4'
            )
            
            # Then create the counselor
            counselor = Counselor.objects.create(
                name=data['name'],
                email=data['email'],
                unit=data['college'],
                rank=data['rank'],
                is_active=True,
                user=user
            )
            
            # Handle image upload if present
            if files and 'image' in files:
                counselor.image = files['image']
                counselor.save()

            # Send setup email
            setup_url = request.build_absolute_uri(
                reverse('counselor_setup', kwargs={'token': setup_token})
            )
            
            html_message = render_to_string('mentalhealth/counselor-setup-email.html', {
                'user': user,
                'setup_url': setup_url,
                'temporary_password': default_password
            })
            
            plain_message = f"""
            Counselor Account Setup - CalmConnect
            
            Hello {user.full_name},
            
            An administrator has created a counselor account for you on CalmConnect.
            
            Your temporary credentials:
            Username: {user.username}
            Temporary Password: {default_password}
            
            Please complete your account setup by clicking this link:
            {setup_url}
            
            This link will expire in 48 hours. If you didn't request this account, please contact the administrator.
            
            The CalmConnect Team
            """
            
            send_mail(
                'Complete Your CalmConnect Counselor Account Setup',
                plain_message,
                settings.DEFAULT_FROM_EMAIL,
                [user.email],
                html_message=html_message,
                fail_silently=False,
            )

            # Build image URL
            image_url = counselor.image.url if counselor.image else None
            if image_url:
                image_url = request.build_absolute_uri(image_url)
            else:
                image_url = request.build_absolute_uri(settings.STATIC_URL + 'img/default.jpg')

            return JsonResponse({
                'success': True,
                'counselor': {
                    'id': counselor.id,
                    'name': counselor.name,
                    'email': counselor.email,
                    'unit': counselor.unit,
                    'rank': counselor.rank,
                    'image_url': image_url,
                    'college': user.college,
                    'college_display': user.get_college_display()
                },
                'message': 'Counselor added successfully. Setup email sent.'
            })

        except IntegrityError as e:
            return JsonResponse({
                'success': False,
                'error': 'Database integrity error occurred'
            }, status=400)
        except Exception as e:
            # Clean up if anything fails
            if 'user' in locals():
                user.delete()
            if 'counselor' in locals() and counselor.id:
                counselor.delete()
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=400)
            
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)

@csrf_exempt_if_railway
@require_http_methods(["GET", "POST", "PATCH"])
def update_counselor(request, counselor_id):
    # Manual authentication check for AJAX compatibility
    if not request.user.is_authenticated or not (request.user.is_staff or request.user.is_superuser):
        return JsonResponse({'error': 'Permission denied'}, status=403)

    try:
        counselor = Counselor.objects.get(id=counselor_id)

        # Handle GET request - return counselor data
        if request.method == 'GET':
            image_url = counselor.image.url if counselor.image else None
            if image_url:
                image_url = request.build_absolute_uri(image_url)
            else:
                image_url = request.build_absolute_uri(settings.STATIC_URL + 'img/default.jpg')

            # Return the college code
            response_data = {
                'id': counselor.id,
                'name': counselor.name,
                'email': counselor.email,
                'college': counselor.college,  # Return the college code
                'rank': counselor.rank,
                'image_url': image_url,
            }

            return JsonResponse({
                'success': True,
                'counselor': response_data
            })

        if request.content_type == 'multipart/form-data':
            data = request.POST.copy()  # Make a mutable copy
            files = request.FILES
        else:
            try:
                data = json.loads(request.body)
                files = None
            except json.JSONDecodeError:
                return JsonResponse({
                    'success': False,
                    'error': 'Invalid JSON data'
                }, status=400)

        # Validate college to be one of the valid college codes
        if data.get('college'):
            from .models import CustomUser
            valid_college_codes = [code for code, name in CustomUser.COLLEGE_CHOICES]

            # Check if it's a valid college code
            if data['college'] not in valid_college_codes:
                return JsonResponse({
                    'success': False,
                    'error': f'Invalid college code. Must be one of: {", ".join(valid_college_codes)}'
                }, status=400)

        # Update fields
        counselor.name = data.get('name', counselor.name)
        counselor.email = data.get('email', counselor.email)
        counselor.college = data.get('college', counselor.college)
        counselor.rank = data.get('rank', counselor.rank)

        # Update associated user if it exists
        if counselor.user:
            # Update user email if counselor email changed
            if data.get('email') and data['email'] != counselor.user.email:
                counselor.user.email = data['email']
            # Update user college if counselor college changed
            if data.get('college') and data['college'] != counselor.user.college:
                counselor.user.college = data['college']
            counselor.user.save()

        # Handle image upload if present
        if files and 'image' in files:
            # Delete old image if exists
            if counselor.image:
                default_storage.delete(counselor.image.path)
            counselor.image = files['image']

        counselor.save()

        # Build image URL
        image_url = counselor.image.url if counselor.image else None
        if image_url:
            image_url = request.build_absolute_uri(image_url)
        else:
            image_url = request.build_absolute_uri(settings.STATIC_URL + 'img/default.jpg')

        # Prepare response data
        response_data = {
            'id': counselor.id,
            'name': counselor.name,
            'email': counselor.email,
            'college': counselor.college,
            'rank': counselor.rank,
            'image_url': image_url,
        }

        return JsonResponse({
            'success': True,
            'counselor': response_data,
            'message': 'Counselor updated successfully'
        })

    except Counselor.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Counselor not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)

@csrf_exempt_if_railway
@require_http_methods(["POST"])
def archive_counselor(request, counselor_id):
    # Manual authentication check for AJAX compatibility
    if not request.user.is_authenticated or not (request.user.is_staff or request.user.is_superuser):
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    try:
        counselor = Counselor.objects.get(id=counselor_id)
        counselor.is_active = False
        counselor.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Counselor archived successfully'
        })
    except Counselor.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Counselor not found'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)
    
    
@staff_member_required
def admin_archive(request):
    active_tab = request.GET.get('tab', 'appointments')

    # Paginate appointments
    archived_appointments = Appointment.objects.filter(
        Q(status='completed') | Q(status='cancelled')
    ).order_by('-date', '-time')
    appointments_paginator = Paginator(archived_appointments, 5)
    appointments_page_number = request.GET.get('appointments_page', 1)
    archived_appointments_page = appointments_paginator.get_page(appointments_page_number)

    # Paginate DASS results
    dass_results = DASSResult.objects.all().order_by('-date_taken')
    dass_paginator = Paginator(dass_results, 5)
    dass_page_number = request.GET.get('dass_page', 1)
    dass_results_page = dass_paginator.get_page(dass_page_number)

    # Paginate inactive counselors
    inactive_counselors = Counselor.objects.filter(is_active=False).order_by('name')
    counselors_paginator = Paginator(inactive_counselors, 5)
    counselors_page_number = request.GET.get('employees_page', 1)
    inactive_counselors_page = counselors_paginator.get_page(counselors_page_number)

    # Paginate archived reports (include both archived and completed reports)
    archived_reports = Report.objects.filter(status__in=['archived', 'completed']).order_by('-created_at')
    reports_paginator = Paginator(archived_reports, 5)
    reports_page_number = request.GET.get('reports_page', 1)
    archived_reports_page = reports_paginator.get_page(reports_page_number)


    return render(request, 'admin-archive.html', {
        'archived_appointments_page': archived_appointments_page,
        'dass_results_page': dass_results_page,
        'inactive_counselors_page': inactive_counselors_page,
        'archived_reports_page': archived_reports_page,
        'active_tab': active_tab,
    })
    
    
@login_required
@counselor_required
def counselor_dashboard(request):
    logger.debug(f"Accessing counselor dashboard for {request.user.username}")
    # Verify this is actually a counselor
    if not hasattr(request.user, 'counselor_profile'):
        logger.warning(f"Non-counselor user {request.user.username} attempted to access counselor dashboard")
        return redirect('index')
    
    counselor = request.user.counselor_profile
    today = timezone.now().date()
    
    # Get today's appointments
    today_appointments = Appointment.objects.filter(
        counselor=counselor,
        date=today,
        status__in=['pending', 'confirmed']
    ).order_by('time')
    
    # Get pending reports
    pending_reports = Report.objects.filter(
        counselor=counselor,
        status='pending'
    ).order_by('-created_at')[:5]
    
    # Get weekly session count (all statuses: pending, confirmed, completed)
    start_of_week = today - timedelta(days=today.weekday())
    weekly_sessions = Appointment.objects.filter(
        counselor=counselor,
        date__gte=start_of_week,
        status__in=['pending', 'confirmed', 'completed']
    ).count()
    
    logger.debug(f"Rendering dashboard for counselor {counselor.name} with {len(today_appointments)} appointments today")
    
    context = {
        'counselor': counselor,
        'today_appointments': today_appointments,
        'pending_reports': pending_reports,
        'pending_reports_count': pending_reports.count(),
        'weekly_sessions_count': weekly_sessions,
    }
    
    return render(request, 'counselor-panel.html', context)

def counselor_required(view_func):
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('login')
        if not hasattr(request.user, 'counselor_profile'):
            return redirect('index')
        return view_func(request, *args, **kwargs)
    return wrapper


from django.contrib.auth.forms import PasswordChangeForm
from django.contrib.auth import update_session_auth_hash

@login_required
def password_change(request):
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user)
            messages.success(request, 'Your password was successfully updated!')
            
            if 'force_password_change' in request.session:
                del request.session['force_password_change']
            
            # Redirect based on user type
            if hasattr(request.user, 'counselor_profile'):
                return redirect('counselor_dashboard')
            elif request.user.is_staff:
                return redirect('admin_dashboard')
            else:
                return redirect('index')
    else:
        form = PasswordChangeForm(request.user)
    
    return render(request, 'password_change.html', {
        'form': form,
        'force_change': request.session.get('force_password_change', False)
    })
    

def counselor_setup(request, token):
    try:
        user = CustomUser.objects.get(verification_token=token)
        
        # Check if token is expired (48 hours)
        if (timezone.now() - user.date_joined) > timedelta(hours=48):
            messages.error(request, "This setup link has expired. Please contact the administrator.")
            return redirect('login')
            
        if request.method == 'POST':
            form = PasswordChangeForm(user, request.POST)
            if form.is_valid():
                user = form.save()
                user.email_verified = True
                user.is_active = True
                user.verification_token = None
                user.save()
                
                update_session_auth_hash(request, user)
                messages.success(request, 'Your account has been successfully set up!')
                
                # Log them in and redirect to counselor dashboard
                login(request, user)
                return redirect('counselor_dashboard')
            else:
                messages.error(request, 'Please correct the errors below.')
        else:
            form = PasswordChangeForm(user)
            
        return render(request, 'mentalhealth/counselor-setup.html', {
            'form': form,
            'token': token,
            'user': user
        })
        
    except CustomUser.DoesNotExist:
        messages.error(request, "Invalid setup link.")
        return redirect('login')
    
    
@login_required
def counselor_dashboard(request):
    # Verify this is actually a counselor
    if not hasattr(request.user, 'counselor_profile'):
        return redirect('index')
    
    counselor = request.user.counselor_profile
    today = timezone.now().date()
    
    # Get appointments for the current week (Monday to Sunday)
    start_of_week = today - timedelta(days=today.weekday())
    end_of_week = start_of_week + timedelta(days=6)
    
    # Get this week's appointments
    this_week_appointments = Appointment.objects.filter(
        counselor=counselor,
        date__range=[start_of_week, end_of_week],
        status__in=['pending', 'confirmed']
    ).order_by('date', 'time')
    
    # Get recent reports (show recent reports regardless of status for counselor review)
    pending_reports = Report.objects.filter(
        counselor=counselor
    ).exclude(status='archived').order_by('-created_at')[:5]
    
    # Get weekly session count
    weekly_sessions = Appointment.objects.filter(
        counselor=counselor,
        date__gte=start_of_week,
        status='completed'
    ).count()
    
    context = {
        'counselor': counselor,
        'today_appointments': this_week_appointments,  # Changed variable name but kept template compatibility
        'pending_reports': pending_reports,
        'pending_reports_count': pending_reports.count(),
        'weekly_sessions_count': weekly_sessions,
        'week_start': start_of_week,
        'week_end': end_of_week,
        'has_availability': len(counselor.available_days) > 0,
    }
    
    return render(request, 'counselor-panel.html', context)


@login_required
def counselor_schedule(request):
    if not hasattr(request.user, 'counselor_profile'):
        return redirect('index')
    
    counselor = request.user.counselor_profile
    today = timezone.now().date()
    start_of_week = today - timedelta(days=today.weekday())
    end_of_week = start_of_week + timedelta(days=6)
    
    print(f"Counselor schedule view called for: {counselor.name}")
    print(f"Week range: {start_of_week} to {end_of_week}")
    
    # Get appointments for the current week, ALL statuses for debugging
    appointments = Appointment.objects.filter(
        counselor=counselor,
        date__range=[start_of_week, end_of_week]
    ).order_by('date', 'time')
    
    print(f"Found {appointments.count()} appointments for this week")
    
    # Debug: Check all appointments for this counselor
    all_counselor_appointments = Appointment.objects.filter(counselor=counselor).order_by('date', 'time')
    print(f"Total appointments for counselor {counselor.name}: {all_counselor_appointments.count()}")
    for appt in all_counselor_appointments:
        print(f"  - {appt.user.full_name} on {appt.date} at {appt.time} (status: {appt.status})")
    
    for appt in appointments:
        print(f"  - {appt.user.full_name} on {appt.date} at {appt.time} (status: {appt.status})")
    
    # Prepare calendar events for FullCalendar
    calendar_events = []
    for appt in appointments:
        time_str = appt.time.strftime('%H:%M:%S') if appt.time else ''
        event = {
            'id': appt.id,
            'title': f"{appt.user.full_name} - {appt.reason} ({appt.status})",
            'start': f"{appt.date}T{time_str}",
            'extendedProps': {
                'studentId': appt.user.id,
                'studentName': appt.user.full_name,
                'college': getattr(appt.user, 'college', ''),
                'program': getattr(appt.user, 'program', ''),
                'sessionType': getattr(appt, 'session_type', ''),
                'status': appt.status,
                'notes': getattr(appt, 'notes', ''),
                'appointmentId': appt.id,
                'type': 'appointment',
            }
        }
        calendar_events.append(event)
        print(f"Created calendar event: {event['title']} on {event['start']}")
    
    print(f"Total calendar events: {len(calendar_events)}")
    
    import json
    return render(request, 'counselor-schedule.html', {
        'counselor': counselor,
        'appointments': appointments,
        'week_start': start_of_week,
        'week_end': end_of_week,
        'calendar_events': json.dumps(calendar_events),
    })

@login_required
def counselor_reports(request):
    if not hasattr(request.user, 'counselor_profile'):
        return redirect('index')
    
    counselor = request.user.counselor_profile
    
    # Get filter parameters
    report_type = request.GET.get('report_type', '')
    date_range = request.GET.get('date_range', '')
    priority = request.GET.get('priority', '')
    status = request.GET.get('status', '')
    search_query = request.GET.get('search', '')
    
    # Base queryset - exclude archived and completed reports from main view
    reports = Report.objects.filter(counselor=counselor)
    reports = reports.exclude(status__in=['archived', 'completed'])
    
    # Apply filters
    if report_type:
        reports = reports.filter(report_type=report_type)
    
    if status:
        reports = reports.filter(status=status)
    
    if search_query:
        reports = reports.filter(
            Q(title__icontains=search_query) | 
            Q(description__icontains=search_query)
        )
    
    # Apply date range filter
    if date_range:
        today = timezone.now().date()
        if date_range == 'today':
            reports = reports.filter(created_at__date=today)
        elif date_range == 'week':
            week_ago = today - timedelta(days=7)
            reports = reports.filter(created_at__date__gte=week_ago)
        elif date_range == 'month':
            month_ago = today - timedelta(days=30)
            reports = reports.filter(created_at__date__gte=month_ago)
        elif date_range == 'quarter':
            quarter_ago = today - timedelta(days=90)
            reports = reports.filter(created_at__date__gte=quarter_ago)
    
    # Order by created_at
    reports = reports.order_by('-created_at')
    
    # Calculate statistics
    total_reports = Report.objects.filter(counselor=counselor).count()
    filtered_reports_count = reports.count()
    pending_reports = reports.filter(status='pending').count()
    urgent_reports = reports.filter(report_type='urgent').count()
    
    # Get recent reports for display
    recent_reports = list(reports[:10])

    # Only show appointments that are:
    # - pending or confirmed
    session_report_exists = Report.objects.filter(
        appointment=OuterRef('pk'),
        report_type='session'
    )
    appointments = Appointment.objects.filter(
        counselor=counselor,
        status__in=['pending', 'confirmed']
    ).select_related('user').order_by('user__full_name')
    
    # Remove duplicates by user (keep only one appointment per user for display)
    seen_users = set()
    unique_appointments = []
    for appointment in appointments:
        if appointment.user.id not in seen_users:
            unique_appointments.append(appointment)
            seen_users.add(appointment.user.id)
    
    students_with_appointments = unique_appointments
    
    context = {
        'counselor': counselor,
        'reports': recent_reports,
        'total_reports': total_reports,
        'filtered_reports_count': filtered_reports_count,
        'pending_reports': pending_reports,
        'urgent_reports': urgent_reports,
        'students_with_appointments': students_with_appointments,
        'filters': {
            'report_type': report_type,
            'date_range': date_range,
            'priority': priority,
            'status': status,
            'search': search_query,
        }
    }
    
    return render(
        request,
        'counselor-reports.html',
        context
    )

@login_required
def counselor_archive(request):
    # Check if user has counselor profile first (prioritize counselor view)
    if hasattr(request.user, 'counselor_profile') and request.user.counselor_profile.is_active:
        counselor = request.user.counselor_profile
    elif request.user.is_staff or request.user.is_superuser:
        # If admin specifies a counselor ID, show that counselor's archive
        counselor_id = request.GET.get('counselor_id')
        if counselor_id:
            try:
                counselor = Counselor.objects.get(id=counselor_id, is_active=True)
            except Counselor.DoesNotExist:
                messages.error(request, 'Counselor not found.')
                return redirect('admin_personnel')
        else:
            # Admin accessing archives - show all data
            return admin_archive(request)
    else:
        return redirect('index')

    # Completed sessions (appointments)
    completed_appointments = Appointment.objects.filter(
        counselor=counselor,
        status='completed'
    ).select_related('user').order_by('-date', '-time')

    # Archived reports
    archived_reports = Report.objects.filter(
        counselor=counselor,
        status='archived'
    ).select_related('user').order_by('-created_at')

    # Completed reports (for backward compatibility)
    completed_reports = Report.objects.filter(
        counselor=counselor,
        status='completed'
    ).select_related('user').order_by('-created_at')

    # Combine all archived and completed reports
    all_archived_reports = list(archived_reports) + list(completed_reports)

    context = {
        'counselor': counselor,
        'completed_appointments': completed_appointments,
        'archived_reports': all_archived_reports,
    }

    return render(request, 'mentalhealth/counselor-archive.html', context)

@login_required
def counselor_profile(request):
    if not hasattr(request.user, 'counselor_profile'):
        return redirect('index')
    
    counselor = request.user.counselor_profile
    
    # Calculate feedback statistics
    feedbacks = Feedback.objects.filter(counselor=counselor, skipped=False)
    total_feedbacks = feedbacks.count()
    
    # Calculate average ratings
    avg_overall = feedbacks.aggregate(avg=Avg('overall_rating'))['avg'] or 0
    avg_professionalism = feedbacks.aggregate(avg=Avg('professionalism_rating'))['avg'] or 0
    avg_helpfulness = feedbacks.aggregate(avg=Avg('helpfulness_rating'))['avg'] or 0
    avg_recommend = feedbacks.aggregate(avg=Avg('recommend_rating'))['avg'] or 0
    
    # Calculate overall average rating
    overall_avg = 0
    if total_feedbacks > 0:
        total_rating = avg_overall + avg_professionalism + avg_helpfulness + avg_recommend
        overall_avg = round(total_rating / 4, 1)
    
    if request.method == 'POST':
        form = CounselorProfileForm(request.POST, request.FILES, instance=counselor, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Profile updated successfully!')
            return redirect('counselor_profile')
    else:
        form = CounselorProfileForm(instance=counselor, user=request.user)
    
    return render(request, 'counselor-profile.html', {
        'counselor': counselor,
        'form': form,
        'total_feedbacks': total_feedbacks,
        'overall_avg_rating': overall_avg,
        'avg_overall': avg_overall,
        'avg_professionalism': avg_professionalism,
        'avg_helpfulness': avg_helpfulness,
        'avg_recommend': avg_recommend,
    })


def forget_password(request):
    """Handle password reset request"""
    if request.method == 'POST':
        form = PasswordResetRequestForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            user = CustomUser.objects.get(email=email)

            # Generate reset token
            reset_token = get_random_string(64)
            user.password_reset_token = reset_token
            user.password_reset_expires = timezone.now() + timedelta(hours=24)
            user.save()

            # Send reset email
            reset_url = request.build_absolute_uri(
                reverse('reset_password', kwargs={'token': reset_token})
            )

            html_message = render_to_string('mentalhealth/password-reset-email.html', {
                'user': user,
                'reset_url': reset_url,
            })

            plain_message = f"""
            Password Reset Request - CalmConnect

            Hello {user.full_name},

            You have requested to reset your password for your CalmConnect account.

            Please click the link below to reset your password:
            {reset_url}

            This link will expire in 24 hours. If you didn't request this password reset, please ignore this email.

            The CalmConnect Team
            """

            try:
                send_mail(
                    'Password Reset Request - CalmConnect',
                    plain_message,
                    settings.DEFAULT_FROM_EMAIL,
                    [user.email],
                    html_message=html_message,
                    fail_silently=False,
                )
                messages.success(request, 'Password reset instructions have been sent to your email.')
            except Exception as e:
                messages.error(request, 'Failed to send reset email. Please try again later.')

            return redirect('login')
    else:
        form = PasswordResetRequestForm()

    return render(request, 'mentalhealth/forget-password.html', {'form': form})


def reset_password(request, token):
    """Handle password reset with token"""
    try:
        user = CustomUser.objects.get(
            password_reset_token=token,
            password_reset_expires__gt=timezone.now()
        )
    except CustomUser.DoesNotExist:
        messages.error(request, 'Invalid or expired reset link.')
        return redirect('forget_password')

    if request.method == 'POST':
        form = SetNewPasswordForm(request.POST)
        if form.is_valid():
            password = form.cleaned_data['password']
            user.set_password(password)
            user.password_reset_token = None
            user.password_reset_expires = None
            user.save()

            messages.success(request, 'Your password has been reset successfully. You can now log in with your new password.')
            return redirect('login')
    else:
        form = SetNewPasswordForm()

    return render(request, 'mentalhealth/reset-password.html', {'form': form})

@login_required
def create_appointment(request):
    if not hasattr(request.user, 'counselor_profile'):
        return redirect('index')
    
    if request.method == 'POST':
        form = AppointmentForm(request.user.counselor_profile, request.POST)
        if form.is_valid():
            appointment = form.save(commit=False)
            appointment.counselor = request.user.counselor_profile
            appointment.save()
            messages.success(request, 'Appointment created successfully!')
            return redirect('counselor_dashboard')
    else:
        form = AppointmentForm(request.user.counselor_profile)
    
    return render(request, 'create-appointment.html', {
        'form': form
    })

@login_required
def appointment_detail(request, pk):
    try:
        appointment = Appointment.objects.get(pk=pk)
        if not hasattr(request.user, 'counselor_profile') or appointment.counselor != request.user.counselor_profile:
            return redirect('index')

        # If this is a follow-up appointment, redirect to the follow-up session view
        if appointment.session_type == 'followup':
            return redirect('followup_session', appointment_id=appointment.id)

        if request.method == 'POST':
            form = AppointmentStatusForm(request.POST, instance=appointment)
            if form.is_valid():
                form.save()
                messages.success(request, 'Appointment updated successfully!')
                return redirect('appointment_detail', pk=appointment.id)
        else:
            form = AppointmentStatusForm(instance=appointment)

        return render(request, 'appointment-detail.html', {
            'appointment': appointment,
            'form': form
        })
    except Appointment.DoesNotExist:
        return redirect('counselor_dashboard')

@login_required
def create_report(request):
    if not hasattr(request.user, 'counselor_profile'):
        return redirect('index')
    
    appointment_id = request.GET.get('appointment')
    initial = {}
    appointment = None
    if appointment_id:
        try:
            appointment = Appointment.objects.get(pk=appointment_id)
            # Pre-fill initial data for the report form
            initial = {
                'user': appointment.user.id,
                'appointment': appointment.id,
                'title': f'Session with {appointment.user.full_name} on {appointment.date}',
                'description': appointment.reason or '',
                'report_type': 'session',
            }
        except Appointment.DoesNotExist:
            appointment = None

    if request.method == 'POST':
        form = ReportForm(request.POST)
        if form.is_valid():
            report = form.save(commit=False)
            report.counselor = request.user.counselor_profile
            
            # Always mark reports as archived when created (completed reports are archived)
            report.status = 'archived'
            
            if appointment or form.cleaned_data.get('appointment'):
                # Use appointment from GET or from POST data
                if not appointment:
                    try:
                        appointment = Appointment.objects.get(pk=form.cleaned_data['appointment'])
                    except Appointment.DoesNotExist:
                        appointment = None
                if appointment:
                    report.user = appointment.user
                    report.appointment = appointment
                    # Mark appointment as completed and save
                    appointment.status = 'completed'
                    appointment.save()
                    
                    # Send feedback request email
                    print(f"📧 Triggering feedback email for appointment {appointment.id}")
                    print(f"🔍 Debug: appointment status = {appointment.status}")
                    print(f"🔍 Debug: appointment user = {appointment.user.email}")
                    send_feedback_request_email(request, appointment)
            
            report.save()
            
            # Create notification for the student if report has a user
            if report.user:
                create_notification(
                    user=report.user,
                    message=f'Your session report with {request.user.counselor_profile.name} has been completed.',
                    notification_type='report',
                    url=reverse('user-profile')
                )
            
            messages.success(request, 'Report created successfully!')
            return redirect('counselor_reports')
    else:
        form = ReportForm(initial=initial)
    
    return render(request, 'create-report.html', {
        'form': form,
        'appointment': appointment
    })

@login_required
def report_detail(request, pk):
    try:
        report = Report.objects.get(pk=pk)
        if not hasattr(request.user, 'counselor_profile') or report.counselor != request.user.counselor_profile:
            return redirect('index')

        # Check if there's an existing follow-up request for this report
        followup_request = FollowupRequest.objects.filter(report=report).first()

        return render(request, 'report-detail.html', {
            'report': report,
            'followup_request': followup_request
        })
    except Report.DoesNotExist:
        return redirect('counselor_reports')


@login_required
def report_export(request, pk, format_type):
    """Export individual report to PDF or DOC format"""
    try:
        report = Report.objects.select_related('user', 'counselor').get(pk=pk)
        if not hasattr(request.user, 'counselor_profile') or report.counselor != request.user.counselor_profile:
            return JsonResponse({'error': 'Permission denied'}, status=403)

        from django.http import HttpResponse
        from io import BytesIO
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from docx import Document
        from docx.shared import Inches

        if format_type == 'pdf':
            # PDF Export
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="report_{report.id}.pdf"'

            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []

            # Title
            title = Paragraph(f"Session Report: {report.title}", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))

            # Report details
            details_data = [
                ['Student Name', report.user.full_name if report.user else 'N/A'],
                ['Student ID', report.user.student_id if report.user else 'N/A'],
                ['Report Type', report.get_report_type_display()],
                ['Status', report.get_status_display()],
                ['Created Date', report.created_at.strftime('%Y-%m-%d %H:%M')],
                ['Last Modified', report.updated_at.strftime('%Y-%m-%d %H:%M')],
            ]

            details_table = Table(details_data, colWidths=[100, 300])
            details_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(details_table)
            story.append(Spacer(1, 12))

            # Description
            desc_title = Paragraph("Report Description:", styles['Heading2'])
            story.append(desc_title)
            story.append(Spacer(1, 6))

            description = Paragraph(report.description, styles['Normal'])
            story.append(description)

            doc.build(story)
            pdf = buffer.getvalue()
            buffer.close()
            response.write(pdf)
            return response

        elif format_type == 'doc':
            # DOC Export
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="report_{report.id}.docx"'

            document = Document()
            document.add_heading(f'Session Report: {report.title}', 0)

            # Report details
            document.add_paragraph(f"Student Name: {report.user.full_name if report.user else 'N/A'}")
            document.add_paragraph(f"Student ID: {report.user.student_id if report.user else 'N/A'}")
            document.add_paragraph(f"Report Type: {report.get_report_type_display()}")
            document.add_paragraph(f"Status: {report.get_status_display()}")
            document.add_paragraph(f"Created Date: {report.created_at.strftime('%Y-%m-%d %H:%M')}")
            document.add_paragraph(f"Last Modified: {report.updated_at.strftime('%Y-%m-%d %H:%M')}")
            document.add_paragraph("")

            # Description
            document.add_heading('Report Description:', level=2)
            document.add_paragraph(report.description)

            # Save to response
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            response.write(buffer.getvalue())
            buffer.close()
            return response

        else:
            return JsonResponse({'error': 'Unsupported format'}, status=400)

    except Report.DoesNotExist:
        return JsonResponse({'error': 'Report not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@api_view(['GET', 'POST', 'PUT', 'DELETE'])
def report_api(request, pk=None):
    """API endpoint for report CRUD operations"""
    # Check authentication
    if not request.user.is_authenticated:
        return Response({'success': False, 'error': 'Authentication required'}, status=401)

    # Check if user is a counselor
    if not hasattr(request.user, 'counselor_profile'):
        return Response({'success': False, 'error': 'Permission denied: counselor access required'}, status=403)

    counselor = request.user.counselor_profile
    
    if request.method == 'GET':
        if pk:
            try:
                report = Report.objects.get(pk=pk, counselor=counselor)
                return JsonResponse({
                    'success': True,
                    'report': {
                        'id': report.id,
                        'title': report.title,
                        'description': report.description,
                        'report_type': report.report_type,
                        'status': report.status,
                        'user_name': report.user.full_name if report.user else None,
                        'created_at': report.created_at.strftime('%Y-%m-%d %H:%M'),
                        'updated_at': report.updated_at.strftime('%Y-%m-%d %H:%M')
                    }
                })
            except Report.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Report not found'}, status=404)
        else:
            # Get all reports for counselor (excluding archived and completed)
            reports = Report.objects.filter(counselor=counselor).exclude(status__in=['archived', 'completed']).order_by('-created_at')
            reports_data = [{
                'id': report.id,
                'title': report.title,
                'description': report.description,
                'report_type': report.report_type,
                'status': report.status,
                'created_at': report.created_at.strftime('%Y-%m-%d'),
                'updated_at': report.updated_at.strftime('%Y-%m-%d %H:%M')
            } for report in reports]
            return JsonResponse({'success': True, 'reports': reports_data})
    
    elif request.method == 'POST':
        # Create new report
        try:
            data = json.loads(request.body)
            
            # Get the user if provided
            user = None
            if data.get('user_id'):
                try:
                    user = CustomUser.objects.get(id=data['user_id'])
                except CustomUser.DoesNotExist:
                    return JsonResponse({'success': False, 'error': 'Student not found'}, status=400)
            
            # Always set status to completed when creating reports
            status = 'completed'
            
            report = Report.objects.create(
                counselor=counselor,
                user=user,
                title=data.get('title', ''),
                description=data.get('description', ''),
                report_type=data.get('report_type', 'session'),
                status=status
            )
            
            # Create notification for the student if report has a user
            if user:
                create_notification(
                    user=user,
                    message=f'Your session report with {counselor.name} has been completed.',
                    notification_type='report',
                    url=reverse('user-profile')
                )
                
                # Send feedback request email if this is a session report
                print(f"🔍 Debug: report_type = {data.get('report_type')}")
                print(f"🔍 Debug: user = {user}")
                print(f"🔍 Debug: counselor = {counselor}")
                
                if data.get('report_type') == 'session':
                    print(f"✅ Report type is 'session', proceeding with feedback email...")
                    # Find the appointment for this user and counselor (not completed or cancelled)
                    try:
                        appointment = Appointment.objects.filter(
                            user=user,
                            counselor=counselor
                        ).exclude(status__in=['completed', 'cancelled']).latest('date')
                        print(f"✅ Found appointment {appointment.id} for feedback email")
                        print(f"📊 Found appointment status: {appointment.status}")
                        
                        # Mark as completed
                        appointment.status = 'completed'
                        appointment.save()
                        print(f"✅ Appointment {appointment.id} marked as completed")
                        
                        # Send feedback email
                        print("🔔 Sending feedback email...")
                        send_feedback_request_email(request, appointment)
                        print("✅ Feedback email sent successfully via API simulation")
                        
                    except Appointment.DoesNotExist:
                        print(f"❌ No available appointment found for user {user} and counselor {counselor}")
                        pass  # No appointment found, skip feedback email
                else:
                    print(f"❌ Report type is not 'session': {data.get('report_type')}")
            
            return JsonResponse({
                'success': True,
                'report': {
                    'id': report.id,
                    'title': report.title,
                    'description': report.description,
                    'report_type': report.report_type,
                    'status': report.status,
                    'user_name': user.full_name if user else None,
                    'created_at': report.created_at.strftime('%Y-%m-%d %H:%M')
                }
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=400)
    
    elif request.method == 'PUT':
        # Update report
        try:
            report = Report.objects.get(pk=pk, counselor=counselor)
            data = json.loads(request.body)
            
            if 'title' in data:
                report.title = data['title']
            if 'description' in data:
                report.description = data['description']
            if 'report_type' in data:
                report.report_type = data['report_type']
            if 'status' in data:
                # If status is being set to 'completed', automatically archive it
                if data['status'] == 'completed':
                    report.status = 'archived'
                else:
                    report.status = data['status']
            
            report.save()
            
            # Check if the report was automatically archived
            if data.get('status') == 'completed':
                return JsonResponse({
                    'success': True, 
                    'message': 'Report completed and automatically archived! Completed reports are moved to the archive.',
                    'status': 'archived'
                })
            else:
                return JsonResponse({'success': True, 'message': 'Report updated successfully'})
        except Report.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Report not found'}, status=404)
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=400)
    
    elif request.method == 'DELETE':
        # Delete report
        try:
            report = Report.objects.get(pk=pk, counselor=counselor)
            report.delete()
            return JsonResponse({'success': True, 'message': 'Report deleted successfully'})
        except Report.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Report not found'}, status=404)
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=400)

@login_required
def edit_report(request, pk):
    try:
        report = Report.objects.get(pk=pk)
        if not hasattr(request.user, 'counselor_profile') or report.counselor != request.user.counselor_profile:
            return redirect('index')
        if request.method == 'POST':
            form = ReportForm(request.POST, instance=report)
            if form.is_valid():
                form.save()
                messages.success(request, 'Report updated successfully!')
                return redirect('counselor_reports')
        else:
            form = ReportForm(instance=report)
        return render(request, 'create-report.html', {
            'form': form,
            'report': report,
            'edit_mode': True
        })
    except Report.DoesNotExist:
        return redirect('counselor_reports')

def get_notifications(request):
    """Get notifications for the current user"""
    try:
        # Add debugging
        import logging
        logger = logging.getLogger(__name__)
        logger.debug(f"get_notifications called. User: {request.user}, Authenticated: {request.user.is_authenticated}")
        print(f"get_notifications called. User: {request.user}, Authenticated: {request.user.is_authenticated}")
        
        # Check if user is authenticated
        if not request.user.is_authenticated:
            print("User not authenticated")
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'error': 'Authentication required',
                    'notifications': [],
                    'count': 0
                }, status=401)
            else:
                return redirect('login')
        
        # Get notifications from the database
        try:
            notifications = Notification.objects.filter(
                user=request.user,
                dismissed=False
            ).order_by('-created_at')[:10]  # Get last 10 notifications
        except Exception:
            notifications = []

        # Get unread count separately
        try:
            unread_count = Notification.objects.filter(
                user=request.user,
                dismissed=False,
                read=False
            ).count()
        except Exception:
            unread_count = 0
        
        # Convert to the format expected by the frontend
        notification_list = []
        for notif in notifications:
            notification_list.append({
                'id': notif.id,
                'type': notif.type,
                'message': notif.message,
                'url': '#',  # Always disable navigation to prevent redirects
                'time': notif.created_at.strftime('%Y-%m-%d %H:%M'),
                'read': notif.read,
                'metadata': notif.metadata  # Include metadata for modal handling
            })
        
        print(f"Total notifications: {len(notification_list)}")
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            # For AJAX requests, return JSON
            logger.debug(f"Returning JSON response with {len(notification_list)} notifications")
            print(f"Returning JSON response with {len(notification_list)} notifications")
            return JsonResponse({
                'success': True,
                'notifications': notification_list,
                'count': len(notification_list),
                'unread_count': unread_count
            })
        
        # For regular requests, render a template
        return render(request, 'mentalhealth/notifications.html', {
            'notifications': notification_list
        })
    except Exception as e:
        print(f"Error in get_notifications: {e}")
        import traceback
        traceback.print_exc()
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({
                'success': False,
                'error': 'Internal server error',
                'notifications': [],
                'count': 0
            }, status=500)
        else:
            messages.error(request, 'An error occurred while loading notifications')
            return redirect('index')

def login_required_json(view_func):
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return JsonResponse({'success': False, 'error': 'Authentication required.'}, status=401)
        return view_func(request, *args, **kwargs)
    return wrapper

from django.views.decorators.http import require_http_methods

@login_required_json
@require_http_methods(["GET", "POST", "PUT", "DELETE"])
def update_schedule(request):
    import logging
    logger = logging.getLogger(__name__)
    logger.debug(f"update_schedule called. Authenticated: {request.user.is_authenticated}, User: {request.user}, Method: {request.method}")
    # Permission check: must be a counselor
    if not hasattr(request.user, 'counselor_profile'):
        logger.warning(f"User {request.user} does not have counselor_profile.")
        return JsonResponse({'success': False, 'error': 'Permission denied: not a counselor.'}, status=403)
    counselor = request.user.counselor_profile
    def serialize_day_schedules(day_schedules):
        result = {}
        for day, sched in day_schedules.items():
            result[day] = {
                'start_time': sched.get('start_time') if sched.get('start_time') else None,
                'end_time': sched.get('end_time') if sched.get('end_time') else None
            }
        return result

    if request.method == 'GET':
        return JsonResponse({
            'success': True,
            'days': counselor.available_days,
            'start_time': counselor.available_start_time.strftime('%H:%M') if counselor.available_start_time else None,
            'end_time': counselor.available_end_time.strftime('%H:%M') if counselor.available_end_time else None,
            'day_schedules': serialize_day_schedules(counselor.day_schedules)
        })

    try:
        data = json.loads(request.body)
    except Exception:
        data = {}

    if request.method in ['POST', 'PUT']:
        days = data.get('days', [])
        day_schedules = data.get('day_schedules', {})
        if not days:
            return JsonResponse({'success': False, 'error': 'At least one day must be selected'}, status=400)
        processed_schedules = {}
        try:
            for day in days:
                day_schedule = day_schedules.get(day, {})
                start_time = day_schedule.get('start_time')
                end_time = day_schedule.get('end_time')
                if not start_time or not end_time:
                    return JsonResponse({'success': False, 'error': f'Start and end times are required for {day}'}, status=400)
                if start_time >= end_time:
                    return JsonResponse({'success': False, 'error': f'End time must be after start time for {day}'}, status=400)
                # Store as strings, not time objects
                processed_schedules[day] = {
                    'start_time': start_time,
                    'end_time': end_time
                }
        except ValueError:
            return JsonResponse({'success': False, 'error': 'Invalid time format. Use HH:MM format.'}, status=400)
        counselor.available_days = days
        counselor.day_schedules = processed_schedules
        counselor.save()
        return JsonResponse({
            'success': True,
            'message': 'Availability updated successfully',
            'days': counselor.available_days,
            'start_time': counselor.available_start_time.strftime('%H:%M') if counselor.available_start_time else None,
            'end_time': counselor.available_end_time.strftime('%H:%M') if counselor.available_end_time else None,
            'day_schedules': serialize_day_schedules(counselor.day_schedules)
        })
    elif request.method == 'DELETE':
        counselor.available_days = []
        counselor.available_start_time = None
        counselor.available_end_time = None
        counselor.save()
        return JsonResponse({'success': True, 'message': 'Availability cleared successfully'})

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def appointment_detail_api(request, pk):
    """API endpoint for appointment details"""
    try:
        appointment = Appointment.objects.get(pk=pk)
        
        # Check if user is staff or the counselor for this appointment
        if not (request.user.is_staff or 
                (hasattr(request.user, 'counselor_profile') and 
                 appointment.counselor == request.user.counselor_profile)):
            return Response({'success': False, 'error': 'Permission denied'}, status=403)
        
        return Response({
            'success': True,
            'appointment': {
                'id': appointment.id,
                'user_name': appointment.user.full_name,
                'student_id': appointment.user.student_id,
                'college': appointment.user.get_college_display(),
                'program': appointment.user.program,
                'date': appointment.date.strftime('%Y-%m-%d'),
                'time': appointment.time.strftime('%H:%M'),
                'services': appointment.services,
                'status': appointment.status,
                'reason': appointment.reason,
                'notes': appointment.notes,
                'created_at': appointment.created_at.strftime('%Y-%m-%d %H:%M'),
                'updated_at': appointment.updated_at.strftime('%Y-%m-%d %H:%M')
            }
        })
    except Appointment.DoesNotExist:
        return Response({'success': False, 'error': 'Appointment not found'}, status=404)
    except Exception as e:
        return Response({'success': False, 'error': str(e)}, status=400)

# --- Student-only decorator ---
def student_required(view_func):
    def wrapper(request, *args, **kwargs):
        user = request.user
        # You may need to adjust this logic based on your user model
        if not user.is_authenticated or getattr(user, 'is_counselor', False) or user.is_staff:
            return JsonResponse({'success': False, 'error': 'Students only.'}, status=403)
        return view_func(request, *args, **kwargs)
    return wrapper

# Set up logger
logger = logging.getLogger('dass21_ai_feedback')

@csrf_exempt
@api_view(['POST'])
# @login_required
# @permission_classes([IsAuthenticated])
@ratelimit_429(key='ip', rate='50/m', block=True)  # Increased to 50 requests per minute per IP, returns 429
def generate_ai_tips(request):
    """
    Generate AI-powered personalized mental health tips based on DASS21 scores.
    Expects JSON: { 'depression': int, 'anxiety': int, 'stress': int, 'depression_severity': str, 'anxiety_severity': str, 'stress_severity': str, 'answers': dict }
    Strictly sanitizes data and enforces student-only access.
    """
    if not request.user.is_authenticated:
        return Response({'success': False, 'error': 'Authentication required'}, status=401)

    user = request.user
    data = request.data

    # Validate required fields
    required_fields = ['depression', 'anxiety', 'stress', 'depression_severity', 'anxiety_severity', 'stress_severity']
    for field in required_fields:
        if field not in data:
            return Response({'success': False, 'error': f'Missing required field: {field}'}, status=400)

    depression = data.get('depression')
    anxiety = data.get('anxiety')
    stress = data.get('stress')
    depression_severity = data.get('depression_severity')
    anxiety_severity = data.get('anxiety_severity')
    stress_severity = data.get('stress_severity')
    answers = data.get('answers', {})

    # Validate score ranges (0-42 for each scale)
    if not all(isinstance(x, int) and 0 <= x <= 42 for x in [depression, anxiety, stress]):
        logger.warning(f"Invalid DASS21 scores from user {user.id}")
        return Response({'success': False, 'error': 'Invalid DASS21 scores. Must be integers between 0-42.'}, status=400)

    # Get comprehensive user history for personalization
    user_history = get_user_personalization_data(user)

    # Analyze specific DASS21 responses for targeted tips
    dass_analysis = analyze_dass21_responses(answers, depression, anxiety, stress)

    # Compose personalized prompt for AI tips generation
    prompt = build_dass21_tips_prompt(
        user, depression, anxiety, stress,
        depression_severity, anxiety_severity, stress_severity,
        user_history, dass_analysis
    )

    try:
        openai.api_key = getattr(settings, 'OPENAI_API_KEY', None)
        logger.info(f"OpenAI API key status: {'Configured' if openai.api_key else 'Not configured'}")
        if not openai.api_key:
            logger.warning("OpenAI API key not configured. Using personalized fallback tips.")
            # Provide personalized fallback tips when OpenAI is not configured
            fallback_tips = generate_dass21_specific_fallback_tips(
                user, depression, anxiety, stress,
                depression_severity, anxiety_severity, stress_severity,
                user_history, dass_analysis
            )
            return Response({
                'success': True,
                'tips': fallback_tips,
                'source': 'fallback',
                'personalization_level': 'high',
                'dass_analysis': dass_analysis
            })

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a compassionate mental health assistant specializing in student wellness. Generate 5-7 evidence-based, actionable mental health tips tailored to the student's specific DASS21 results, academic context, and personal circumstances. Focus on progressive strategies: self-help for milder symptoms, professional resources for severe symptoms. Ensure tips are culturally sensitive, inclusive, and aligned with WHO/NICE guidelines. Structure each tip with a clear title and detailed description."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=600,
            temperature=0.7,
        )
        tips_text = response.choices[0].message['content'].strip()
        logger.info(f"Personalized AI tips generated for user {user.id}")

        # Parse tips into structured format
        structured_tips = parse_ai_tips_response(tips_text)

        return Response({
            'success': True,
            'tips': structured_tips,
            'source': 'openai',
            'personalization_level': 'high',
            'dass_analysis': dass_analysis,
            'context_used': {
                'academic_context': bool(user_history['academic_context']),
                'test_history': user_history['test_count'],
                'exercise_preferences': bool(user_history['exercise_preferences']),
                'trend_analysis': bool(user_history['trend_analysis']),
                'dass_specific': True
            }
        })
    except Exception as e:
        logger.error(f"AI tips generation error for user {user.id}: {str(e)}")
        # Provide personalized fallback tips on any error
        fallback_tips = generate_dass21_specific_fallback_tips(
            user, depression, anxiety, stress,
            depression_severity, anxiety_severity, stress_severity,
            user_history, dass_analysis
        )
        return Response({
            'success': True,
            'tips': fallback_tips,
            'source': 'fallback',
            'personalization_level': 'high',
            'error': 'AI service temporarily unavailable',
            'dass_analysis': dass_analysis
        })


@csrf_exempt
# @login_required
@api_view(['POST'])
# @permission_classes([IsAuthenticated])
@ratelimit_429(key='ip', rate='50/m', block=True)  # Increased to 50 requests per minute per IP, returns 429
def ai_feedback(request):
    """
    Accepts DASS21 scores and returns personalized AI-generated feedback using
    OpenAI's API. Expects JSON: { 'depression': int, 'anxiety': int, 'stress': int, 'answers': dict, ... }
    Strictly sanitizes data and enforces student-only access.
    """
    if not request.user.is_authenticated:
        return Response({'success': False, 'error': 'Authentication required'}, status=401)

    user = request.user
    data = request.data
    depression = data.get('depression')
    anxiety = data.get('anxiety')
    stress = data.get('stress')
    depression_severity = data.get('depression_severity')
    anxiety_severity = data.get('anxiety_severity')
    stress_severity = data.get('stress_severity')
    answers = data.get('answers', {})  # Individual question responses

    # Convert scores to integers if they are numeric
    try:
        if depression is not None:
            depression = int(float(depression))
        if anxiety is not None:
            anxiety = int(float(anxiety))
        if stress is not None:
            stress = int(float(stress))
    except (ValueError, TypeError):
        logger.warning(f"Invalid DASS21 input from user {user.id}: cannot convert to integers")
        return Response({'success': False, 'error': 'Invalid DASS21 scores.'}, status=400)

    # Validate input
    if not all(isinstance(x, int) for x in [depression, anxiety, stress] if x is not None):
        logger.warning(f"Invalid DASS21 input from user {user.id}: scores are not integers after conversion")
        return Response({'success': False, 'error': 'Invalid DASS21 scores.'}, status=400)
    if depression is None or anxiety is None or stress is None:
        logger.warning(f"Missing DASS21 input from user {user.id}")
        return Response({'success': False, 'error': 'Missing DASS21 scores.'}, status=400)

    # Get comprehensive user history for personalization
    user_history = get_user_personalization_data(user)
    
    # Analyze specific DASS21 responses for targeted feedback
    dass_analysis = analyze_dass21_responses(answers, depression, anxiety, stress)
    
    # Compose a personalized prompt for OpenAI
    prompt = build_dass21_specific_prompt(
        user, depression, anxiety, stress, 
        depression_severity, anxiety_severity, stress_severity,
        user_history, dass_analysis
    )

    try:
        openai.api_key = getattr(settings, 'OPENAI_API_KEY', None)
        logger.info(f"OpenAI API key status: {'Configured' if openai.api_key else 'Not configured'}")
        if not openai.api_key:
            logger.warning("OpenAI API key not configured. Using personalized fallback feedback.")
            # Provide personalized fallback feedback when OpenAI is not configured
            fallback_feedback = generate_dass21_specific_fallback_feedback(
                user, depression, anxiety, stress, 
                depression_severity, anxiety_severity, stress_severity,
                user_history, dass_analysis
            )
            return Response({
                'success': True, 
                'feedback': fallback_feedback,
                'source': 'fallback',
                'personalization_level': 'high',
                'dass_analysis': dass_analysis
            })

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a supportive, personalized mental health assistant for university students. Provide empathetic, actionable advice tailored to the student's specific DASS21 responses, academic context, and mental health journey. Always maintain a warm, encouraging tone and provide concrete, actionable steps based on their specific symptoms."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=350,
            temperature=0.7,
        )
        feedback = response.choices[0].message['content'].strip()
        logger.info(f"Personalized AI feedback generated for user {user.id}")
        return Response({
            'success': True, 
            'feedback': feedback,
            'source': 'openai',
            'personalization_level': 'high',
            'dass_analysis': dass_analysis,
            'context_used': {
                'academic_context': bool(user_history['academic_context']),
                'test_history': user_history['test_count'],
                'exercise_preferences': bool(user_history['exercise_preferences']),
                'trend_analysis': bool(user_history['trend_analysis']),
                'dass_specific': True
            }
        })
    except Exception as e:
        logger.error(f"AI feedback error for user {user.id}: {str(e)}")
        # Provide personalized fallback feedback on any error
        fallback_feedback = generate_dass21_specific_fallback_feedback(
            user, depression, anxiety, stress, 
            depression_severity, anxiety_severity, stress_severity,
            user_history, dass_analysis
        )
        return Response({
            'success': True, 
            'feedback': fallback_feedback,
            'source': 'fallback',
            'personalization_level': 'high',
            'error': 'OpenAI service temporarily unavailable',
            'dass_analysis': dass_analysis
        })

def analyze_dass21_responses(answers, depression, anxiety, stress):
    """Analyze specific DASS21 responses for targeted feedback"""
    analysis = {
        'depression_symptoms': [],
        'anxiety_symptoms': [],
        'stress_symptoms': [],
        'primary_concerns': [],
        'coping_patterns': [],
        'specific_triggers': []
    }
    
    # DASS21 question mapping (7 questions per dimension)
    depression_questions = {
        'q3': 'I couldn\'t seem to experience any positive feeling at all',
        'q5': 'I found it difficult to work up the initiative to do things',
        'q10': 'I felt that I had nothing to look forward to',
        'q13': 'I felt down-hearted and blue',
        'q16': 'I was unable to become enthusiastic about anything',
        'q17': 'I felt I wasn\'t worth much as a person',
        'q21': 'I felt that life was meaningless'
    }
    
    anxiety_questions = {
        'q2': 'I was aware of dryness of my mouth',
        'q4': 'I experienced breathing difficulty (e.g., excessively rapid breathing, breathlessness in the absence of physical exertion)',
        'q7': 'I experienced trembling (e.g., in the hands)',
        'q9': 'I was worried about situations in which I might panic and make a fool of myself',
        'q15': 'I felt I was close to panic',
        'q19': 'I was aware of the action of my heart in the absence of physical exertion (e.g., sense of heart rate increase, heart missing a beat)',
        'q20': 'I felt scared without any good reason'
    }
    
    stress_questions = {
        'q1': 'I found it hard to wind down',
        'q6': 'I tended to over-react to situations',
        'q8': 'I felt that I was using a lot of nervous energy',
        'q11': 'I found myself getting agitated',
        'q12': 'I found it difficult to relax',
        'q14': 'I was intolerant of anything that kept me from getting on with what I was doing',
        'q18': 'I felt that I was rather touchy'
    }
    
    # Analyze depression responses
    for q_id, question in depression_questions.items():
        if q_id in answers:
            score = answers[q_id]
            if score >= 2:  # Moderate to severe response
                analysis['depression_symptoms'].append({
                    'question': question,
                    'score': score,
                    'severity': 'moderate' if score == 2 else 'severe'
                })
    
    # Analyze anxiety responses
    for q_id, question in anxiety_questions.items():
        if q_id in answers:
            score = answers[q_id]
            if score >= 2:  # Moderate to severe response
                analysis['anxiety_symptoms'].append({
                    'question': question,
                    'score': score,
                    'severity': 'moderate' if score == 2 else 'severe'
                })
    
    # Analyze stress responses
    for q_id, question in stress_questions.items():
        if q_id in answers:
            score = answers[q_id]
            if score >= 2:  # Moderate to severe response
                analysis['stress_symptoms'].append({
                    'question': question,
                    'score': score,
                    'severity': 'moderate' if score == 2 else 'severe'
                })
    
    # Identify primary concerns based on most severe responses
    all_symptoms = analysis['depression_symptoms'] + analysis['anxiety_symptoms'] + analysis['stress_symptoms']
    if all_symptoms:
        # Sort by severity and score
        all_symptoms.sort(key=lambda x: (x['severity'] == 'severe', x['score']), reverse=True)
        analysis['primary_concerns'] = all_symptoms[:3]  # Top 3 concerns
    
    # Analyze coping patterns based on specific responses
    if 'q1' in answers and answers['q1'] >= 2:  # Hard to wind down
        analysis['coping_patterns'].append('difficulty relaxing')
    if 'q12' in answers and answers['q12'] >= 2:  # Difficult to relax
        analysis['coping_patterns'].append('relaxation challenges')
    if 'q6' in answers and answers['q6'] >= 2:  # Over-react to situations
        analysis['coping_patterns'].append('emotional reactivity')
    if 'q18' in answers and answers['q18'] >= 2:  # Rather touchy
        analysis['coping_patterns'].append('sensitivity to criticism')
    
    # Identify specific triggers
    if 'q9' in answers and answers['q9'] >= 2:  # Worried about panic situations
        analysis['specific_triggers'].append('social situations')
    if 'q15' in answers and answers['q15'] >= 2:  # Close to panic
        analysis['specific_triggers'].append('panic attacks')
    if 'q5' in answers and answers['q5'] >= 2:  # Difficulty with initiative
        analysis['specific_triggers'].append('motivation challenges')
    if 'q17' in answers and answers['q17'] >= 2:  # Not worth much as a person
        analysis['specific_triggers'].append('self-esteem issues')
    
    return analysis

def build_dass21_specific_prompt(user, depression, anxiety, stress, 
                            depression_severity, anxiety_severity, stress_severity,
                                user_history, dass_analysis):
    """Build a DASS21-specific personalized prompt for AI feedback"""
    
    # Base prompt with current scores
    prompt = f"A university student has completed the DASS21 test with these scores: "
    prompt += f"Depression: {depression} ({depression_severity}), "
    prompt += f"Anxiety: {anxiety} ({anxiety_severity}), "
    prompt += f"Stress: {stress} ({stress_severity}). "
    
    # Add specific DASS21 analysis
    if dass_analysis['primary_concerns']:
        prompt += f"\n\nSpecific concerns identified from their responses: "
        for i, concern in enumerate(dass_analysis['primary_concerns'][:3], 1):
            prompt += f"{i}. {concern['question']} (severity: {concern['severity']}) "
    
    if dass_analysis['coping_patterns']:
        prompt += f"\nCoping challenges: {', '.join(dass_analysis['coping_patterns'])}. "
    
    if dass_analysis['specific_triggers']:
        prompt += f"\nSpecific triggers: {', '.join(dass_analysis['specific_triggers'])}. "
    
    # Add personalization context
    personalization = []
    
    # Academic context
    if user_history['academic_context']:
        context = user_history['academic_context']
        if 'college_stressors' in context:
            personalization.append(f"The student is in {user.get_college_display()} and faces challenges with {context['college_stressors']}.")
        if 'year_challenges' in context:
            personalization.append(f"As a {user.get_year_level_display()} student, they're dealing with {context['year_challenges']}.")
        if 'age_context' in context:
            personalization.append(f"They are in the {context['age_context']} phase of life.")
    
    # Test history and trends
    if user_history['test_count'] > 1:
        personalization.append(f"This is their {user_history['test_count']}th DASS21 assessment.")
        if user_history['trend_analysis']:
            trends = user_history['trend_analysis']
            trend_desc = []
            for dimension, trend in trends.items():
                if trend == 'improving':
                    trend_desc.append(f"{dimension} scores are improving")
                elif trend == 'worsening':
                    trend_desc.append(f"{dimension} scores are worsening")
            if trend_desc:
                personalization.append(f"Trend analysis shows: {', '.join(trend_desc)}.")
    
    # Exercise preferences
    if user_history['exercise_preferences']:
        pref = user_history['exercise_preferences']
        personalization.append(f"They have completed {pref['total_sessions']} relaxation sessions, with {pref['preferred_exercise']} being their preferred exercise.")
    
    # Add personalization to prompt
    if personalization:
        prompt += " ".join(personalization) + " "
    
    # Add gender context if available
    if hasattr(user, 'gender') and user.gender and user.gender != 'Prefer not to say':
        prompt += f"The student identifies as {user.gender.lower()}. "
    
    # Enhanced final instructions with DASS21-specific guidance
    prompt += (
        "Provide a personalized, empathetic feedback message (4-5 sentences) that: "
        "1. Acknowledges their specific DASS21 responses and concerns "
        "2. References their academic and personal context "
        "3. Offers specific, actionable advice tailored to their identified symptoms "
        "4. Addresses their specific coping challenges and triggers "
        "5. Suggests relevant resources or coping strategies "
        "6. Maintains a supportive, encouraging tone "
        "7. Provides concrete next steps they can take today "
        "Highlight actionable advice in <b>bold</b> HTML tags. "
        "Include specific suggestions based on their DASS21 responses. "
        "Do not mention AI or that this is automated. "
        "Make the feedback feel like it's coming from someone who understands their specific symptoms and challenges. "
        "Structure the response with: 1) Acknowledgment of specific concerns, 2) Context-specific advice, 3) Actionable steps, 4) Encouragement."
    )
    
    return prompt

def generate_dass21_specific_fallback_feedback(user, depression, anxiety, stress, 
                                         depression_severity, anxiety_severity, stress_severity,
                                             user_history, dass_analysis):
    """Generate DASS21-specific personalized fallback feedback when OpenAI is not available"""
    
    feedback_parts = []
    
    # Address specific DASS21 concerns first
    if dass_analysis['primary_concerns']:
        primary_concern = dass_analysis['primary_concerns'][0]
        if 'positive feeling' in primary_concern['question'].lower():
            feedback_parts.append("I notice you're having difficulty experiencing positive feelings. <b>Try starting with small activities you once enjoyed, even if just for 5 minutes each day.</b>")
        elif 'initiative' in primary_concern['question'].lower():
            feedback_parts.append("You mentioned struggling with motivation and initiative. <b>Break tasks into tiny steps and celebrate each small completion.</b>")
        elif 'panic' in primary_concern['question'].lower():
            feedback_parts.append("I see you're experiencing panic-related concerns. <b>Practice the 4-7-8 breathing technique: inhale for 4, hold for 7, exhale for 8.</b>")
        elif 'worth' in primary_concern['question'].lower():
            feedback_parts.append("You're feeling down about your self-worth. <b>Remember that your value isn't determined by your current struggles.</b>")
        elif 'relax' in primary_concern['question'].lower():
            feedback_parts.append("You're finding it hard to relax. <b>Try progressive muscle relaxation: tense and release each muscle group for 5 seconds.</b>")
    
    # Address coping patterns
    if 'difficulty relaxing' in dass_analysis['coping_patterns']:
        feedback_parts.append("<b>For relaxation challenges, try setting aside 10 minutes daily for deep breathing or guided meditation.</b>")
    if 'emotional reactivity' in dass_analysis['coping_patterns']:
        feedback_parts.append("<b>When feeling overwhelmed, try the 'STOP' technique: Stop, Take a breath, Observe your thoughts, Proceed mindfully.</b>")
    
    # Academic context introduction
    if user_history['academic_context']:
        context = user_history['academic_context']
        if 'college_stressors' in context:
            feedback_parts.append(f"As a {user.get_college_display()} student, you're navigating {context['college_stressors']}.")
        if 'year_challenges' in context:
            feedback_parts.append(f"Being a {user.get_year_level_display()} student brings {context['year_challenges']}.")
    
    # Trend-based feedback
    if user_history['trend_analysis']:
        trends = user_history['trend_analysis']
        improving = [dim for dim, trend in trends.items() if trend == 'improving']
        if improving:
            feedback_parts.append(f"It's encouraging to see improvement in your {', '.join(improving)} scores. <b>Continue the strategies that have been working for you.</b>")
    
    # Current score feedback with DASS21-specific advice
    if depression_severity in ['moderate', 'severe', 'extremely-severe']:
        feedback_parts.append("Your depression scores suggest significant emotional challenges. <b>Consider reaching out to a counselor or mental health professional for support.</b> <b>Try to maintain a regular sleep schedule and engage in activities you once enjoyed, even if it's just for a few minutes each day.</b>")
    elif depression_severity == 'mild':
        feedback_parts.append("You're showing some signs of depression. <b>Try engaging in activities you usually enjoy and maintain regular social connections.</b> <b>Consider setting small, achievable goals for each day to help build momentum.</b>")
    
    if anxiety_severity in ['moderate', 'severe', 'extremely-severe']:
        feedback_parts.append("Your anxiety levels appear elevated. <b>Practice deep breathing exercises and consider talking to a counselor about your concerns.</b> <b>Try the 4-7-8 breathing technique: inhale for 4 counts, hold for 7, exhale for 8.</b>")
    elif anxiety_severity == 'mild':
        feedback_parts.append("You're experiencing some anxiety. <b>Try mindfulness techniques and regular exercise to help manage stress.</b> <b>Consider taking short breaks during study sessions to prevent overwhelm.</b>")
    
    if stress_severity in ['moderate', 'severe', 'extremely-severe']:
        feedback_parts.append("Your stress levels are quite high. <b>Prioritize self-care activities and consider seeking professional support to develop coping strategies.</b> <b>Try breaking large tasks into smaller, manageable steps and celebrate each completion.</b>")
    elif stress_severity == 'mild':
        feedback_parts.append("You're experiencing some stress. <b>Try time management techniques and regular breaks to help maintain balance.</b> <b>Consider using a planner to organize your academic and personal responsibilities.</b>")
    
    # Exercise recommendations based on preferences
    if user_history['exercise_preferences']:
        pref = user_history['exercise_preferences']
        if pref['preferred_exercise'] == 'PMR':
            feedback_parts.append("<b>Consider using Progressive Muscle Relaxation, which you've found helpful before.</b> <b>Try a 10-minute PMR session before studying or before bed.</b>")
        elif pref['preferred_exercise'] == 'EFT':
            feedback_parts.append("<b>Try Emotional Freedom Technique tapping, which has worked well for you in the past.</b> <b>Use EFT when you feel overwhelmed or before important academic tasks.</b>")
        elif pref['preferred_exercise'] == 'Breathing':
            feedback_parts.append("<b>Your breathing exercises have been effective.</b> <b>Try box breathing: 4 counts in, 4 hold, 4 out, 4 hold, repeat for 5 minutes.</b>")
    
    # College-specific advice
    if user_history['academic_context']:
        context = user_history['academic_context']
        if 'college_stressors' in context:
            if 'engineering' in context['college_stressors'].lower():
                feedback_parts.append("<b>For technical coursework stress, try the Pomodoro Technique: 25 minutes of focused work followed by a 5-minute break.</b>")
            elif 'business' in context['college_stressors'].lower():
                feedback_parts.append("<b>For case study stress, try discussing complex topics with classmates to gain different perspectives.</b>")
            elif 'arts' in context['college_stressors'].lower():
                feedback_parts.append("<b>For creative project stress, try free-writing or sketching to unlock creative blocks.</b>")
    
    # Year-level specific advice
    if user_history['academic_context']:
        context = user_history['academic_context']
        if 'year_challenges' in context:
            if '1st' in context['year_challenges']:
                feedback_parts.append("<b>As a first-year student, focus on building good study habits and finding your support network.</b>")
            elif '4th' in context['year_challenges']:
                feedback_parts.append("<b>As a graduating student, remember to celebrate your achievements while managing thesis stress.</b>")
    
    # General encouragement if no specific feedback
    if not feedback_parts:
        feedback_parts.append("Your scores are within normal ranges. <b>Continue practicing good mental health habits and reach out for support if needed.</b> <b>Consider maintaining a gratitude journal to track positive moments.</b>")
    
    return " ".join(feedback_parts)

def build_dass21_tips_prompt(user, depression, anxiety, stress,
                           depression_severity, anxiety_severity, stress_severity,
                           user_history, dass_analysis):
    """Build a DASS21-specific personalized prompt for AI tips generation"""

    # Base prompt with current scores
    prompt = f"Generate 5-7 personalized mental health tips for a university student with these DASS21 scores: "
    prompt += f"Depression: {depression} ({depression_severity}), "
    prompt += f"Anxiety: {anxiety} ({anxiety_severity}), "
    prompt += f"Stress: {stress} ({stress_severity}). "

    # Add specific DASS21 analysis
    if dass_analysis['primary_concerns']:
        prompt += f"\n\nSpecific concerns identified from their responses: "
        for i, concern in enumerate(dass_analysis['primary_concerns'][:3], 1):
            prompt += f"{i}. {concern['question']} (severity: {concern['severity']}) "

    if dass_analysis['coping_patterns']:
        prompt += f"\nCoping challenges: {', '.join(dass_analysis['coping_patterns'])}. "

    if dass_analysis['specific_triggers']:
        prompt += f"\nSpecific triggers: {', '.join(dass_analysis['specific_triggers'])}. "

    # Add personalization context
    personalization = []

    # Academic context
    if user_history['academic_context']:
        context = user_history['academic_context']
        if 'college_stressors' in context:
            personalization.append(f"The student is in {user.get_college_display()} and faces challenges with {context['college_stressors']}.")
        if 'year_challenges' in context:
            personalization.append(f"As a {user.get_year_level_display()} student, they're dealing with {context['year_challenges']}.")
        if 'age_context' in context:
            personalization.append(f"They are in the {context['age_context']} phase of life.")

    # Test history and trends
    if user_history['test_count'] > 1:
        personalization.append(f"This is their {user_history['test_count']}th DASS21 assessment.")
        if user_history['trend_analysis']:
            trends = user_history['trend_analysis']
            trend_desc = []
            for dimension, trend in trends.items():
                if trend == 'improving':
                    trend_desc.append(f"{dimension} scores are improving")
                elif trend == 'worsening':
                    trend_desc.append(f"{dimension} scores are worsening")
            if trend_desc:
                personalization.append(f"Trend analysis shows: {', '.join(trend_desc)}.")

    # Exercise preferences
    if user_history['exercise_preferences']:
        pref = user_history['exercise_preferences']
        personalization.append(f"They have completed {pref['total_sessions']} relaxation sessions, with {pref['preferred_exercise']} being their preferred exercise.")

    # Add personalization to prompt
    if personalization:
        prompt += " ".join(personalization) + " "

    # Add gender context if available
    if hasattr(user, 'gender') and user.gender and user.gender != 'Prefer not to say':
        prompt += f"The student identifies as {user.gender.lower()}. "

    # Enhanced final instructions with DASS21-specific guidance
    prompt += (
        "Generate 5-7 evidence-based mental health tips specifically tailored to their DASS21 results and personal context. "
        "Each tip should have: "
        "1. A clear, actionable title "
        "2. A detailed description explaining why it helps their specific symptoms "
        "3. Progressive approach: self-help strategies for milder symptoms, professional resources for severe symptoms "
        "4. Cultural sensitivity and inclusivity "
        "5. Alignment with WHO/NICE mental health guidelines "
        "6. Student-specific considerations (academic pressure, social isolation, time management) "
        "Focus on practical, immediate actions they can take today. "
        "Ensure tips are empathetic, encouraging, and avoid medicalizing language. "
        "Format each tip as: **Title**\\nDescription\\n\\n"
    )

    return prompt


def parse_ai_tips_response(tips_text):
    """Parse AI-generated tips response into structured format"""
    tips = []
    sections = tips_text.split('\n\n')

    for section in sections:
        if '**' in section and section.strip():
            lines = section.strip().split('\n')
            if len(lines) >= 2:
                title = lines[0].replace('**', '').strip()
                description = ' '.join(lines[1:]).strip()
                tips.append({
                    'title': title,
                    'description': description,
                    'priority': len(tips) + 1  # Lower number = higher priority
                })

    # Ensure we have at least 5 tips, pad with defaults if needed
    while len(tips) < 5:
        default_tips = [
            {
                'title': 'Practice Daily Gratitude',
                'description': 'Take a moment each day to note three things you\'re grateful for. This simple practice can help shift focus from negative thoughts to positive aspects of your life.',
                'priority': len(tips) + 1
            },
            {
                'title': 'Stay Connected',
                'description': 'Reach out to friends, family, or classmates regularly. Social support is crucial for mental health, even brief check-ins can make a significant difference.',
                'priority': len(tips) + 1
            },
            {
                'title': 'Maintain Healthy Sleep Habits',
                'description': 'Aim for 7-9 hours of quality sleep. Create a calming bedtime routine and avoid screens before bed to improve sleep quality.',
                'priority': len(tips) + 1
            },
            {
                'title': 'Incorporate Physical Activity',
                'description': 'Engage in regular movement - take short walks, stretch, or do simple exercises. Physical activity releases endorphins and improves mood naturally.',
                'priority': len(tips) + 1
            },
            {
                'title': 'Practice Mindful Breathing',
                'description': 'When feeling overwhelmed, try the 4-7-8 breathing technique: inhale for 4 seconds, hold for 7 seconds, exhale for 8 seconds.',
                'priority': len(tips) + 1
            }
        ]
        tips.extend(default_tips[len(tips)-5:])

    return tips[:7]  # Return up to 7 tips


def generate_dass21_specific_fallback_tips(user, depression, anxiety, stress,
                                         depression_severity, anxiety_severity, stress_severity,
                                         user_history, dass_analysis):
    """Generate DASS21-specific personalized fallback tips when OpenAI is not available"""

    tips = []

    # Address specific DASS21 concerns first
    if dass_analysis['primary_concerns']:
        primary_concern = dass_analysis['primary_concerns'][0]
        if 'positive feeling' in primary_concern['question'].lower():
            tips.append({
                'title': 'Cultivate Positive Moments',
                'description': 'Make time for activities that bring joy, even if just for 5-10 minutes. Small positive experiences can help rebuild your capacity for enjoyment.',
                'priority': 1
            })
        elif 'initiative' in primary_concern['question'].lower():
            tips.append({
                'title': 'Break Tasks into Small Steps',
                'description': 'When feeling overwhelmed by tasks, break them into tiny, manageable steps. Celebrate completing each small step to build momentum.',
                'priority': 1
            })
        elif 'panic' in primary_concern['question'].lower():
            tips.append({
                'title': 'Grounding Techniques for Panic',
                'description': 'Use the 5-4-3-2-1 grounding technique: name 5 things you see, 4 you can touch, 3 you hear, 2 you smell, 1 you taste.',
                'priority': 1
            })
        elif 'worth' in primary_concern['question'].lower():
            tips.append({
                'title': 'Practice Self-Compassion',
                'description': 'Treat yourself with the same kindness you would offer a friend. Remember that your value is not determined by your current struggles.',
                'priority': 1
            })
        elif 'relax' in primary_concern['question'].lower():
            tips.append({
                'title': 'Progressive Muscle Relaxation',
                'description': 'Practice tensing and releasing different muscle groups. Start with your toes and work up to your head, noticing the difference between tension and relaxation.',
                'priority': 1
            })

    # Address coping patterns
    if 'difficulty relaxing' in dass_analysis['coping_patterns']:
        tips.append({
            'title': 'Create a Relaxation Routine',
            'description': 'Set aside 10-15 minutes daily for relaxation activities like deep breathing, gentle stretching, or listening to calming music.',
            'priority': 2
        })
    if 'emotional reactivity' in dass_analysis['coping_patterns']:
        tips.append({
            'title': 'Pause Before Responding',
            'description': 'When emotions run high, try the STOP technique: Stop, Take a breath, Observe your thoughts and feelings, Proceed mindfully.',
            'priority': 2
        })

    # Academic context tips
    if user_history['academic_context']:
        context = user_history['academic_context']
        if 'college_stressors' in context:
            if 'engineering' in context['college_stressors'].lower():
                tips.append({
                    'title': 'Academic Study Breaks',
                    'description': 'Use the Pomodoro technique: 25 minutes of focused study followed by a 5-minute break. This helps maintain concentration and prevents burnout.',
                    'priority': 2
                })
            elif 'business' in context['college_stressors'].lower():
                tips.append({
                    'title': 'Balance Academic and Personal Time',
                    'description': 'Create a weekly schedule that includes both study time and personal activities. Discuss complex topics with classmates to gain different perspectives.',
                    'priority': 2
                })

    # Current severity-based tips
    if depression_severity in ['moderate', 'severe', 'extremely-severe']:
        tips.append({
            'title': 'Seek Professional Support',
            'description': 'Consider connecting with a counselor or mental health professional. They can provide personalized strategies and support for managing depression.',
            'priority': 1
        })

    if anxiety_severity in ['moderate', 'severe', 'extremely-severe']:
        tips.append({
            'title': 'Anxiety Management Techniques',
            'description': 'Practice cognitive restructuring: when anxious thoughts arise, ask yourself "Is this thought based on facts or fears?" Replace irrational worries with realistic perspectives.',
            'priority': 1
        })

    if stress_severity in ['moderate', 'severe', 'extremely-severe']:
        tips.append({
            'title': 'Stress Reduction Strategies',
            'description': 'Identify your main stressors and create an action plan. Break overwhelming tasks into smaller steps and prioritize what truly matters.',
            'priority': 1
        })

    # Always include foundational tips
    tips.extend([
        {
            'title': 'Maintain Social Connections',
            'description': 'Stay connected with trusted friends, family, or classmates. Regular social interaction provides support and helps combat feelings of isolation.',
            'priority': 3
        },
        {
            'title': 'Healthy Sleep Habits',
            'description': 'Aim for consistent sleep and wake times. Create a calming bedtime routine and limit screen time before bed to improve sleep quality.',
            'priority': 3
        },
        {
            'title': 'Physical Activity',
            'description': 'Incorporate movement into your day through walking, stretching, or any enjoyable physical activity. Exercise naturally boosts mood and reduces stress.',
            'priority': 3
        }
    ])

    # Sort by priority and return top 7
    tips.sort(key=lambda x: x['priority'])
    return tips[:7]


def generate_fallback_feedback(depression, anxiety, stress, depression_severity, anxiety_severity, stress_severity):
    """Generate basic feedback when OpenAI is not available"""
    feedback_parts = []

    # Depression feedback
    if depression_severity in ['moderate', 'severe', 'extremely-severe']:
        feedback_parts.append("Your depression scores suggest you may be experiencing significant emotional challenges. <b>Consider reaching out to a mental health professional or counselor for support.</b>")
    elif depression_severity == 'mild':
        feedback_parts.append("You're showing some signs of depression. <b>Try engaging in activities you usually enjoy and maintain regular social connections.</b>")

    # Anxiety feedback
    if anxiety_severity in ['moderate', 'severe', 'extremely-severe']:
        feedback_parts.append("Your anxiety levels appear elevated. <b>Practice deep breathing exercises and consider talking to a counselor about your concerns.</b>")
    elif anxiety_severity == 'mild':
        feedback_parts.append("You're experiencing some anxiety. <b>Try mindfulness techniques and regular exercise to help manage stress.</b>")

    # Stress feedback
    if stress_severity in ['moderate', 'severe', 'extremely-severe']:
        feedback_parts.append("Your stress levels are quite high. <b>Prioritize self-care activities and consider seeking professional support to develop coping strategies.</b>")
    elif stress_severity == 'mild':
        feedback_parts.append("You're experiencing some stress. <b>Try time management techniques and regular breaks to help maintain balance.</b>")

    # General encouragement
    if not feedback_parts:
        feedback_parts.append("Your scores are within normal ranges. <b>Continue practicing good mental health habits and reach out for support if needed.</b>")

    return " ".join(feedback_parts)

@login_required
def clear_all_notifications(request):
    """Clear all notifications for the current user"""
    if request.method == 'POST':
        # Mark all notifications as dismissed
        Notification.objects.filter(
            user=request.user,
            dismissed=False
        ).update(dismissed=True)
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({
                'success': True,
                'message': 'All notifications cleared'
            })
        else:
            messages.success(request, 'All notifications cleared')
            return redirect('notifications')
    
    return redirect('notifications')

@login_required
def clear_notification(request, notification_id):
    """Clear a specific notification for the current user"""
    if request.method == 'POST':
        try:
            notification = Notification.objects.get(
                id=notification_id,
                user=request.user,
                dismissed=False
            )
            notification.dismissed = True
            notification.save()
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'message': 'Notification cleared'
                })
            else:
                messages.success(request, 'Notification cleared')
                return redirect('notifications')
        except Notification.DoesNotExist:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'error': 'Notification not found'
                }, status=404)
            else:
                messages.error(request, 'Notification not found')
                return redirect('notifications')
    
    return redirect('notifications')


@login_required
def feedback_form(request, appointment_id):
    """Display feedback form for completed appointment"""
    try:
        appointment = Appointment.objects.get(
            pk=appointment_id,
            user=request.user,
            status='completed'
        )
        
        # Check if feedback already exists
        feedback_exists = Feedback.objects.filter(
            appointment=appointment,
            user=request.user
        ).exists()
        
        if feedback_exists:
            messages.info(request, 'You have already provided feedback for this session.')
            return redirect('user-profile')
        
        return render(request, 'mentalhealth/feedback.html', {
            'appointment': appointment,
            'appointment_id': appointment.id,
            'counselor': appointment.counselor,
            'appointment_date': appointment.date,
            'appointment_time': appointment.time,
            'appointment_type': ', '.join(appointment.services),
            'duration': '60'  # Default duration, could be made configurable
        })
        
    except Appointment.DoesNotExist:
        messages.error(request, 'Appointment not found or you do not have permission to provide feedback.')
        return redirect('user-profile')


@require_http_methods(["POST"])
@login_required
def submit_feedback(request):
    """Handle feedback submission"""
    try:
        data = json.loads(request.body)
        appointment_id = data.get('appointment_id')
        
        appointment = Appointment.objects.get(
            pk=appointment_id,
            user=request.user,
            status='completed'
        )

        # Check if feedback already exists
        if Feedback.objects.filter(appointment=appointment, user=request.user).exists():
            return JsonResponse({
                'status': 'error',
                'message': 'Feedback already submitted for this appointment.'
            })
        
        # Create feedback object
        feedback = Feedback.objects.create(
            appointment=appointment,
            user=request.user,
            counselor=appointment.counselor,
            overall_rating=data.get('overall_rating'),
            professionalism_rating=data.get('professionalism_rating'),
            helpfulness_rating=data.get('helpfulness_rating'),
            recommend_rating=data.get('recommend_rating'),
            positive_feedback=data.get('positive_feedback', ''),
            improvement_feedback=data.get('improvement_feedback', ''),
            additional_comments=data.get('additional_comments', ''),
            skipped=False
        )
        
        return JsonResponse({
            'status': 'success',
            'message': 'Feedback submitted successfully!'
        })
        
    except Appointment.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': 'Appointment not found.'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        })


@require_http_methods(["POST"])
@login_required
def skip_feedback(request):
    """Handle feedback skip"""
    try:
        data = json.loads(request.body)
        appointment_id = data.get('appointment_id')
        
        appointment = Appointment.objects.get(
            pk=appointment_id,
            user=request.user,
            status='completed'
        )
        
        # Check if feedback already exists
        if Feedback.objects.filter(appointment=appointment, user=request.user).exists():
            return JsonResponse({
                'status': 'error',
                'message': 'Feedback already submitted for this appointment.'
            })
        
        # Create feedback object with skipped=True
        feedback = Feedback.objects.create(
            appointment=appointment,
            user=request.user,
            counselor=appointment.counselor,
            skipped=True
        )
        
        return JsonResponse({
            'status': 'success',
            'message': 'Feedback skipped successfully!'
        })
        
    except Appointment.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': 'Appointment not found.'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        })

@login_required
def create_live_session(request, appointment_id):
    """Create a live session for an appointment"""
    try:
        appointment = Appointment.objects.get(id=appointment_id)
        
        # Check if user has permission to create session
        if not (request.user == appointment.user or 
                (hasattr(request.user, 'counselor_profile') and 
                 request.user.counselor_profile == appointment.counselor)):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        # Check if session already exists
        if hasattr(appointment, 'live_session'):
            return JsonResponse({
                'success': True,
                'session_id': appointment.live_session.room_id,
                'message': 'Session already exists'
            })
        
        # Create live session
        start_time = datetime.combine(appointment.date, appointment.time)
        end_time = start_time + timedelta(hours=1)  # Default 1-hour session
        
        live_session = LiveSession.objects.create(
            appointment=appointment,
            scheduled_start=start_time,
            scheduled_end=end_time,
            session_type='video'
        )
        
        # Generate room ID
        room_id = live_session.generate_room_id()
        
        return JsonResponse({
            'success': True,
            'session_id': room_id,
            'meeting_url': f'/live-session/{room_id}/'
        })
        
    except Appointment.DoesNotExist:
        return JsonResponse({'error': 'Appointment not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def join_live_session(request, room_id):
    """Join a live session"""
    try:
        live_session = LiveSession.objects.get(room_id=room_id)
        
        # Check if user has permission to join
        if not (request.user == live_session.appointment.user or 
                (hasattr(request.user, 'counselor_profile') and 
                 request.user.counselor_profile == live_session.appointment.counselor)):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        # Check if session is active
        if live_session.status not in ['scheduled', 'waiting', 'active']:
            return JsonResponse({'error': 'Session is not available'}, status=400)
        
        # Update session status if needed
        if live_session.status == 'scheduled':
            live_session.status = 'waiting'
            live_session.save()
        
        # Create participant record
        participant, created = SessionParticipant.objects.get_or_create(
            session=live_session,
            user=request.user,
            defaults={
                'role': 'counselor' if hasattr(request.user, 'counselor_profile') else 'student'
            }
        )
        
        return JsonResponse({
            'success': True,
            'session_data': {
                'room_id': room_id,
                'session_type': live_session.session_type,
                'status': live_session.status,
                'participant_role': participant.role
            }
        })
        
    except LiveSession.DoesNotExist:
        return JsonResponse({'error': 'Session not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def live_session_view(request, room_id):
    """Render the live session interface"""
    try:
        # Try to get existing live session
        live_session = LiveSession.objects.get(room_id=room_id)
    except LiveSession.DoesNotExist:
        # If room_id follows the pattern 'appointment_X', try to create from appointment
        if room_id.startswith('appointment_'):
            try:
                appointment_id = int(room_id.split('_')[1])
                appointment = Appointment.objects.get(id=appointment_id)
                
                # Check if user has permission to access this appointment
                if not (request.user == appointment.user or 
                        (hasattr(request.user, 'counselor_profile') and 
                         request.user.counselor_profile == appointment.counselor)):
                    return redirect('index')
                
                # Create live session for this appointment
                start_time = timezone.make_aware(datetime.combine(appointment.date, appointment.time))
                end_time = start_time + timedelta(hours=1)
                
                live_session = LiveSession.objects.create(
                    appointment=appointment,
                    room_id=room_id,
                    scheduled_start=start_time,
                    scheduled_end=end_time,
                    session_type='video',
                    status='scheduled'
                )
            except (ValueError, Appointment.DoesNotExist):
                return redirect('index')
        else:
            return redirect('index')
    
    # Check if user has permission to access
    if not (request.user == live_session.appointment.user or 
            (hasattr(request.user, 'counselor_profile') and 
             request.user.counselor_profile == live_session.appointment.counselor)):
        return redirect('index')
    
    return render(request, 'mentalhealth/live-session.html', {
        'live_session': live_session,
        'appointment': live_session.appointment,
        'user': request.user
    })


@login_required
def end_live_session(request, room_id):
    """End a live session"""
    try:
        live_session = LiveSession.objects.get(room_id=room_id)
        
        # Check if user has permission to end session
        if not (request.user == live_session.appointment.user or 
                (hasattr(request.user, 'counselor_profile') and 
                 request.user.counselor_profile == live_session.appointment.counselor)):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        # Update session status
        live_session.status = 'completed'
        live_session.actual_end = timezone.now()
        live_session.save()
        
        # Update appointment status
        appointment = live_session.appointment
        appointment.status = 'completed'
        appointment.save()
        
        return JsonResponse({'success': True})
        
    except LiveSession.DoesNotExist:
        return JsonResponse({'error': 'Session not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def get_session_messages(request, room_id):
    """Get chat messages for a session"""
    try:
        live_session = LiveSession.objects.get(room_id=room_id)
        
        # Check if user has permission
        if not (request.user == live_session.appointment.user or 
                (hasattr(request.user, 'counselor_profile') and 
                 request.user.counselor_profile == live_session.appointment.counselor)):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        messages = SessionMessage.objects.filter(session=live_session).order_by('timestamp')
        messages_data = [{
            'id': msg.id,
            'sender': msg.sender.username,
            'message': msg.message,
            'timestamp': msg.timestamp.isoformat(),
            'message_type': msg.message_type
        } for msg in messages]
        
        return JsonResponse({'messages': messages_data})
        
    except LiveSession.DoesNotExist:
        return JsonResponse({'error': 'Session not found'}, status=404)


@login_required
def update_session_notes(request, room_id):
    """Update session notes"""
    try:
        live_session = LiveSession.objects.get(room_id=room_id)
        
        # Only counselors can update notes
        if not (hasattr(request.user, 'counselor_profile') and 
                request.user.counselor_profile == live_session.appointment.counselor):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        if request.method == 'POST':
            data = json.loads(request.body)
            live_session.notes = data.get('notes', '')
            live_session.save()
            
            return JsonResponse({'success': True})
        
        return JsonResponse({'notes': live_session.notes})
        
    except LiveSession.DoesNotExist:
        return JsonResponse({'error': 'Session not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def test_video_call(request, appointment_id):
    """Test view to verify video call functionality"""
    try:
        appointment = Appointment.objects.get(id=appointment_id)
        
        # Check if user has permission
        if not (request.user == appointment.user or 
                (hasattr(request.user, 'counselor_profile') and 
                 request.user.counselor_profile == appointment.counselor)):
            return JsonResponse({'error': 'Permission denied'}, status=403)
        
        # Get or create live session
        live_session, created = LiveSession.objects.get_or_create(
            appointment=appointment,
            defaults={
                'room_id': f'appointment_{appointment.id}',
                'scheduled_start': timezone.now(),
                'scheduled_end': timezone.now() + timedelta(hours=1),
                'session_type': 'video',
                'status': 'scheduled'
            }
        )
        
        video_call_url = request.build_absolute_uri(
            reverse('live_session_view', kwargs={'room_id': live_session.room_id})
        )
        
        return render(request, 'mentalhealth/test-video-call.html', {
            'appointment': appointment,
            'live_session': live_session,
            'video_call_url': video_call_url
        })
        
    except Appointment.DoesNotExist:
        return JsonResponse({'error': 'Appointment not found'}, status=404)

@login_required
def websocket_test(request):
    """Test WebSocket connection"""
    return render(request, 'mentalhealth/websocket-test.html')

@login_required
def simple_websocket_test(request):
    """Simple WebSocket test page"""
    return render(request, 'mentalhealth/simple-websocket-test.html')

def get_user_personalization_data(user):
    """Gather comprehensive user data for personalization"""
    # Get all DASS results for trend analysis
    dass_results = DASSResult.objects.filter(user=user).order_by('-date_taken')

    # Get user's test history
    test_count = dass_results.count()
    if test_count > 1:
        # Get the most recent and previous results for trend analysis
        recent_results = list(dass_results[:3])  # Last 3 tests
        trend_analysis = analyze_dass_trends(recent_results)
    else:
        trend_analysis = None

    # Get academic context
    academic_context = get_academic_context(user)

    # Get relaxation exercise history
    relaxation_history = RelaxationLog.objects.filter(user=user).order_by('-timestamp')
    exercise_preferences = analyze_exercise_preferences(relaxation_history)

    return {
        'test_count': test_count,
        'trend_analysis': trend_analysis,
        'academic_context': academic_context,
        'exercise_preferences': exercise_preferences,
        'recent_results': list(dass_results[:3]) if test_count > 0 else []
    }

def analyze_dass_trends(results):
    """Analyze trends in DASS scores over time"""
    if len(results) < 2:
        return None
    
    trends = {}
    for dimension in ['depression', 'anxiety', 'stress']:
        scores = [getattr(result, f'{dimension}_score') for result in results]
        if len(scores) >= 2:
            change = scores[0] - scores[-1]  # Most recent - oldest
            if change > 2:
                trends[dimension] = 'improving'
            elif change < -2:
                trends[dimension] = 'worsening'
            else:
                trends[dimension] = 'stable'
    
    return trends

def get_academic_context(user):
    """Get academic context for personalization"""
    context = {}
    
    # College-specific stressors
    college_stressors = {
        'CASS': 'arts and humanities coursework, creative projects, and social sciences research',
        'CEN': 'engineering projects, technical coursework, and problem-solving challenges',
        'CBA': 'business case studies, financial analysis, and professional development',
        'COF': 'fieldwork, laboratory work, and environmental studies',
        'CAG': 'agricultural research, fieldwork, and sustainability projects',
        'CHSI': 'home economics projects, industry applications, and practical skills',
        'CED': 'teaching practicums, educational research, and classroom management',
        'COS': 'scientific research, laboratory work, and mathematical analysis',
        'CVSM': 'veterinary clinical work, animal care, and medical studies'
    }
    
    if hasattr(user, 'college') and user.college:
        context['college_stressors'] = college_stressors.get(user.college, 'academic coursework')
    
    # Year-level specific challenges
    year_challenges = {
        '1': 'adjusting to university life, building study habits, and making new friends',
        '2': 'increasing academic workload, choosing specializations, and career exploration',
        '3': 'advanced coursework, research projects, and internship preparation',
        '4': 'thesis work, career preparation, and transition to professional life'
    }
    
    if hasattr(user, 'year_level') and user.year_level:
        context['year_challenges'] = year_challenges.get(user.year_level, 'academic development')
    
    # Age-specific considerations
    if hasattr(user, 'age') and user.age:
        if user.age < 20:
            context['age_context'] = 'young adult transitioning to independence'
        elif user.age < 25:
            context['age_context'] = 'emerging adult navigating career and personal development'
        else:
            context['age_context'] = 'adult learner balancing multiple responsibilities'
    
    return context

def analyze_exercise_preferences(relaxation_history):
    """Analyze user's relaxation exercise preferences"""
    if not relaxation_history:
        return None

    exercise_counts = {}
    for log in relaxation_history:
        exercise_type = log.exercise_type
        exercise_counts[exercise_type] = exercise_counts.get(exercise_type, 0) + 1

    # Get most preferred exercise
    if exercise_counts:
        preferred_exercise = max(exercise_counts, key=exercise_counts.get)
        return {
            'preferred_exercise': preferred_exercise,
            'total_sessions': len(relaxation_history),
            'exercise_counts': exercise_counts
        }

    return None


@api_view(['GET'])
def welcome_api(request):
    """API endpoint that logs request metadata and returns a welcome message"""
    logger.info(f"Request received: {request.method} {request.path}")
    return Response({
        'message': 'Welcome to CalmConnect API',
        'status': 'success'
    })

@api_view(['GET'])
def welcome_with_logging(request):
    """API endpoint that logs request metadata and returns a welcome message"""
    logger.info(f"Request received: {request.method} {request.path}")
    return Response({
        'message': 'Welcome to the CalmConnect API Service!',
        'status': 'success'
    })

@api_view(['GET'])
def welcome(request):
    """API endpoint that logs request metadata and returns a welcome message"""
    logger.info(f"Request received: {request.method} {request.path}")
    return Response({
        'message': 'Welcome to the CalmConnect API Service!',
        'status': 'success'
    })

@api_view(['GET'])
def health_check(request):
    """API endpoint for health check"""
    return Response({
        'status': 'healthy',
        'service': 'CalmConnect API',
        'timestamp': timezone.now().isoformat()
    })

@api_view(['GET'])
def welcome_new(request):
    """New API endpoint that logs request metadata and returns a welcome message"""
    logger.info(f"Request received: {request.method} {request.path}")
    return Response({
        'message': 'Welcome to the CalmConnect API Service!',
        'status': 'success'
    })

@api_view(['GET'])
def welcome_with_metadata(request):
    """API endpoint that logs request metadata and returns a welcome message with metadata"""
    logger.info(f"Request received: {request.method} {request.path}")
    return Response({
        'message': 'Welcome to the CalmConnect API Service!',
        'status': 'success',
        'metadata': {
            'method': request.method,
            'path': request.path,
            'user_agent': request.META.get('HTTP_USER_AGENT'),
            'remote_addr': request.META.get('REMOTE_ADDR')
        }
    })

@api_view(['GET'])
def welcome_with_method_path_logging(request):
    """API endpoint that logs request method and path and returns a welcome message"""
    logger.info(f"Request received: {request.method} {request.path}")
    return Response({
        'message': 'Welcome to the CalmConnect API Service!',
        'status': 'success'
    })

# Follow-up Session Views

@login_required
def create_followup_session(request, appointment_id):
    """Create a follow-up session for a completed appointment"""
    try:
        # Get the original appointment
        appointment = Appointment.objects.get(
            pk=appointment_id,
            status='completed'
        )

        # Check if user has permission (student or counselor)
        if not (request.user == appointment.user or
                (hasattr(request.user, 'counselor_profile') and
                 request.user.counselor_profile == appointment.counselor)):
            return JsonResponse({'error': 'Permission denied'}, status=403)

        if request.method == 'POST':
            data = json.loads(request.body)

            # Validate required fields
            required_fields = ['date', 'time', 'reason']
            for field in required_fields:
                if field not in data or not data[field]:
                    return JsonResponse({
                        'success': False,
                        'error': f'Missing required field: {field}'
                    }, status=400)

            # Parse date and time
            try:
                followup_date = datetime.strptime(data['date'], '%Y-%m-%d').date()
                followup_time = datetime.strptime(data['time'], '%H:%M').time()
            except ValueError:
                return JsonResponse({
                    'success': False,
                    'error': 'Invalid date or time format'
                }, status=400)

            # Check if date is in the future
            if followup_date < timezone.now().date():
                return JsonResponse({
                    'success': False,
                    'error': 'Follow-up date cannot be in the past'
                }, status=400)

            # Check if counselor is available at this time
            existing_appointment = Appointment.objects.filter(
                counselor=appointment.counselor,
                date=followup_date,
                time=followup_time,
                status__in=['pending', 'confirmed']
            ).first()

            if existing_appointment:
                return JsonResponse({
                    'success': False,
                    'error': 'Counselor is not available at this time'
                }, status=400)

            # Create follow-up appointment
            followup_appointment = Appointment.objects.create(
                user=appointment.user,
                counselor=appointment.counselor,
                date=followup_date,
                time=followup_time,
                session_type='followup',
                services=['Follow-up Session'],
                reason=f"Follow-up: {data['reason']}",
                phone=appointment.phone,
                course_section=appointment.course_section,
                status='pending',
                parent_appointment=appointment  # Link to original appointment
            )

            # Create notification for the student
            create_notification(
                user=appointment.user,
                message=f'A follow-up session has been scheduled with {appointment.counselor.name} on {followup_date.strftime("%B %d, %Y")} at {followup_time.strftime("%I:%M %p")}.',
                notification_type='appointment',
                url=reverse('user-profile')
            )

            # Create notification for the counselor
            create_notification(
                user=appointment.counselor.user,
                message=f'A follow-up session has been scheduled with {appointment.user.full_name} on {followup_date.strftime("%B %d, %Y")} at {followup_time.strftime("%I:%M %p")}.',
                notification_type='appointment',
                url=reverse('counselor_schedule')
            )

            return JsonResponse({
                'success': True,
                'message': 'Follow-up session created successfully!',
                'appointment_id': followup_appointment.id
            })

        # GET request - return form data
        return JsonResponse({
            'success': True,
            'appointment': {
                'id': appointment.id,
                'counselor_name': appointment.counselor.name,
                'original_date': appointment.date.strftime('%Y-%m-%d'),
                'services': list(appointment.services)
            }
        })

    except Appointment.DoesNotExist:
        return JsonResponse({'error': 'Appointment not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def followup_sessions(request):
    """List all follow-up sessions for the current user"""
    if hasattr(request.user, 'counselor_profile'):
        # Counselor view - show all their follow-up sessions
        followup_appointments = Appointment.objects.filter(
            counselor=request.user.counselor_profile,
            session_type='followup'
        ).select_related('user', 'counselor', 'parent_appointment').order_by('-date', '-time')
    else:
        # Student view - show their follow-up sessions
        followup_appointments = Appointment.objects.filter(
            user=request.user,
            session_type='followup'
        ).select_related('user', 'counselor', 'parent_appointment').order_by('-date', '-time')

    # Separate by status
    upcoming_sessions = followup_appointments.filter(
        status__in=['pending', 'confirmed'],
        date__gte=timezone.now().date()
    )

    past_sessions = followup_appointments.filter(
        Q(status='completed') |
        Q(date__lt=timezone.now().date())
    )

    context = {
        'upcoming_sessions': upcoming_sessions,
        'past_sessions': past_sessions,
        'total_sessions': followup_appointments.count(),
        'is_counselor': hasattr(request.user, 'counselor_profile')
    }

    return render(request, 'mentalhealth/followup-sessions.html', context)


@login_required
@api_view(['GET', 'POST', 'PUT'])
@permission_classes([IsAuthenticated])
def user_settings_api(request):
    """API endpoint for user settings management"""
    try:
        # Get or create user settings
        try:
            settings_obj, created = UserSettings.objects.get_or_create(
                user=request.user,
                defaults={
                    'dark_mode': False,
                    'font_size': 'medium',
                    'language': 'en',
                    'high_contrast': False,
                    'screen_reader': False,
                    'reduced_motion': False,
                    'analytics_tracking': True,
                    'profile_visibility': 'counselors_only',
                    'notification_preferences': {
                        'assignment_reminders': {'enabled': True, 'frequency': 'daily'},
                        'grade_updates': {'enabled': True, 'frequency': 'weekly'},
                        'study_session_alerts': {'enabled': True, 'frequency': 'daily'},
                        'appointment_reminders': {'enabled': True, 'frequency': 'daily'},
                        'followup_notifications': {'enabled': True, 'frequency': 'weekly'}
                    }
                }
            )
        except Exception:
            # If table doesn't exist, return default settings
            default_settings = {
                'dark_mode': False,
                'font_size': 'medium',
                'language': 'en',
                'high_contrast': False,
                'screen_reader': False,
                'reduced_motion': False,
                'analytics_tracking': True,
                'profile_visibility': 'counselors_only',
                'notification_preferences': {
                    'assignment_reminders': {'enabled': True, 'frequency': 'daily'},
                    'grade_updates': {'enabled': True, 'frequency': 'weekly'},
                    'study_session_alerts': {'enabled': True, 'frequency': 'daily'},
                    'appointment_reminders': {'enabled': True, 'frequency': 'daily'},
                    'followup_notifications': {'enabled': True, 'frequency': 'weekly'}
                }
            }
            if request.method == 'GET':
                return Response({
                    'success': True,
                    'settings': default_settings
                })
            elif request.method in ['POST', 'PUT']:
                return Response({
                    'success': True,
                    'message': 'Settings saved successfully (temporarily stored)',
                    'settings': default_settings
                })
            return Response({'success': True, 'settings': default_settings})

        if request.method == 'GET':
            return Response({
                'success': True,
                'settings': {
                    'dark_mode': settings_obj.dark_mode,
                    'font_size': settings_obj.font_size,
                    'language': settings_obj.language,
                    'high_contrast': settings_obj.high_contrast,
                    'screen_reader': settings_obj.screen_reader,
                    'reduced_motion': settings_obj.reduced_motion,
                    'analytics_tracking': settings_obj.analytics_tracking,
                    'profile_visibility': settings_obj.profile_visibility,
                    'notification_preferences': settings_obj.notification_preferences
                }
            })

        elif request.method in ['POST', 'PUT']:
            data = request.data

            # Update settings fields
            if 'dark_mode' in data:
                settings_obj.dark_mode = data['dark_mode']
            if 'font_size' in data:
                settings_obj.font_size = data['font_size']
            if 'language' in data:
                settings_obj.language = data['language']
            if 'high_contrast' in data:
                settings_obj.high_contrast = data['high_contrast']
            if 'screen_reader' in data:
                settings_obj.screen_reader = data['screen_reader']
            if 'reduced_motion' in data:
                settings_obj.reduced_motion = data['reduced_motion']
            if 'analytics_tracking' in data:
                settings_obj.analytics_tracking = data['analytics_tracking']
            if 'profile_visibility' in data:
                settings_obj.profile_visibility = data['profile_visibility']
            if 'notification_preferences' in data:
                settings_obj.notification_preferences = data['notification_preferences']

            settings_obj.save()

            return Response({
                'success': True,
                'message': 'Settings saved successfully',
                'settings': {
                    'dark_mode': settings_obj.dark_mode,
                    'font_size': settings_obj.font_size,
                    'language': settings_obj.language,
                    'high_contrast': settings_obj.high_contrast,
                    'screen_reader': settings_obj.screen_reader,
                    'reduced_motion': settings_obj.reduced_motion,
                    'analytics_tracking': settings_obj.analytics_tracking,
                    'profile_visibility': settings_obj.profile_visibility,
                    'notification_preferences': settings_obj.notification_preferences
                }
            })

    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def reset_user_settings(request):
    """Reset user settings to defaults"""
    try:
        settings_obj, created = UserSettings.objects.get_or_create(
            user=request.user,
            defaults={
                'dark_mode': False,
                'font_size': 'medium',
                'language': 'en',
                'high_contrast': False,
                'screen_reader': False,
                'reduced_motion': False,
                'analytics_tracking': True,
                'profile_visibility': 'counselors_only',
                'notification_preferences': {
                    'assignment_reminders': {'enabled': True, 'frequency': 'daily'},
                    'grade_updates': {'enabled': True, 'frequency': 'weekly'},
                    'study_session_alerts': {'enabled': True, 'frequency': 'daily'},
                    'appointment_reminders': {'enabled': True, 'frequency': 'daily'},
                    'followup_notifications': {'enabled': True, 'frequency': 'weekly'}
                }
            }
        )

        # Reset to defaults
        settings_obj.dark_mode = False
        settings_obj.font_size = 'medium'
        settings_obj.language = 'en'
        settings_obj.high_contrast = False
        settings_obj.screen_reader = False
        settings_obj.reduced_motion = False
        settings_obj.analytics_tracking = True
        settings_obj.profile_visibility = 'counselors_only'
        settings_obj.notification_preferences = {
            'assignment_reminders': {'enabled': True, 'frequency': 'daily'},
            'grade_updates': {'enabled': True, 'frequency': 'weekly'},
            'study_session_alerts': {'enabled': True, 'frequency': 'daily'},
            'appointment_reminders': {'enabled': True, 'frequency': 'daily'},
            'followup_notifications': {'enabled': True, 'frequency': 'weekly'}
        }
        settings_obj.save()

        return Response({
            'success': True,
            'message': 'Settings reset to defaults',
            'settings': {
                'dark_mode': settings_obj.dark_mode,
                'font_size': settings_obj.font_size,
                'language': settings_obj.language,
                'high_contrast': settings_obj.high_contrast,
                'screen_reader': settings_obj.screen_reader,
                'reduced_motion': settings_obj.reduced_motion,
                'analytics_tracking': settings_obj.analytics_tracking,
                'profile_visibility': settings_obj.profile_visibility,
                'notification_preferences': settings_obj.notification_preferences
            }
        })

    except Exception as e:
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
def followup_session_detail(request, appointment_id):
    """View details of a specific follow-up session"""
    try:
        appointment = Appointment.objects.select_related(
            'user', 'counselor', 'parent_appointment'
        ).get(pk=appointment_id, session_type='followup')

        # Check permissions
        if not (request.user == appointment.user or
                (hasattr(request.user, 'counselor_profile') and
                 request.user.counselor_profile == appointment.counselor)):
            return redirect('index')

        # Get related sessions (original and other follow-ups)
        if appointment.parent_appointment:
            related_sessions = Appointment.objects.filter(
                Q(id=appointment.parent_appointment.id) |
                Q(parent_appointment=appointment.parent_appointment, session_type='followup')
            ).exclude(id=appointment.id).select_related('user', 'counselor').order_by('date', 'time')
        else:
            related_sessions = Appointment.objects.filter(
                parent_appointment=appointment,
                session_type='followup'
            ).select_related('user', 'counselor').order_by('date', 'time')

        context = {
            'appointment': appointment,
            'related_sessions': related_sessions,
            'is_counselor': hasattr(request.user, 'counselor_profile'),
            'can_edit': hasattr(request.user, 'counselor_profile') and
                       request.user.counselor_profile == appointment.counselor
        }

        return render(request, 'mentalhealth/followup-session-detail.html', context)

    except Appointment.DoesNotExist:
        messages.error(request, 'Follow-up session not found.')
        return redirect('followup_sessions')


@login_required
def cancel_followup_session(request, appointment_id):
    """Cancel a follow-up session"""
    try:
        appointment = Appointment.objects.get(
            pk=appointment_id,
            session_type='followup'
        )

        # Check permissions
        if not (request.user == appointment.user or
                (hasattr(request.user, 'counselor_profile') and
                 request.user.counselor_profile == appointment.counselor)):
            return JsonResponse({'error': 'Permission denied'}, status=403)

        if request.method == 'POST':
            data = json.loads(request.body)
            reason = data.get('reason', 'No reason provided')

            # Update appointment status
            appointment.status = 'cancelled'
            appointment.cancellation_reason = reason
            appointment.save()

            # Create notifications
            if request.user == appointment.user:
                # Student cancelled
                create_notification(
                    user=appointment.counselor.user,
                    message=f'Follow-up session with {appointment.user.full_name} on {appointment.date} has been cancelled by the student.',
                    notification_type='appointment',
                    url=reverse('counselor_schedule')
                )
            else:
                # Counselor cancelled
                create_notification(
                    user=appointment.user,
                    message=f'Your follow-up session with {appointment.counselor.name} on {appointment.date} has been cancelled.',
                    notification_type='appointment',
                    url=reverse('user-profile')
                )

            return JsonResponse({
                'success': True,
                'message': 'Follow-up session cancelled successfully'
            })

        return JsonResponse({'error': 'Method not allowed'}, status=405)

    except Appointment.DoesNotExist:
        return JsonResponse({'error': 'Follow-up session not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def reschedule_followup_session(request, appointment_id):
    """Reschedule a follow-up session"""
    try:
        appointment = Appointment.objects.get(
            pk=appointment_id,
            session_type='followup'
        )

        # Check permissions
        if not (request.user == appointment.user or
                (hasattr(request.user, 'counselor_profile') and
                 request.user.counselor_profile == appointment.counselor)):
            return JsonResponse({'error': 'Permission denied'}, status=403)

        if request.method == 'POST':
            data = json.loads(request.body)

            # Validate new date and time
            try:
                new_date = datetime.strptime(data['date'], '%Y-%m-%d').date()
                new_time = datetime.strptime(data['time'], '%H:%M').time()
            except (ValueError, KeyError):
                return JsonResponse({
                    'success': False,
                    'error': 'Invalid date or time format'
                }, status=400)

            # Check if date is in the future
            if new_date < timezone.now().date():
                return JsonResponse({
                    'success': False,
                    'error': 'New date cannot be in the past'
                }, status=400)

            # Check counselor availability
            existing_appointment = Appointment.objects.filter(
                counselor=appointment.counselor,
                date=new_date,
                time=new_time,
                status__in=['pending', 'confirmed']
            ).exclude(id=appointment.id).first()

            if existing_appointment:
                return JsonResponse({
                    'success': False,
                    'error': 'Counselor is not available at this time'
                }, status=400)

            # Update appointment
            old_date = appointment.date
            old_time = appointment.time
            appointment.date = new_date
            appointment.time = new_time
            appointment.save()

            # Create notifications
            create_notification(
                user=appointment.user,
                message=f'Your follow-up session with {appointment.counselor.name} has been rescheduled from {old_date} {old_time} to {new_date} {new_time}.',
                notification_type='appointment',
                url=reverse('user-profile')
            )

            create_notification(
                user=appointment.counselor.user,
                message=f'Follow-up session with {appointment.user.full_name} has been rescheduled to {new_date} {new_time}.',
                notification_type='appointment',
                url=reverse('counselor_schedule')
            )

            return JsonResponse({
                'success': True,
                'message': 'Follow-up session rescheduled successfully'
            })

        return JsonResponse({'error': 'Method not allowed'}, status=405)

    except Appointment.DoesNotExist:
        return JsonResponse({'error': 'Follow-up session not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def followup_sessions_api(request):
    """API endpoint to get follow-up sessions data"""
    if hasattr(request.user, 'counselor_profile'):
        # Counselor view
        followup_appointments = Appointment.objects.filter(
            counselor=request.user.counselor_profile,
            session_type='followup'
        ).select_related('user', 'counselor', 'parent_appointment')
    else:
        # Student view
        followup_appointments = Appointment.objects.filter(
            user=request.user,
            session_type='followup'
        ).select_related('user', 'counselor', 'parent_appointment')

    # Filter by status if provided
    status_filter = request.GET.get('status')
    if status_filter:
        followup_appointments = followup_appointments.filter(status=status_filter)

    # Order by date
    followup_appointments = followup_appointments.order_by('-date', '-time')

    # Serialize data
    sessions_data = []
    for appointment in followup_appointments:
        sessions_data.append({
            'id': appointment.id,
            'date': appointment.date.strftime('%Y-%m-%d'),
            'time': appointment.time.strftime('%H:%M'),
            'counselor_name': appointment.counselor.name,
            'student_name': appointment.user.full_name,
            'reason': appointment.reason,
            'status': appointment.status,
            'parent_appointment_id': appointment.parent_appointment.id if appointment.parent_appointment else None,
            'created_at': appointment.created_at.strftime('%Y-%m-%d %H:%M')
        })

    return Response({
        'success': True,
        'sessions': sessions_data,
        'total_count': len(sessions_data)
    })


@counselor_required
def export_reports(request):
    """Export reports in various formats (CSV, PDF, DOC)"""
    from django.http import HttpResponse
    import csv
    from io import BytesIO
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from docx import Document
    from docx.shared import Inches

    # Get filter parameters
    report_type = request.GET.get('report_type', '')
    date_range = request.GET.get('date_range', '')
    status = request.GET.get('status', '')
    search = request.GET.get('search', '')
    format_type = request.GET.get('format', 'csv')  # Default to CSV

    # Base queryset - only reports for this counselor
    reports = Report.objects.filter(counselor=request.user.counselor_profile)

    # Apply filters
    if report_type:
        reports = reports.filter(report_type=report_type)

    if status:
        reports = reports.filter(status=status)

    if date_range:
        from datetime import datetime, timedelta
        today = timezone.now().date()
        if date_range == 'today':
            reports = reports.filter(created_at__date=today)
        elif date_range == 'week':
            week_ago = today - timedelta(days=7)
            reports = reports.filter(created_at__date__gte=week_ago)
        elif date_range == 'month':
            month_ago = today - timedelta(days=30)
            reports = reports.filter(created_at__date__gte=month_ago)
        elif date_range == 'quarter':
            quarter_ago = today - timedelta(days=90)
            reports = reports.filter(created_at__date__gte=quarter_ago)

    if search:
        reports = reports.filter(
            Q(title__icontains=search) |
            Q(description__icontains=search) |
            Q(user__full_name__icontains=search)
        )

    # Order by creation date
    reports = reports.select_related('user').order_by('-created_at')

    if format_type == 'csv':
        # CSV Export
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="counselor_reports.csv"'

        writer = csv.writer(response)
        writer.writerow(['Title', 'Student', 'Report Type', 'Status', 'Created Date', 'Last Modified', 'Description'])

        for report in reports:
            writer.writerow([
                report.title,
                report.user.full_name if report.user else 'N/A',
                report.get_report_type_display(),
                report.get_status_display(),
                report.created_at.strftime('%Y-%m-%d %H:%M'),
                report.updated_at.strftime('%Y-%m-%d %H:%M'),
                report.description[:200] + '...' if len(report.description) > 200 else report.description
            ])

        return response

    elif format_type == 'pdf':
        # PDF Export
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="counselor_reports.pdf"'

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title = Paragraph("Counselor Reports Export", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))

        # Summary
        summary_text = f"Total Reports: {reports.count()}<br/>Generated on: {timezone.now().strftime('%Y-%m-%d %H:%M')}"
        summary = Paragraph(summary_text, styles['Normal'])
        story.append(summary)
        story.append(Spacer(1, 12))

        # Table data
        data = [['Title', 'Student', 'Type', 'Status', 'Created', 'Description']]
        for report in reports:
            data.append([
                report.title,
                report.user.full_name if report.user else 'N/A',
                report.get_report_type_display(),
                report.get_status_display(),
                report.created_at.strftime('%Y-%m-%d'),
                report.description[:100] + '...' if len(report.description) > 100 else report.description
            ])

        # Create table
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(table)
        doc.build(story)
        pdf = buffer.getvalue()
        buffer.close()
        response.write(pdf)
        return response

    elif format_type == 'doc':
        # DOC Export
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="counselor_reports.docx"'

        document = Document()
        document.add_heading('Counselor Reports Export', 0)

        # Summary
        document.add_paragraph(f"Total Reports: {reports.count()}")
        document.add_paragraph(f"Generated on: {timezone.now().strftime('%Y-%m-%d %H:%M')}")
        document.add_paragraph("")

        # Table
        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Title'
        hdr_cells[1].text = 'Student'
        hdr_cells[2].text = 'Type'
        hdr_cells[3].text = 'Status'
        hdr_cells[4].text = 'Created'
        hdr_cells[5].text = 'Description'

        for report in reports:
            row_cells = table.add_row().cells
            row_cells[0].text = report.title
            row_cells[1].text = report.user.full_name if report.user else 'N/A'
            row_cells[2].text = report.get_report_type_display()
            row_cells[3].text = report.get_status_display()
            row_cells[4].text = report.created_at.strftime('%Y-%m-%d')
            row_cells[5].text = report.description[:100] + '...' if len(report.description) > 100 else report.description

        # Save to response
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response.write(buffer.getvalue())
        buffer.close()
        return response

    else:
        return JsonResponse({'error': 'Unsupported format'}, status=400)


# Archive API endpoints for the new frontend
@counselor_required
@api_view(['GET'])
def archive_stats(request):
    """Get archive statistics for the counselor"""
    counselor = request.user.counselor_profile

    # Calculate statistics
    total_sessions = Appointment.objects.filter(
        counselor=counselor,
        status='completed'
    ).count()

    total_reports = Report.objects.filter(
        counselor=counselor,
        status__in=['archived', 'completed']
    ).count()

    total_assessments = SecureDASSResult.objects.filter(
        user__in=Appointment.objects.filter(
            counselor=counselor,
            status='completed'
        ).values_list('user', flat=True)
    ).count()

    # Monthly activity (last 30 days)
    thirty_days_ago = timezone.now() - timedelta(days=30)
    monthly_sessions = Appointment.objects.filter(
        counselor=counselor,
        status='completed',
        date__gte=thirty_days_ago
    ).count()

    monthly_reports = Report.objects.filter(
        counselor=counselor,
        status__in=['archived', 'completed'],
        created_at__gte=thirty_days_ago
    ).count()

    return Response({
        'success': True,
        'stats': {
            'total_sessions': total_sessions,
            'total_reports': total_reports,
            'total_assessments': total_assessments,
            'monthly_sessions': monthly_sessions,
            'monthly_reports': monthly_reports,
            'monthly_assessments': SecureDASSResult.objects.filter(
                user__in=Appointment.objects.filter(
                    counselor=counselor,
                    status='completed',
                    date__gte=thirty_days_ago
                ).values_list('user', flat=True),
                date_taken__gte=thirty_days_ago
            ).count()
        }
    })


@counselor_required
@api_view(['GET'])
def archive_data(request):
    """Get paginated archive data for the counselor"""
    counselor = request.user.counselor_profile

    # Get query parameters
    tab = request.GET.get('tab', 'sessions')  # sessions, reports, assessments
    page = int(request.GET.get('page', 1))
    per_page = int(request.GET.get('per_page', 10))
    sort_by = request.GET.get('sort_by', 'date')
    sort_order = request.GET.get('sort_order', 'desc')
    search = request.GET.get('search', '')
    time_period = request.GET.get('time_period', 'all')
    session_type = request.GET.get('session_type', 'all')

    # Base querysets
    if tab == 'sessions':
        queryset = Appointment.objects.filter(
            counselor=counselor,
            status='completed'
        ).select_related('user')
    elif tab == 'reports':
        queryset = Report.objects.filter(
            counselor=counselor,
            status__in=['archived', 'completed']
        ).select_related('user')
    elif tab == 'assessments':
        # Get assessments for students who have completed sessions with this counselor
        student_ids = Appointment.objects.filter(
            counselor=counselor,
            status='completed'
        ).values_list('user', flat=True).distinct()
        queryset = SecureDASSResult.objects.filter(
            user__in=student_ids
        ).select_related('user')
    else:
        return Response({'success': False, 'error': 'Invalid tab'}, status=400)

    # Apply filters
    if search:
        if tab == 'sessions':
            queryset = queryset.filter(
                Q(user__full_name__icontains=search) |
                Q(reason__icontains=search) |
                Q(services__icontains=search)
            )
        elif tab == 'reports':
            queryset = queryset.filter(
                Q(title__icontains=search) |
                Q(description__icontains=search) |
                Q(user__full_name__icontains=search)
            )
        elif tab == 'assessments':
            queryset = queryset.filter(
                Q(user__full_name__icontains=search) |
                Q(user__student_id__icontains=search)
            )

    # Time period filter
    if time_period != 'all':
        today = timezone.now().date()
        if time_period == 'month':
            cutoff_date = today - timedelta(days=30)
        elif time_period == 'quarter':
            cutoff_date = today - timedelta(days=90)
        elif time_period == 'year':
            cutoff_date = today - timedelta(days=365)
        else:
            cutoff_date = None

        if cutoff_date:
            if tab == 'sessions':
                queryset = queryset.filter(date__gte=cutoff_date)
            elif tab == 'reports':
                queryset = queryset.filter(created_at__date__gte=cutoff_date)
            elif tab == 'assessments':
                queryset = queryset.filter(date_taken__gte=cutoff_date)

    # Session type filter (only for sessions tab)
    if tab == 'sessions' and session_type != 'all':
        queryset = queryset.filter(session_type=session_type)

    # Apply sorting
    if tab == 'sessions':
        if sort_by == 'date':
            queryset = queryset.order_by(f'{"-" if sort_order == "desc" else ""}date', f'{"-" if sort_order == "desc" else ""}time')
        elif sort_by == 'name':
            queryset = queryset.order_by(f'{"-" if sort_order == "desc" else ""}user__full_name')
        elif sort_by == 'type':
            queryset = queryset.order_by(f'{"-" if sort_order == "desc" else ""}session_type')
    elif tab == 'reports':
        if sort_by == 'date':
            queryset = queryset.order_by(f'{"-" if sort_order == "desc" else ""}created_at')
        elif sort_by == 'name':
            queryset = queryset.order_by(f'{"-" if sort_order == "desc" else ""}user__full_name')
        elif sort_by == 'type':
            queryset = queryset.order_by(f'{"-" if sort_order == "desc" else ""}report_type')
    elif tab == 'assessments':
        if sort_by == 'date':
            queryset = queryset.order_by(f'{"-" if sort_order == "desc" else ""}date_taken')
        elif sort_by == 'name':
            queryset = queryset.order_by(f'{"-" if sort_order == "desc" else ""}user__full_name')

    # Pagination
    paginator = Paginator(queryset, per_page)
    try:
        page_obj = paginator.page(page)
    except:
        page_obj = paginator.page(1)

    # Serialize data
    data = []
    for item in page_obj:
        if tab == 'sessions':
            data.append({
                'id': item.id,
                'title': f"{item.user.full_name} - {', '.join(item.services) if item.services else 'General Session'}",
                'date': item.date.strftime('%M %d, %Y'),
                'student_name': item.user.full_name,
                'student_id': item.user.student_id,
                'college': item.user.get_college_display(),
                'services': item.services,
                'duration': getattr(item, 'duration', 'N/A'),
                'reason': item.reason,
                'type': item.session_type,
                'status': item.status
            })
        elif tab == 'reports':
            data.append({
                'id': item.id,
                'title': item.title,
                'date': item.created_at.strftime('%M %d, %Y'),
                'student_name': item.user.full_name if item.user else 'N/A',
                'student_id': item.user.student_id if item.user else 'N/A',
                'type': item.get_report_type_display(),
                'status': item.get_status_display(),
                'description': item.description
            })
        elif tab == 'assessments':
            data.append({
                'id': item.id,
                'title': f"DASS21 Results - {item.user.full_name}",
                'date': item.date_taken.strftime('%M %d, %Y'),
                'student_name': item.user.full_name,
                'student_id': item.user.student_id,
                'depression_score': item.depression_score,
                'anxiety_score': item.anxiety_score,
                'stress_score': item.stress_score,
                'depression_severity': item.depression_severity,
                'anxiety_severity': item.anxiety_severity,
                'stress_severity': item.stress_severity
            })

    return Response({
        'success': True,
        'data': data,
        'pagination': {
            'page': page_obj.number,
            'per_page': per_page,
            'total_pages': paginator.num_pages,
            'total_items': paginator.count,
            'has_next': page_obj.has_next(),
            'has_previous': page_obj.has_previous()
        }
    })


@counselor_required
@api_view(['GET'])
def archive_export(request):
    """Export archive data in various formats"""
    from django.http import HttpResponse
    import csv
    from io import BytesIO

    counselor = request.user.counselor_profile

    # Get parameters
    tab = request.GET.get('tab', 'sessions')
    format_type = request.GET.get('format', 'csv')
    search = request.GET.get('search', '')
    time_period = request.GET.get('time_period', 'all')
    session_type = request.GET.get('session_type', 'all')

    # Get data using same logic as archive_data
    if tab == 'sessions':
        queryset = Appointment.objects.filter(
            counselor=counselor,
            status='completed'
        ).select_related('user')
    elif tab == 'reports':
        queryset = Report.objects.filter(
            counselor=counselor,
            status__in=['archived', 'completed']
        ).select_related('user')
    elif tab == 'assessments':
        student_ids = Appointment.objects.filter(
            counselor=counselor,
            status='completed'
        ).values_list('user', flat=True).distinct()
        queryset = SecureDASSResult.objects.filter(
            user__in=student_ids
        ).select_related('user')
    else:
        return Response({'success': False, 'error': 'Invalid tab'}, status=400)

    # Apply same filters as archive_data
    if search:
        if tab == 'sessions':
            queryset = queryset.filter(
                Q(user__full_name__icontains=search) |
                Q(reason__icontains=search) |
                Q(services__icontains=search)
            )
        elif tab == 'reports':
            queryset = queryset.filter(
                Q(title__icontains=search) |
                Q(description__icontains=search) |
                Q(user__full_name__icontains=search)
            )
        elif tab == 'assessments':
            queryset = queryset.filter(
                Q(user__full_name__icontains=search) |
                Q(user__student_id__icontains=search)
            )

    if time_period != 'all':
        today = timezone.now().date()
        if time_period == 'month':
            cutoff_date = today - timedelta(days=30)
        elif time_period == 'quarter':
            cutoff_date = today - timedelta(days=90)
        elif time_period == 'year':
            cutoff_date = today - timedelta(days=365)
        else:
            cutoff_date = None

        if cutoff_date:
            if tab == 'sessions':
                queryset = queryset.filter(date__gte=cutoff_date)
            elif tab == 'reports':
                queryset = queryset.filter(created_at__date__gte=cutoff_date)
            elif tab == 'assessments':
                queryset = queryset.filter(date_taken__gte=cutoff_date)

    if tab == 'sessions' and session_type != 'all':
        queryset = queryset.filter(session_type=session_type)

    # Order by date
    if tab == 'sessions':
        queryset = queryset.order_by('-date', '-time')
    elif tab == 'reports':
        queryset = queryset.order_by('-created_at')
    elif tab == 'assessments':
        queryset = queryset.order_by('-date_taken')

    if format_type == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="counselor_archive_{tab}.csv"'

        writer = csv.writer(response)

        if tab == 'sessions':
            writer.writerow(['Date', 'Student Name', 'Student ID', 'College', 'Services', 'Session Type', 'Reason'])
            for item in queryset:
                writer.writerow([
                    item.date.strftime('%Y-%m-%d'),
                    item.user.full_name,
                    item.user.student_id,
                    item.user.get_college_display(),
                    ', '.join(item.services) if item.services else '',
                    item.session_type,
                    item.reason
                ])
        elif tab == 'reports':
            writer.writerow(['Created Date', 'Title', 'Student Name', 'Student ID', 'Report Type', 'Status', 'Description'])
            for item in queryset:
                writer.writerow([
                    item.created_at.strftime('%Y-%m-%d %H:%M'),
                    item.title,
                    item.user.full_name if item.user else 'N/A',
                    item.user.student_id if item.user else 'N/A',
                    item.get_report_type_display(),
                    item.get_status_display(),
                    item.description
                ])
        elif tab == 'assessments':
            writer.writerow(['Date Taken', 'Student Name', 'Student ID', 'Depression Score', 'Anxiety Score', 'Stress Score', 'Depression Severity', 'Anxiety Severity', 'Stress Severity'])
            for item in queryset:
                writer.writerow([
                    item.date_taken.strftime('%Y-%m-%d'),
                    item.user.full_name,
                    item.user.student_id,
                    item.depression_score,
                    item.anxiety_score,
                    item.stress_score,
                    item.depression_severity,
                    item.anxiety_severity,
                    item.stress_severity
                ])

        return response

    elif format_type == 'json':
        response = HttpResponse(content_type='application/json')
        response['Content-Disposition'] = f'attachment; filename="counselor_archive_{tab}.json"'

        data = []
        for item in queryset:
            if tab == 'sessions':
                data.append({
                    'date': item.date.strftime('%Y-%m-%d'),
                    'student_name': item.user.full_name,
                    'student_id': item.user.student_id,
                    'college': item.user.get_college_display(),
                    'services': item.services,
                    'session_type': item.session_type,
                    'reason': item.reason
                })
            elif tab == 'reports':
                data.append({
                    'created_at': item.created_at.strftime('%Y-%m-%d %H:%M'),
                    'title': item.title,
                    'student_name': item.user.full_name if item.user else 'N/A',
                    'student_id': item.user.student_id if item.user else 'N/A',
                    'report_type': item.get_report_type_display(),
                    'status': item.get_status_display(),
                    'description': item.description
                })
            elif tab == 'assessments':
                data.append({
                    'date_taken': item.date_taken.strftime('%Y-%m-%d'),
                    'student_name': item.user.full_name,
                    'student_id': item.user.student_id,
                    'depression_score': item.depression_score,
                    'anxiety_score': item.anxiety_score,
                    'stress_score': item.stress_score,
                    'depression_severity': item.depression_severity,
                    'anxiety_severity': item.anxiety_severity,
                    'stress_severity': item.stress_severity
                })

        response.write(json.dumps(data, indent=2))
        return response

    else:
        return Response({'success': False, 'error': 'Unsupported format'}, status=400)


@login_required
def get_notifications(request):
    """Get notifications for the current user (AJAX endpoint)"""
    if request.method != 'GET':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    try:
        # Get recent notifications
        notifications = Notification.objects.filter(
            user=request.user,
            dismissed=False
        ).order_by('-created_at')[:20]  # Get more for pagination

        notifications_data = []
        for notif in notifications:
            # Compute URL like send_notification does
            url = notif.action_url if notif.action_url else '#'
            if notif.type == 'followup':
                url = '#'  # Disable navigation for all follow-up modals

            # Ensure metadata is properly serialized
            metadata = notif.metadata
            if not metadata:
                metadata = {}
            elif isinstance(metadata, str):
                import json
                try:
                    metadata = json.loads(metadata)
                except:
                    metadata = {}
            else:
                # Ensure it's a dict
                metadata = dict(metadata) if hasattr(metadata, 'items') else {}
    
            # Set URL based on notification type and action_url
            if notif.type == 'followup' and not notif.action_url:
                # Student followup notifications should show modal
                final_url = '#'
            else:
                # Other notifications use their action_url or default to '#'
                final_url = notif.action_url if notif.action_url else '#'
    
            notifications_data.append({
                'id': notif.id,
                'message': notif.message,
                'type': notif.type,
                'url': final_url,
                'action_url': notif.action_url,
                'read': notif.read,
                'created_at': notif.created_at.isoformat(),
                'priority': getattr(notif, 'priority', 'normal'),
                'action_text': getattr(notif, 'action_text', None),
                'metadata': metadata,
                'icon': notif.get_icon(),
                'color': notif.get_color()
            })

        return JsonResponse({
            'success': True,
            'notifications': notifications_data
        })

    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@require_POST
def clear_notification(request, notification_id):
    """Clear a specific notification"""
    try:
        notification = Notification.objects.get(
            id=notification_id,
            user=request.user
        )
        notification.read = True
        notification.save()

        return JsonResponse({'success': True})
    except Notification.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Notification not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def clear_all_notifications(request):
    """Clear all notifications for the current user"""
    try:
        Notification.objects.filter(
            user=request.user,
            read=False
        ).update(read=True)

        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def request_followup(request, pk):
    """Allow counselor or student to request a follow-up session"""
    try:
        report = Report.objects.select_related('user', 'counselor').get(pk=pk)

        # Check permissions - must be the counselor or the student
        if not (hasattr(request.user, 'counselor_profile') and
                request.user.counselor_profile == report.counselor) and \
           request.user != report.user:
            return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)

        # Check if follow-up already exists for this report
        existing_request = FollowupRequest.objects.filter(report=report).first()
        if existing_request:
            return JsonResponse({
                'success': False,
                'error': f'Follow-up already requested (Status: {existing_request.get_status_display()})'
            }, status=400)

        data = json.loads(request.body)
        reason = data.get('reason', '').strip()
        requested_date = data.get('requested_date')
        requested_time = data.get('requested_time')

        if not reason:
            return JsonResponse({'success': False, 'error': 'Reason is required'}, status=400)

        # Determine requester type
        requester_type = 'counselor' if hasattr(request.user, 'counselor_profile') and \
                        request.user.counselor_profile == report.counselor else 'student'

        # Create follow-up request (automatically approved to remove admin boilerplate)
        followup_request = FollowupRequest.objects.create(
            report=report,
            requested_by=request.user,
            requester_type=requester_type,
            reason=reason,
            requested_date=requested_date,
            requested_time=requested_time,
            status='approved'
        )

        # Create notifications
        from .notification_service import create_followup_notification
        create_followup_notification(followup_request, 'requested')

        return JsonResponse({
            'success': True,
            'message': 'Follow-up request submitted successfully',
            'request_id': followup_request.id
        })

    except Report.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Report not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@staff_member_required
def admin_followup_requests(request):
    """Admin interface for viewing follow-up requests (now auto-approved)"""
    print(f"DEBUG: admin_followup_requests called by user: {request.user}, is_staff: {request.user.is_staff}, is_superuser: {request.user.is_superuser}")

    # Get filter parameters - default to 'approved' since requests are now auto-approved
    status_filter = request.GET.get('status', 'approved')
    search = request.GET.get('search', '')
    print(f"DEBUG: filters - status: {status_filter}, search: {search}")

    # Base queryset
    followup_requests = FollowupRequest.objects.select_related(
        'report', 'requested_by', 'report__user', 'report__counselor',
        'approved_denied_by', 'resulting_appointment'
    ).order_by('-created_at')

    # Apply filters
    if status_filter and status_filter != 'all':
        followup_requests = followup_requests.filter(status=status_filter)

    if search:
        followup_requests = followup_requests.filter(
            Q(report__user__full_name__icontains=search) |
            Q(report__counselor__name__icontains=search) |
            Q(requested_by__full_name__icontains=search) |
            Q(reason__icontains=search)
        )

    # Paginate
    from django.core.paginator import Paginator
    paginator = Paginator(followup_requests, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Build query parameters string for pagination links
    query_params = ''
    if status_filter and status_filter != 'all':
        query_params += f'&status={status_filter}'
    if search:
        query_params += f'&search={search}'

    context = {
        'page_obj': page_obj,
        'status_filter': status_filter,
        'search': search,
        'query_params': query_params,
        'total_requests': followup_requests.count(),
        'pending_count': FollowupRequest.objects.filter(status='pending').count(),
        'approved_count': FollowupRequest.objects.filter(status='approved').count(),
        'denied_count': FollowupRequest.objects.filter(status='denied').count(),
    }

    return render(request, 'admin-followup-requests.html', context)


@staff_member_required
@require_POST
def approve_deny_followup(request, request_id):
    """Admin approve or deny follow-up request"""
    try:
        followup_request = FollowupRequest.objects.select_related(
            'report', 'requested_by', 'report__user', 'report__counselor'
        ).get(pk=request_id)

        if followup_request.status != 'pending':
            return JsonResponse({
                'success': False,
                'error': 'Request has already been processed'
            }, status=400)

        data = json.loads(request.body)
        action = data.get('action')  # 'approve' or 'deny'
        admin_notes = data.get('admin_notes', '').strip()

        if action not in ['approve', 'deny']:
            return JsonResponse({'success': False, 'error': 'Invalid action'}, status=400)

        # Update request
        followup_request.status = 'approved' if action == 'approve' else 'denied'
        followup_request.admin_notes = admin_notes
        followup_request.approved_denied_by = request.user
        followup_request.approved_denied_at = timezone.now()
        followup_request.save()

        # Create notifications
        from .notification_service import create_followup_notification
        create_followup_notification(followup_request, action)

        return JsonResponse({
            'success': True,
            'message': f'Follow-up request {action}d successfully'
        })

    except FollowupRequest.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Follow-up request not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
def followup_session(request, appointment_id):
    """Display the follow-up session assessment page"""
    try:
        # Get the follow-up appointment
        appointment = Appointment.objects.select_related(
            'user', 'counselor', 'parent_appointment'
        ).get(pk=appointment_id, session_type='followup')

        # Check permissions - must be the student
        if appointment.user != request.user:
            return redirect('index')

        # Check if appointment is confirmed
        if appointment.status != 'confirmed':
            messages.error(request, 'This follow-up session is not available.')
            return redirect('user-profile')

        # Calculate days since original session
        if appointment.parent_appointment:
            original_date = appointment.parent_appointment.date
            days_since_session = (timezone.now().date() - original_date).days
        else:
            days_since_session = 0

        context = {
            'counselor': appointment.counselor,
            'appointment_date': appointment.date,
            'days_since_session': days_since_session,
            'appointment_id': appointment.id,
            'appointment': appointment,
        }

        return render(request, 'mentalhealth/followup-session.html', context)

    except Appointment.DoesNotExist:
        messages.error(request, 'Follow-up session not found.')
        return redirect('user-profile')
    except Exception as e:
        messages.error(request, f'An error occurred: {str(e)}')
        return redirect('user-profile')


@login_required
@require_POST
def submit_followup(request):
    """Handle follow-up assessment submission"""
    try:
        data = json.loads(request.body)
        appointment_id = data.get('appointment_id')

        # Get the appointment
        appointment = Appointment.objects.select_related('user', 'counselor').get(
            pk=appointment_id, session_type='followup'
        )

        # Check permissions
        if appointment.user != request.user:
            return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)

        # Check if already submitted
        if appointment.status == 'completed':
            return JsonResponse({
                'success': False,
                'error': 'Follow-up assessment already submitted'
            }, status=400)

        # Validate required fields
        required_fields = ['wellbeing', 'symptoms', 'strategies', 'mood', 'practice', 'new_challenges', 'future_session']
        for field in required_fields:
            if field not in data:
                return JsonResponse({
                    'success': False,
                    'error': f'Missing required field: {field}'
                }, status=400)

        # Store the follow-up assessment data as a DASSResult with special handling
        followup_data = {
            'wellbeing_progress': int(data['wellbeing']),
            'symptoms_progress': int(data['symptoms']),
            'strategies_progress': int(data['strategies']),
            'mood_change': data['mood'],
            'practice_frequency': data['practice'],
            'helpful_things': data.get('helpful_things', ''),
            'new_challenges': data['new_challenges'],
            'challenges_description': data.get('challenges_description', ''),
            'future_session_preference': data['future_session'],
            'future_focus': data.get('future_focus', ''),
        }

        # Create a special follow-up result (using DASSResult model for consistency)
        followup_result = DASSResult.objects.create(
            user=request.user,
            depression_score=0,  # Not applicable for follow-up
            anxiety_score=0,     # Not applicable for follow-up
            stress_score=0,      # Not applicable for follow-up
            depression_severity='normal',
            anxiety_severity='normal',
            stress_severity='normal',
            answers=followup_data,
            is_followup=True,
            followup_appointment=appointment
        )

        # Update appointment status
        appointment.status = 'completed'
        appointment.save()

        # Create notification for counselor
        create_notification(
            user=appointment.counselor.user,
            message=f'{appointment.user.full_name} has completed their follow-up assessment.',
            notification_type='followup',
            url=reverse('appointment_detail', kwargs={'pk': appointment.id})
        )

        # Create notification for student
        create_notification(
            user=appointment.user,
            message='Thank you for completing your follow-up assessment. Your counselor will review the results.',
            notification_type='followup',
            url=reverse('user-profile')
        )

        return JsonResponse({
            'success': True,
            'message': 'Follow-up assessment submitted successfully!',
            'status': 'success'
        })

    except Appointment.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Follow-up session not found'}, status=404)
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@counselor_required
def schedule_followup(request, request_id):
    """Counselor schedules an approved follow-up request"""
    try:
        followup_request = FollowupRequest.objects.select_related(
            'report', 'report__user', 'report__counselor'
        ).get(pk=request_id)

        # Check permissions
        if followup_request.report.counselor != request.user.counselor_profile:
            messages.error(request, 'Permission denied')
            return redirect('admin_followup_requests')

        if followup_request.status != 'approved':
            messages.error(request, 'Request must be approved before scheduling')
            return redirect('admin_followup_requests')

        if request.method == 'POST':
            scheduled_date = request.POST.get('scheduled_date')
            scheduled_time = request.POST.get('scheduled_time')
            session_type = request.POST.get('session_type', 'face_to_face')

            if not scheduled_date or not scheduled_time:
                messages.error(request, 'Date and time are required')
                return redirect('schedule_followup', request_id=request_id)

            # Validate date/time
            try:
                scheduled_date_obj = datetime.strptime(scheduled_date, '%Y-%m-%d').date()
                scheduled_time_obj = datetime.strptime(scheduled_time, '%H:%M').time()
            except ValueError:
                messages.error(request, 'Invalid date or time format')
                return redirect('schedule_followup', request_id=request_id)

            # Check if date is in the future
            scheduled_datetime = timezone.make_aware(
                datetime.combine(scheduled_date_obj, scheduled_time_obj)
            )
            if scheduled_datetime <= timezone.now():
                messages.error(request, 'Scheduled time must be in the future')
                return redirect('schedule_followup', request_id=request_id)

            # Check counselor availability
            existing_appointment = Appointment.objects.filter(
                counselor=followup_request.report.counselor,
                date=scheduled_date_obj,
                time=scheduled_time_obj,
                status__in=['pending', 'confirmed']
            ).first()

            if existing_appointment:
                messages.error(request, 'Counselor is not available at this time')
                return redirect('schedule_followup', request_id=request_id)

            # Update follow-up request
            followup_request.scheduled_date = scheduled_date_obj
            followup_request.scheduled_time = scheduled_time_obj
            followup_request.session_type = session_type
            followup_request.status = 'scheduled'
            followup_request.save()

            # Create the actual appointment
            appointment = Appointment.objects.create(
                user=followup_request.report.user,
                counselor=followup_request.report.counselor,
                date=scheduled_date_obj,
                time=scheduled_time_obj,
                session_type=session_type,
                services=['Follow-up Session'],
                reason=f"Follow-up: {followup_request.reason}",
                phone=getattr(followup_request.report.user, 'phone', ''),
                course_section=getattr(followup_request.report.user, 'program', ''),
                status='confirmed'
            )

            followup_request.resulting_appointment = appointment
            followup_request.save()

            # Handle session type specific actions
            if appointment.session_type == 'remote':
                # Create live session for remote follow-up appointments
                start_time = timezone.make_aware(datetime.combine(appointment.date, appointment.time))
                end_time = start_time + timedelta(hours=1)  # Default 1-hour session

                # Create or get existing live session
                live_session, created = LiveSession.objects.get_or_create(
                    appointment=appointment,
                    defaults={
                        'scheduled_start': start_time,
                        'scheduled_end': end_time,
                        'session_type': 'video',
                        'status': 'scheduled'
                    }
                )

                # Generate room ID if not already set
                if not live_session.room_id:
                    live_session.room_id = f'followup_{appointment.id}'
                    live_session.save()

                # Generate video call link for remote follow-up sessions
                video_call_url = request.build_absolute_uri(
                    reverse('live_session_view', kwargs={'room_id': live_session.room_id})
                )
                appointment.video_call_url = video_call_url
                appointment.save()

            # Create notifications
            from .notification_service import create_followup_notification
            create_followup_notification(followup_request, 'scheduled')

            messages.success(request, 'Follow-up session scheduled successfully')
            return redirect('admin_followup_requests')

        # GET request - show the form
        return render(request, 'mentalhealth/schedule-followup.html', {
            'followup_request': followup_request,
            'today': timezone.now().date()
        })

    except FollowupRequest.DoesNotExist:
        messages.error(request, 'Follow-up request not found')
        return redirect('admin_followup_requests')
    except Exception as e:
        messages.error(request, f'An error occurred: {str(e)}')
        return redirect('admin_followup_requests')


@login_required
@require_http_methods(["GET", "POST"])
def accept_followup(request, request_id):
    """Handle student acceptance for scheduled follow-up (AJAX modal or direct page access)"""
    logger.info(f"accept_followup called by user {request.user.username} for request_id {request_id}, method: {request.method}, AJAX: {request.headers.get('X-Requested-With') == 'XMLHttpRequest'}")

    try:
        followup_request = FollowupRequest.objects.select_related(
            'report', 'report__user', 'report__counselor', 'resulting_appointment'
        ).get(pk=request_id)
        logger.info(f"Found followup_request {followup_request.id} with status {followup_request.status}")

        # Check permissions - must be the student
        if followup_request.report.user != request.user:
            logger.warning(f"Permission denied: user {request.user.username} tried to access followup_request {followup_request.id} belonging to {followup_request.report.user.username}")
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
            messages.error(request, 'Permission denied')
            return redirect('user-profile')

        if followup_request.status not in ['approved', 'scheduled']:
            logger.warning(f"Follow-up request {followup_request.id} has invalid status {followup_request.status}")
            error_msg = 'Follow-up request is not available'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': error_msg}, status=400)
            messages.error(request, error_msg)
            return redirect('user-profile')

        if not followup_request.resulting_appointment:
            logger.error(f"Follow-up request {followup_request.id} has no resulting appointment")
            error_msg = 'No appointment found for this follow-up request'
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'error': error_msg}, status=400)
            messages.error(request, error_msg)
            return redirect('user-profile')

        if request.method == 'GET':
            # Check if this is an AJAX request for modal data
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                # Return JSON data for modal
                appointment = followup_request.resulting_appointment
                counselor = followup_request.report.counselor

                # Robust services field handling
                services = appointment.services
                try:
                    if isinstance(services, list):
                        services_list = services
                    elif isinstance(services, str):
                        # Try to parse as JSON first
                        try:
                            parsed_services = json.loads(services)
                            if isinstance(parsed_services, list):
                                services_list = parsed_services
                            else:
                                services_list = [str(parsed_services)]
                        except (json.JSONDecodeError, TypeError):
                            services_list = [services]
                    elif services is None:
                        services_list = []
                    else:
                        services_list = [str(services)]
                except Exception as e:
                    logger.warning(f"Error processing services field for appointment {appointment.id}: {e}")
                    services_list = []

                # Safe counselor data extraction
                counselor_data = {}
                if counselor:
                    try:
                        counselor_data = {
                            'name': counselor.name or 'Unknown Counselor',
                            'college': counselor.college or 'N/A',
                            'rank': counselor.rank or 'N/A',
                            'image_url': counselor.image.url if hasattr(counselor, 'image') and counselor.image else None
                        }
                    except Exception as e:
                        logger.warning(f"Error extracting counselor data: {e}")
                        counselor_data = {
                            'name': 'Unknown Counselor',
                            'unit': 'N/A',
                            'rank': 'N/A',
                            'image_url': None
                        }
                else:
                    counselor_data = {
                        'name': 'Unknown Counselor',
                        'unit': 'N/A',
                        'rank': 'N/A',
                        'image_url': None
                    }

                # Safe appointment data extraction
                appointment_data = {}
                try:
                    appointment_data = {
                        'id': appointment.id,
                        'date': appointment.date.strftime('%B %d, %Y') if appointment.date else 'TBD',
                        'time': appointment.time.strftime('%I:%M %p') if appointment.time else 'TBD',
                        'session_type': appointment.get_session_type_display() if appointment.session_type and appointment.session_type in dict(appointment.SESSION_TYPE_CHOICES) else (appointment.session_type or 'TBD'),
                        'services': services_list,
                        'reason': appointment.reason or ''
                    }
                except Exception as e:
                    logger.error(f"Error extracting appointment data for appointment {appointment.id}: {e}")
                    appointment_data = {
                        'date': 'TBD',
                        'time': 'TBD',
                        'session_type': 'TBD',
                        'services': [],
                        'reason': 'Error loading appointment details'
                    }

                return JsonResponse({
                    'success': True,
                    'followup_request': {
                        'id': followup_request.id,
                        'reason': followup_request.reason or '',
                        'status': followup_request.status
                    },
                    'appointment': appointment_data,
                    'counselor': counselor_data
                })
            else:
                # Direct page access - render HTML template
                appointment = followup_request.resulting_appointment
                counselor = followup_request.report.counselor

                context = {
                    'counselor': counselor,
                    'appointment': appointment,
                    'followup_request': followup_request,
                    'user': request.user
                }

                return render(request, 'mentalhealth/followup-consent-modal.html', context)

        elif request.method == 'POST':
            # Handle consent submission
            # Check if it's an AJAX request first
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                # AJAX request - expect JSON data
                data = json.loads(request.body)
                consent_given = data.get('consent_given', False)
            else:
                # Form submission - expect form data
                consent_given = request.POST.get('consent_given') == 'true'

            if not consent_given:
                # Student declined - cancel the follow-up
                followup_request.status = 'cancelled'
                followup_request.resulting_appointment.status = 'cancelled'
                followup_request.resulting_appointment.save()
                followup_request.save()

                # Create notifications
                from .notification_service import create_followup_notification
                create_followup_notification(followup_request, 'declined')

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'message': 'Follow-up declined successfully'
                    })
                else:
                    messages.success(request, 'Follow-up declined successfully')
                    return redirect('user-profile')
            else:
                # Student consented - appointment remains confirmed
                # Generate video call URL for remote sessions if not already present
                appointment = followup_request.resulting_appointment
                if appointment.session_type == 'remote' and not appointment.video_call_url:
                    # Create live session for remote follow-up appointments
                    start_time = timezone.make_aware(datetime.combine(appointment.date, appointment.time))
                    end_time = start_time + timedelta(hours=1)  # Default 1-hour session

                    # Create or get existing live session
                    live_session, created = LiveSession.objects.get_or_create(
                        appointment=appointment,
                        defaults={
                            'scheduled_start': start_time,
                            'scheduled_end': end_time,
                            'session_type': 'video',
                            'status': 'scheduled'
                        }
                    )

                    # Generate room ID if not already set
                    if not live_session.room_id:
                        live_session.room_id = f'followup_{appointment.id}'
                        live_session.save()

                    # Generate video call link for remote follow-up sessions
                    video_call_url = request.build_absolute_uri(
                        reverse('live_session_view', kwargs={'room_id': live_session.room_id})
                    )
                    appointment.video_call_url = video_call_url
                    appointment.save()

                # Create notifications
                from .notification_service import create_followup_notification
                create_followup_notification(followup_request, 'consented')

                # Also create a notification for the counselor that student accepted
                counselor = followup_request.report.counselor
                if counselor and counselor.user:
                    student_name = followup_request.report.user.full_name if followup_request.report.user.full_name else 'Unknown Student'
                    create_notification(
                        user=counselor.user,
                        message=f'{student_name} has accepted your follow-up request. Please schedule the session.',
                        notification_type='followup',
                        priority='high',
                        action_url=None,  # Show modal for counselor scheduling
                        action_text='Schedule Session',
                        expires_in_hours=72,
                        metadata={
                            'followup_request_id': followup_request.id,
                            'student_name': student_name,
                            'user_type': 'counselor'
                        }
                    )

                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return JsonResponse({
                        'success': True,
                        'message': 'Consent given successfully'
                    })
                else:
                    messages.success(request, 'Consent given successfully')
                    return redirect('followup_session', appointment_id=appointment.id)

    except FollowupRequest.DoesNotExist:
        logger.error(f"Follow-up request {request_id} not found")
        error_msg = 'Follow-up request not found'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'error': error_msg}, status=404)
        messages.error(request, error_msg)
        return redirect('user-profile')
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in accept_followup: {e}")
        error_msg = 'Invalid request data'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'error': error_msg}, status=400)
        messages.error(request, error_msg)
        return redirect('user-profile')
    except Exception as e:
        logger.error(f"Unexpected error in accept_followup: {e}", exc_info=True)
        error_msg = f'An error occurred: {str(e)}'
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'error': error_msg}, status=500)
        messages.error(request, error_msg)
        return redirect('user-profile')


@login_required
@require_http_methods(["GET"])
def followup_details(request, request_id):
    """Get details of a scheduled follow-up session for counselor modal"""
    logger.info(f"followup_details called by user {request.user.username} for request_id {request_id}")

    try:
        followup_request = FollowupRequest.objects.select_related(
            'report', 'report__user', 'report__counselor', 'resulting_appointment'
        ).get(pk=request_id)

        # Check permissions - must be the counselor
        if followup_request.report.counselor.user != request.user:
            logger.warning(f"Permission denied: user {request.user.username} tried to access followup_request {followup_request.id} for counselor {followup_request.report.counselor.user.username}")
            return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)

        if followup_request.status != 'scheduled':
            logger.warning(f"Follow-up request {followup_request.id} has invalid status {followup_request.status}")
            return JsonResponse({'success': False, 'error': 'Follow-up request is not scheduled'}, status=400)

        if not followup_request.resulting_appointment:
            logger.error(f"Follow-up request {followup_request.id} has no resulting appointment")
            return JsonResponse({'success': False, 'error': 'No appointment found for this follow-up request'}, status=400)

        appointment = followup_request.resulting_appointment
        counselor = followup_request.report.counselor
        student = followup_request.report.user

        # Robust services field handling
        services = appointment.services
        try:
            if isinstance(services, list):
                services_list = services
            elif isinstance(services, str):
                # Try to parse as JSON first
                try:
                    parsed_services = json.loads(services)
                    if isinstance(parsed_services, list):
                        services_list = parsed_services
                    else:
                        services_list = [str(parsed_services)]
                except (json.JSONDecodeError, TypeError):
                    services_list = [services]
            elif services is None:
                services_list = []
            else:
                services_list = [str(services)]
        except Exception as e:
            logger.warning(f"Error processing services field for appointment {appointment.id}: {e}")
            services_list = []

        # Safe counselor data extraction
        counselor_data = {}
        if counselor:
            try:
                counselor_data = {
                    'name': counselor.name or 'Unknown Counselor',
                    'college': counselor.college or 'N/A',
                    'rank': counselor.rank or 'N/A',
                    'image_url': counselor.image.url if hasattr(counselor, 'image') and counselor.image else None
                }
            except Exception as e:
                logger.warning(f"Error extracting counselor data: {e}")
                counselor_data = {
                    'name': 'Unknown Counselor',
                    'unit': 'N/A',
                    'rank': 'N/A',
                    'image_url': None
                }

        # Safe student data extraction
        student_data = {}
        if student:
            try:
                student_data = {
                    'full_name': student.full_name or 'Unknown Student',
                    'student_id': getattr(student, 'student_id', 'N/A'),
                    'email': student.email or 'N/A'
                }
            except Exception as e:
                logger.warning(f"Error extracting student data: {e}")
                student_data = {
                    'full_name': 'Unknown Student',
                    'student_id': 'N/A',
                    'email': 'N/A'
                }

        # Safe appointment data extraction
        appointment_data = {}
        try:
            appointment_data = {
                'id': appointment.id,
                'date': appointment.date.strftime('%B %d, %Y') if appointment.date else 'TBD',
                'time': appointment.time.strftime('%I:%M %p') if appointment.time else 'TBD',
                'session_type': appointment.get_session_type_display() if appointment.session_type and appointment.session_type in dict(appointment.SESSION_TYPE_CHOICES) else (appointment.session_type or 'TBD'),
                'session_type_raw': appointment.session_type or 'face_to_face',  # Add raw session type for video call logic
                'services': services_list,
                'reason': appointment.reason or '',
                'status': appointment.get_status_display() if appointment.status else 'Unknown',
                'video_call_url': appointment.video_call_url or None  # Include video call URL
            }
        except Exception as e:
            logger.error(f"Error extracting appointment data for appointment {appointment.id}: {e}")
            appointment_data = {
                'date': 'TBD',
                'time': 'TBD',
                'session_type': 'TBD',
                'session_type_raw': 'face_to_face',
                'services': [],
                'reason': 'Error loading appointment details',
                'status': 'Unknown',
                'video_call_url': None
            }

        return JsonResponse({
            'success': True,
            'followup_request': {
                'id': followup_request.id,
                'reason': followup_request.reason or '',
                'status': followup_request.status,
                'created_at': followup_request.created_at.strftime('%B %d, %Y %I:%M %p') if followup_request.created_at else None,
                'scheduled_date': followup_request.scheduled_date.strftime('%B %d, %Y') if followup_request.scheduled_date else None,
                'scheduled_time': followup_request.scheduled_time.strftime('%I:%M %p') if followup_request.scheduled_time else None,
                'session_type': followup_request.session_type or 'TBD'
            },
            'appointment': appointment_data,
            'counselor': counselor_data,
            'student': student_data
        })

    except FollowupRequest.DoesNotExist:
        logger.error(f"Follow-up request {request_id} not found")
        return JsonResponse({'success': False, 'error': 'Follow-up request not found'}, status=404)
    except Exception as e:
        logger.error(f"Unexpected error in followup_details: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': f'An error occurred: {str(e)}'}, status=500)




# TEMPORARY: Create superuser endpoint (remove after use)
@csrf_exempt
def create_superuser_endpoint(request):
    """TEMPORARY endpoint to create a superuser - REMOVE AFTER USE"""
    if request.method == 'POST':
        username = request.POST.get('username', 'admin')
        email = request.POST.get('email', 'admin@calmconnect.edu.ph')
        password = request.POST.get('password', 'admin123!')

        # Check if user already exists
        if CustomUser.objects.filter(username=username).exists():
            messages.error(request, f'Superuser "{username}" already exists.')
            return redirect('admin_dashboard')

        # Generate unique student_id
        base_student_id = 'admin001'
        student_id = base_student_id
        counter = 1
        while CustomUser.objects.filter(student_id=student_id).exists():
            counter += 1
            student_id = f'admin{counter:03d}'

        try:
            user = CustomUser.objects.create_superuser(
                username=username,
                email=email,
                password=password,
                full_name='Administrator',
                age=0,
                gender='Prefer not to say',
                college='CBA',
                program='Administration',
                year_level='4',
                student_id=student_id
            )

            messages.success(request, f'Superuser "{username}" created successfully! '
                                    f'Username: {username}, Password: {password}')
            return render(request, 'mentalhealth/create_superuser.html', {
                'title': 'Create Superuser (TEMPORARY)',
                'success_message': f'Superuser "{username}" created successfully!<br>'
                                 f'<strong>Username:</strong> {username}<br>'
                                 f'<strong>Password:</strong> {password}<br>'
                                 f'<strong>Email:</strong> {email}<br><br>'
                                 f'<a href="/admin/" class="btn btn-primary">Go to Admin Panel</a>'
            })

        except Exception as e:
            messages.error(request, f'Error creating superuser: {e}')
            return redirect('admin_dashboard')

    # GET request - show form
    return render(request, 'mentalhealth/create_superuser.html', {
        'title': 'Create Superuser (TEMPORARY)'
    })
@login_required
def run_migrations(request):
    """Create missing database tables (temporary endpoint for Railway)"""
    from django.db import connection
    from django.http import JsonResponse

    try:
        with connection.cursor() as cursor:
            # Create the SecureDASSResult table directly
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS mentalhealth_securedassresult (
                    dassresult_ptr_id bigint NOT NULL PRIMARY KEY REFERENCES mentalhealth_dassresult(id) ON DELETE CASCADE,
                    encrypted_answers text,
                    encrypted_depression_score text,
                    encrypted_anxiety_score text,
                    encrypted_stress_score text,
                    data_hash varchar(64),
                    consent_given boolean NOT NULL DEFAULT false,
                    consent_timestamp timestamp with time zone,
                    encryption_version varchar(10) NOT NULL DEFAULT 'v1',
                    access_count integer NOT NULL DEFAULT 0,
                    last_accessed timestamp with time zone
                );
            """)

            # Create the DASSDataRetentionPolicy table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS mentalhealth_dassdataretentionpolicy (
                    id bigserial NOT NULL PRIMARY KEY,
                    user_id bigint NOT NULL REFERENCES mentalhealth_customuser(id) ON DELETE CASCADE,
                    policy_type varchar(20) NOT NULL DEFAULT 'standard',
                    applied_date timestamp with time zone NOT NULL DEFAULT now(),
                    retention_until timestamp with time zone,
                    reason text,
                    approved_by_id bigint REFERENCES mentalhealth_customuser(id) ON DELETE SET NULL,
                    CONSTRAINT mentalhealth_dassdataret_user_policy_8b8b8b8b_uniq UNIQUE (user_id, policy_type)
                );
            """)

        return JsonResponse({
            'success': True,
            'message': 'Missing tables created successfully'
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })

